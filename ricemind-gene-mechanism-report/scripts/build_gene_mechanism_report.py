#!/usr/bin/env python3
"""Build a fixed-template RiceMind gene mechanism DOCX report from full API data."""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import sys
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from urllib.parse import parse_qs, quote, urlencode, urljoin, urlparse
from urllib.request import Request, urlopen


DEFAULT_API_BASE = "http://lit-evi.hzau.edu.cn/ricemind-api/"
DEFAULT_PAGE_LIMIT = 500
DEFAULT_MAX_PAGES = 10000
REPORT_TRAIT_ROWS = 20
REPORT_API_ROWS = 6
REPORT_MECHANISM_THEMES = 8

FONT_EAST_ASIA = "SimSun"
FONT_LATIN = "Times New Roman"
FONT_SIZES = {
    "title": 16,
    "heading1": 14,
    "heading2": 12,
    "body": 10.5,
    "table": 9,
    "caption": 9,
}

DEFAULT_ENDPOINTS: Dict[str, Dict[str, Any]] = {
    "ricemind_get_gene_profile": {
        "method": "GET",
        "path": "gene-profile/",
        "params": {"gene": "{gene}"},
    },
    "ricemind_get_traits_by_gene": {
        "method": "GET",
        "path": "traits-by-gene/",
        "params": {
            "gene": "{gene}",
            "confidence": "{confidence}",
            "onto_type": "{onto_type}",
            "page": "{page}",
            "limit": "{limit}",
        },
        "record_keys": ["associated_traits"],
    },
    "ricemind_get_varieties_by_gene": {
        "method": "GET",
        "path": "varieties-by-gene/",
        "params": {"gene": "{gene}", "page": "{page}", "limit": "{limit}"},
        "record_keys": ["varieties"],
    },
    "ricemind_search_by_gene": {
        "method": "GET",
        "path": "search-by-gene/",
        "params": {"gene": "{gene}", "page": "{page}", "limit": "{limit}"},
        "record_keys": ["results"],
    },
    "ricemind_search_by_trait_and_gene": {
        "method": "GET",
        "path": "search-by-trait-and-gene/",
        "params": {"gene": "{gene}", "trait": "{trait}", "page": "{page}", "limit": "{limit}"},
        "record_keys": ["sentence_evidence"],
    },
    "ricemind_get_sentence_context": {
        "method": "GET",
        "path": "sentence-context/",
        "params": {"pmid": "{pmid}", "sent_id": "{sent_id}", "window": "{window}"},
    },
    "ricemind_get_gene_omics_sequence": {
        "method": "GET",
        "path": "gene-omics-sequence/",
        "params": {"rap_id": "{rap_id}"},
    },
}

FIELD_ALIASES = {
    "gene": ["Gene_Symbol", "gene_symbol", "geneSymbol", "gene", "symbol", "Original_Gene_ID", "canonical_symbol"],
    "rap_id": ["RAP_ID", "rap_id", "rapId", "RAP", "locus", "locus_id", "standardized_id", "standard_rap_id", "primary_id"],
    "trait": ["Trait_Description", "trait_description", "traitDescription", "trait", "trait_name", "name", "trait_query"],
    "ontology_type": ["Ontology_Type", "ontology_type", "ontologyType", "ontology"],
    "ontology_id": ["Ontology_ID", "ontology_id", "ontologyId", "term_id", "id"],
    "evidence_code": ["Evidence_Code", "evidence_code", "evidenceCode", "evidence_codes"],
    "source_db": ["Source_DB", "source_db", "sourceDb", "source", "sources", "source_databases", "database"],
    "confidence": ["confidence", "Confidence", "confidence_tier", "evidence_tier", "tier", "category"],
    "support": [
        "supporting_article_count",
        "supportingArticleCount",
        "support_count",
        "supportCount",
        "article_count",
        "literature_support_count",
        "pmid_count",
        "row_count",
    ],
    "year": ["year", "Year", "publication_year", "publicationYear", "earliest_year", "first_mention_year"],
    "pmid": ["PMID", "pmid", "pubmed_id", "pubmedId"],
    "sentence_id": ["Sentence_ID", "sentence_id", "sentenceId", "sent_id", "sentId"],
    "sentence": ["text", "sentence", "evidence_sentence", "evidenceText", "Evidence_Text"],
    "title": ["title", "Title", "article_title"],
    "journal": ["journal", "Journal"],
    "doi": ["doi", "DOI"],
}

HIGH_CODES = {"ORYZABASE_CURATED", "RAP-DB_CURATED", "EXP", "IDA", "IMP", "IEP", "HEP", "TAS"}
MEDIUM_NAMES = {"MEDIUM", "TIER2", "TIER 2", "TIER_2", "TIER2_HIGH_CONFIDENCE_NOVEL"}
LOW_NAMES = {"LOW", "TIER3", "TIER 3", "TIER_3", "TIER3_LIMITED_SUPPORT"}


def stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "; ".join(stringify(item) for item in value if item not in (None, ""))
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def get_value(record: Any, canonical: str, default: str = "") -> str:
    if not isinstance(record, dict):
        return default
    lower_map = {str(k).lower(): k for k in record}
    for alias in FIELD_ALIASES.get(canonical, []):
        key = lower_map.get(alias.lower())
        if key is not None and record.get(key) not in (None, ""):
            return stringify(record[key])
    return default


def deep_find(obj: Any, candidate_keys: Sequence[str]) -> str:
    candidates = {key.lower() for key in candidate_keys}
    if isinstance(obj, dict):
        for key, value in obj.items():
            if str(key).lower() in candidates and value not in (None, ""):
                return stringify(value)
        for value in obj.values():
            found = deep_find(value, candidate_keys)
            if found:
                return found
    elif isinstance(obj, list):
        for item in obj:
            found = deep_find(item, candidate_keys)
            if found:
                return found
    return ""


def flatten_records(obj: Any) -> List[Dict[str, Any]]:
    if obj is None:
        return []
    if isinstance(obj, list):
        rows: List[Dict[str, Any]] = []
        for item in obj:
            rows.extend(flatten_records(item))
        return rows
    if isinstance(obj, dict):
        candidate_keys = [
            "records",
            "associated_traits",
            "sentence_evidence",
            "gta_metadata",
            "results",
            "items",
            "data",
            "associations",
            "traits",
            "evidence",
            "sentences",
            "rows",
            "varieties",
            "transcripts",
        ]
        for key in candidate_keys:
            if key in obj and isinstance(obj[key], list):
                return flatten_records(obj[key])
        return [obj]
    return []


def read_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    if not isinstance(data, dict):
        return {"data": data}
    return data


def write_json(path: Path, payload: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)


def render_params(params: Dict[str, Any], values: Dict[str, Any]) -> Dict[str, str]:
    rendered: Dict[str, str] = {}
    for key, value in params.items():
        if value is None or value == "":
            continue
        text = stringify(value)
        for name, replacement in values.items():
            text = text.replace("{" + name + "}", stringify(replacement))
        if text != "":
            rendered[key] = text
    return rendered


def build_url(base: str, endpoint: Dict[str, Any], values: Dict[str, Any]) -> str:
    path = stringify(endpoint["path"])
    for name, replacement in values.items():
        path = path.replace("{" + name + "}", quote(stringify(replacement)))
    url = urljoin(base.rstrip("/") + "/", path.lstrip("/"))
    params = render_params(endpoint.get("params", {}), values)
    if params:
        url += "?" + urlencode(params)
    return url


def call_get(base: str, endpoint: Dict[str, Any], values: Dict[str, Any], api_calls: List[Dict[str, Any]]) -> Any:
    method = stringify(endpoint.get("method", "GET")).upper()
    if method != "GET":
        raise ValueError(f"Only GET endpoints are supported by this helper, got {method}")
    url = build_url(base, endpoint, values)
    record_api_call(api_calls, endpoint, values, url)
    request = Request(url, headers={"Accept": "application/json"})
    with urlopen(request, timeout=60) as resp:
        return json.loads(resp.read().decode("utf-8"))


def record_api_call(api_calls: List[Dict[str, Any]], endpoint: Dict[str, Any], values: Dict[str, Any], url: str) -> None:
    tool = stringify(endpoint.get("tool", ""))
    path = stringify(endpoint.get("path", ""))
    group_values = {k: stringify(v) for k, v in values.items() if k not in {"page", "limit"}}
    query_value = group_values.get("trait") if tool == "ricemind_search_by_trait_and_gene" else ""
    if tool == "ricemind_search_by_trait_and_gene":
        group_values["trait"] = "<multiple traits>"
    key = json.dumps({"tool": tool, "path": path, "values": group_values}, sort_keys=True, ensure_ascii=False)
    page = parse_int(values.get("page"))
    limit = parse_int(values.get("limit"))
    for entry in api_calls:
        if entry.get("key") == key:
            entry["call_count"] = parse_int(entry.get("call_count")) + 1
            if page:
                entry["page_min"] = min(parse_int(entry.get("page_min")) or page, page)
                entry["page_max"] = max(parse_int(entry.get("page_max")) or page, page)
            if limit:
                entry["limit"] = limit
            entry["last_url"] = url
            if query_value:
                examples = entry.setdefault("query_examples", [])
                if query_value not in examples and len(examples) < 8:
                    examples.append(query_value)
                entry["query_count"] = parse_int(entry.get("query_count")) + (0 if query_value in examples[:-1] else 1)
            return
    entry = {
        "key": key,
        "tool": tool,
        "path": path,
        "parameters": group_values,
        "example_url": url,
        "call_count": 1,
        "page_min": page or "",
        "page_max": page or "",
        "limit": limit or "",
    }
    if query_value:
        entry["query_examples"] = [query_value]
        entry["query_count"] = 1
    api_calls.append(entry)


def compact_api_calls(api_calls: Any) -> List[Dict[str, Any]]:
    if not isinstance(api_calls, list):
        return []
    compacted: List[Dict[str, Any]] = []
    for raw in api_calls:
        if not isinstance(raw, dict):
            continue
        if "example_url" in raw and "path" in raw:
            item = {k: v for k, v in raw.items() if k != "key"}
            compacted.append(item)
            continue
        url = raw.get("url") or raw.get("example_url") or ""
        if not url:
            continue
        parsed = urlparse(url)
        query = {k: values[0] if values else "" for k, values in parse_qs(parsed.query).items()}
        path = parsed.path.strip("/").split("/")[-1] + "/"
        tool = raw.get("tool") or tool_from_path(path)
        group_params = {k: v for k, v in query.items() if k not in {"page", "limit"}}
        if tool == "ricemind_search_by_trait_and_gene":
            group_params["trait"] = "<multiple traits>"
        key = json.dumps({"tool": tool, "path": path, "values": group_params}, sort_keys=True, ensure_ascii=False)
        page = parse_int(query.get("page"))
        limit = parse_int(query.get("limit"))
        for entry in compacted:
            if entry.get("key") == key:
                entry["call_count"] = parse_int(entry.get("call_count")) + 1
                if page:
                    entry["page_min"] = min(parse_int(entry.get("page_min")) or page, page)
                    entry["page_max"] = max(parse_int(entry.get("page_max")) or page, page)
                if limit:
                    entry["limit"] = limit
                entry["last_url"] = url
                trait = query.get("trait", "")
                if trait:
                    examples = entry.setdefault("query_examples", [])
                    if trait not in examples and len(examples) < 8:
                        examples.append(trait)
                break
        else:
            entry = {
                "key": key,
                "tool": tool,
                "path": path,
                "parameters": group_params,
                "example_url": url,
                "call_count": 1,
                "page_min": page or "",
                "page_max": page or "",
                "limit": limit or "",
            }
            trait = query.get("trait", "")
            if trait:
                entry["query_examples"] = [trait]
            compacted.append(entry)
    return [{k: v for k, v in entry.items() if k != "key"} for entry in compacted]


def tool_from_path(path: str) -> str:
    mapping = {
        "gene-profile/": "ricemind_get_gene_profile",
        "gene-omics-sequence/": "ricemind_get_gene_omics_sequence",
        "traits-by-gene/": "ricemind_get_traits_by_gene",
        "varieties-by-gene/": "ricemind_get_varieties_by_gene",
        "search-by-gene/": "ricemind_search_by_gene",
        "search-by-trait-and-gene/": "ricemind_search_by_trait_and_gene",
        "sentence-context/": "ricemind_get_sentence_context",
    }
    return mapping.get(path, path.rstrip("/"))


def extract_records(payload: Any, record_keys: Sequence[str]) -> List[Dict[str, Any]]:
    if not isinstance(payload, dict):
        return flatten_records(payload)
    for key in record_keys:
        if isinstance(payload.get(key), list):
            return flatten_records(payload[key])
    if isinstance(payload.get("sentence_pagination"), dict) and isinstance(payload.get("sentence_evidence"), list):
        return flatten_records(payload["sentence_evidence"])
    return flatten_records(payload)


def find_total_pages(payload: Any) -> int:
    if not isinstance(payload, dict):
        return 0
    for key in ("total_pages", "totalPages"):
        if key in payload:
            return parse_int(payload[key])
    if isinstance(payload.get("sentence_pagination"), dict):
        for key in ("total_pages", "totalPages"):
            if key in payload["sentence_pagination"]:
                return parse_int(payload["sentence_pagination"][key])
    return 0


def find_total_count(payload: Any) -> int:
    if not isinstance(payload, dict):
        return 0
    for key in ("total_count", "total_associated_traits", "total_associated_varieties", "total_sentences"):
        if key in payload:
            return parse_int(payload[key])
    if isinstance(payload.get("sentence_pagination"), dict):
        for key in ("total_sentences", "total_count"):
            if key in payload["sentence_pagination"]:
                return parse_int(payload["sentence_pagination"][key])
    return 0


def collect_paginated(
    base: str,
    endpoint: Dict[str, Any],
    values: Dict[str, Any],
    record_keys: Sequence[str],
    page_limit: int,
    max_pages: int,
    api_calls: List[Dict[str, str]],
) -> Dict[str, Any]:
    records: List[Dict[str, Any]] = []
    pages: List[Dict[str, Any]] = []
    page = 1
    total_count = 0
    total_pages = 0
    while page <= max_pages:
        page_values = dict(values)
        page_values.update({"page": page, "limit": page_limit})
        payload = call_get(base, endpoint, page_values, api_calls)
        if isinstance(payload, dict):
            pages.append(payload)
        page_records = extract_records(payload, record_keys)
        records.extend(page_records)
        total_count = total_count or find_total_count(payload)
        total_pages = total_pages or find_total_pages(payload)
        if total_pages and page >= total_pages:
            break
        if not total_pages:
            if not page_records:
                break
            if total_count and len(records) >= total_count:
                break
            if len(page_records) < page_limit:
                break
        page += 1
    else:
        raise RuntimeError(f"Exceeded max_pages={max_pages} for endpoint {endpoint.get('path')}")
    return {
        "records": records,
        "pages": pages,
        "total_count": total_count or len(records),
        "total_pages": total_pages or page,
    }


def fetch_from_api(
    base: str,
    gene: str,
    endpoint_map_path: Optional[Path],
    page_limit: int,
    max_pages: int,
    fetch_trait_evidence: bool,
) -> Dict[str, Any]:
    endpoints = dict(DEFAULT_ENDPOINTS)
    if endpoint_map_path:
        endpoints.update(read_json(endpoint_map_path))
    for name, endpoint in endpoints.items():
        endpoint.setdefault("tool", name)

    api_calls: List[Dict[str, str]] = []
    bundle: Dict[str, Any] = {"api_base": base, "api_calls": api_calls}
    bundle["gene_profile"] = call_get(base, endpoints["ricemind_get_gene_profile"], {"gene": gene}, api_calls)

    profile = choose_profile(bundle)
    rap_id = extract_standard_rap_id(profile) or gene
    if rap_id:
        try:
            bundle["omics_sequence"] = call_get(
                base,
                endpoints["ricemind_get_gene_omics_sequence"],
                {"rap_id": rap_id},
                api_calls,
            )
        except Exception as exc:
            bundle["omics_sequence_error"] = str(exc)

    try:
        bundle["traits_by_gene"] = collect_paginated(
            base,
            endpoints["ricemind_get_traits_by_gene"],
            {"gene": gene, "confidence": "All", "onto_type": "ALL"},
            endpoints["ricemind_get_traits_by_gene"].get("record_keys", ["associated_traits"]),
            page_limit,
            max_pages,
            api_calls,
        )
    except Exception as exc:
        bundle["traits_by_gene"] = {"records": [], "total_count": 0, "total_pages": 0, "error": str(exc)}
        bundle["traits_by_gene_error"] = str(exc)
    try:
        bundle["varieties_by_gene"] = collect_paginated(
            base,
            endpoints["ricemind_get_varieties_by_gene"],
            {"gene": gene},
            endpoints["ricemind_get_varieties_by_gene"].get("record_keys", ["varieties"]),
            page_limit,
            max_pages,
            api_calls,
        )
    except Exception as exc:
        bundle["varieties_by_gene_error"] = str(exc)
    try:
        bundle["gene_sentences"] = collect_paginated(
            base,
            endpoints["ricemind_search_by_gene"],
            {"gene": gene},
            endpoints["ricemind_search_by_gene"].get("record_keys", ["results"]),
            page_limit,
            max_pages,
            api_calls,
        )
    except Exception as exc:
        bundle["gene_sentences_error"] = str(exc)

    if fetch_trait_evidence:
        trait_evidence: Dict[str, Any] = {}
        for trait_record in normalize_traits(bundle):
            trait = trait_record["trait"]
            if not trait or trait in trait_evidence:
                continue
            try:
                trait_evidence[trait] = collect_paginated(
                    base,
                    endpoints["ricemind_search_by_trait_and_gene"],
                    {"gene": gene, "trait": trait},
                    endpoints["ricemind_search_by_trait_and_gene"].get("record_keys", ["sentence_evidence"]),
                    page_limit,
                    max_pages,
                    api_calls,
                )
            except Exception as exc:
                trait_evidence[trait] = {"records": [], "error": str(exc)}
        bundle["trait_evidence"] = trait_evidence
    return bundle


def infer_confidence(record: Dict[str, Any]) -> str:
    explicit = get_value(record, "confidence")
    normalized = explicit.upper().replace("-", "_")
    if normalized in {"HIGH", "TIER1", "TIER_1", "TIER1_KNOWN_RECONSTRUCTED"}:
        return "High"
    if normalized in MEDIUM_NAMES:
        return "Medium"
    if normalized in LOW_NAMES:
        return "Low"
    code_text = get_value(record, "evidence_code").upper()
    codes = {part.strip() for part in re.split(r"[;,|]\s*", code_text) if part.strip()}
    if codes & HIGH_CODES:
        return "High"
    support = parse_int(get_value(record, "support"))
    source_text = get_value(record, "source_db").upper()
    if "RICEMIND_NLP" in source_text and support > 10:
        return "Medium"
    if "NLP_COOCCURRENCE" in codes and support > 10:
        return "Medium"
    if code_text or source_text:
        return "Low" if "NLP" in code_text or "NLP" in source_text else "Unspecified"
    return explicit or "Unspecified"


def parse_int(value: Any) -> int:
    try:
        return int(float(stringify(value)))
    except Exception:
        return 0


def normalize_traits(bundle: Dict[str, Any]) -> List[Dict[str, str]]:
    source = (
        bundle.get("traits_by_gene")
        or bundle.get("ricemind_get_traits_by_gene")
        or bundle.get("traits")
        or bundle.get("associations")
    )
    rows = []
    for rec in flatten_records(source):
        trait = get_value(rec, "trait")
        if not trait:
            continue
        ontology_id = get_value(rec, "ontology_id")
        rows.append(
            {
                "gene": get_value(rec, "gene") or stringify(bundle.get("gene", "")),
                "rap_id": get_value(rec, "rap_id"),
                "trait": trait,
                "ontology_type": get_value(rec, "ontology_type") or infer_ontology_type(ontology_id),
                "ontology_id": ontology_id,
                "evidence_code": get_value(rec, "evidence_code"),
                "source_db": get_value(rec, "source_db"),
                "confidence": infer_confidence(rec),
                "support": stringify(parse_int(get_value(rec, "support"))),
                "year": get_value(rec, "year"),
            }
        )
    return dedupe_rows(rows, ["trait", "ontology_id", "confidence", "evidence_code", "source_db"])


def infer_ontology_type(ontology_id: str) -> str:
    if ":" not in ontology_id:
        return ""
    prefix = ontology_id.split(":", 1)[0].strip().upper()
    return prefix or ""


def normalize_varieties(bundle: Dict[str, Any]) -> List[Dict[str, str]]:
    rows = []
    source = bundle.get("varieties_by_gene")
    if isinstance(source, dict) and isinstance(source.get("records"), list):
        raw_records = source["records"]
    elif isinstance(source, dict) and isinstance(source.get("varieties"), list):
        raw_records = source["varieties"]
    elif isinstance(source, list):
        raw_records = source
    else:
        raw_records = flatten_records(source)
    for rec in raw_records:
        if isinstance(rec, dict):
            variety = get_value(rec, "trait") or get_value(rec, "gene") or stringify(rec.get("normalization", ""))
        else:
            variety = stringify(rec)
        if variety:
            rows.append({"variety": variety})
    return dedupe_rows(rows, ["variety"])


def trait_lookup(traits: List[Dict[str, str]]) -> Dict[str, Dict[str, str]]:
    lookup: Dict[str, Dict[str, str]] = {}
    for row in traits:
        key = row["trait"].strip().lower()
        if key and key not in lookup:
            lookup[key] = row
    return lookup


def normalize_evidence(bundle: Dict[str, Any], traits: Optional[List[Dict[str, str]]] = None) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    evidence_sources: List[Tuple[str, Any]] = []
    for key in ("trait_evidence", "evidence_by_trait"):
        if isinstance(bundle.get(key), dict):
            evidence_sources.extend(bundle[key].items())
    for key in ("gene_sentences", "evidence", "sentences", "ricemind_search_by_gene"):
        if key in bundle:
            evidence_sources.append(("", bundle[key]))

    lookup = trait_lookup(traits or [])
    for trait_name, source in evidence_sources:
        trait_meta = lookup.get(trait_name.strip().lower(), {})
        for rec in flatten_records(source):
            sentence = get_value(rec, "sentence")
            pmid = get_value(rec, "pmid")
            if not sentence and not pmid:
                continue
            row_trait = get_value(rec, "trait", trait_name) or trait_name
            meta = lookup.get(row_trait.strip().lower(), trait_meta)
            rows.append(
                {
                    "trait": row_trait,
                    "ontology_type": get_value(rec, "ontology_type") or meta.get("ontology_type", ""),
                    "ontology_id": get_value(rec, "ontology_id") or meta.get("ontology_id", ""),
                    "evidence_code": get_value(rec, "evidence_code") or meta.get("evidence_code", ""),
                    "source_db": get_value(rec, "source_db") or meta.get("source_db", ""),
                    "confidence": infer_confidence(rec) if (get_value(rec, "confidence") or get_value(rec, "evidence_code")) else meta.get("confidence", ""),
                    "pmid": pmid,
                    "sentence_id": get_value(rec, "sentence_id"),
                    "year": get_value(rec, "year"),
                    "title": get_value(rec, "title"),
                    "journal": get_value(rec, "journal"),
                    "doi": get_value(rec, "doi"),
                    "sentence": sentence,
                }
            )
    return dedupe_rows(rows, ["trait", "pmid", "sentence_id", "sentence"])


def dedupe_rows(rows: List[Dict[str, str]], keys: Sequence[str]) -> List[Dict[str, str]]:
    seen = set()
    deduped = []
    for row in rows:
        key = tuple(row.get(k, "") for k in keys)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def choose_profile(bundle: Dict[str, Any]) -> Dict[str, Any]:
    source = bundle.get("gene_profile") or bundle.get("ricemind_get_gene_profile") or bundle.get("profile") or {}
    if isinstance(source, dict) and isinstance(source.get("data"), dict):
        return source["data"]
    if isinstance(source, dict):
        return source
    records = flatten_records(source)
    return records[0] if records else {}


def extract_standard_rap_id(profile: Dict[str, Any]) -> str:
    return deep_find(profile, ["standard_rap_id", "rap_id", "primary_id", "gene_id"])


def safe_filename(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", text).strip("_") or "gene"


def write_csv(path: Path, rows: List[Dict[str, str]], fieldnames: Optional[Sequence[str]] = None) -> None:
    if not rows and not fieldnames:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    columns = list(fieldnames or rows[0].keys())
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=columns)
        writer.writeheader()
        if rows:
            writer.writerows(rows)


def plot_counter(counter: Counter, title: str, out_path: Path, xlabel: str = "Count", top_n: int = 20) -> Optional[Path]:
    if not counter:
        return None
    try:
        import matplotlib.pyplot as plt
    except Exception:
        return plot_counter_pillow(counter, title, out_path, xlabel=xlabel, top_n=top_n)
    labels, values = zip(*counter.most_common(top_n))
    fig_height = max(3.2, 0.32 * len(labels) + 1.4)
    fig, ax = plt.subplots(figsize=(7.4, fig_height))
    ax.barh(range(len(labels)), values, color="#3f6f8f")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels)
    ax.invert_yaxis()
    ax.set_xlabel(xlabel)
    ax.set_title(title)
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=220)
    plt.close(fig)
    return out_path


def plot_years(years: Iterable[str], out_path: Path) -> Optional[Path]:
    clean = [parse_int(y) for y in years if parse_int(y) > 0]
    if not clean:
        return None
    try:
        import matplotlib.pyplot as plt
    except Exception:
        return plot_years_pillow(Counter(clean), out_path)
    counts = Counter(clean)
    xs = sorted(counts)
    ys = [counts[x] for x in xs]
    fig, ax = plt.subplots(figsize=(7.4, 3.8))
    ax.bar(xs, ys, color="#6b8f3f")
    ax.set_xlabel("Publication year")
    ax.set_ylabel("Evidence records")
    ax.set_title("RiceMind evidence by publication year")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=220)
    plt.close(fig)
    return out_path


def plot_top_traits(traits: List[Dict[str, str]], out_path: Path) -> Optional[Path]:
    counter = Counter({row["trait"]: parse_int(row["support"]) for row in traits if row["trait"]})
    return plot_counter(counter, "Top traits by supporting articles", out_path, xlabel="Supporting articles", top_n=25)


def load_font(size: int, bold: bool = False) -> Any:
    try:
        from PIL import ImageFont

        candidates = [
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
            "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except Exception:
                continue
        return ImageFont.load_default()
    except Exception:
        return None


def text_width(draw: Any, text: str, font: Any) -> int:
    try:
        box = draw.textbbox((0, 0), text, font=font)
        return box[2] - box[0]
    except Exception:
        return len(text) * 8


def ellipsize(draw: Any, text: str, font: Any, max_width: int) -> str:
    text = stringify(text)
    if text_width(draw, text, font) <= max_width:
        return text
    while text and text_width(draw, text + "...", font) > max_width:
        text = text[:-1]
    return text + "..." if text else "..."


def plot_counter_pillow(counter: Counter, title: str, out_path: Path, xlabel: str = "Count", top_n: int = 20) -> Optional[Path]:
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None
    items = counter.most_common(top_n)
    if not items:
        return None
    label_font = load_font(14)
    title_font = load_font(18, bold=True)
    axis_font = load_font(12)
    left = 250
    right = 70
    top = 70
    row_h = 28
    bottom = 55
    width = 900
    height = top + len(items) * row_h + bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 18), title, fill="#1f2933", font=title_font)
    max_value = max(value for _, value in items) or 1
    bar_area = width - left - right
    for idx, (label, value) in enumerate(items):
        y = top + idx * row_h
        draw.text((20, y + 4), ellipsize(draw, label, label_font, left - 35), fill="#111827", font=label_font)
        bar_w = max(2, int(bar_area * (value / max_value)))
        draw.rectangle((left, y + 5, left + bar_w, y + row_h - 6), fill="#3f6f8f")
        draw.text((left + bar_w + 6, y + 3), str(value), fill="#111827", font=axis_font)
    draw.line((left, top - 8, left, top + len(items) * row_h), fill="#9ca3af", width=1)
    draw.text((left, height - 32), xlabel, fill="#374151", font=axis_font)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(out_path)
    return out_path


def plot_years_pillow(counts: Counter, out_path: Path) -> Optional[Path]:
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None
    if not counts:
        return None
    title_font = load_font(18, bold=True)
    axis_font = load_font(12)
    width = 950
    height = 420
    left = 60
    right = 30
    top = 60
    bottom = 70
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 18), "RiceMind evidence by publication year", fill="#1f2933", font=title_font)
    years = sorted(counts)
    max_value = max(counts.values()) or 1
    plot_w = width - left - right
    plot_h = height - top - bottom
    bar_gap = 2
    bar_w = max(3, int(plot_w / max(1, len(years))) - bar_gap)
    draw.line((left, top, left, top + plot_h), fill="#9ca3af", width=1)
    draw.line((left, top + plot_h, left + plot_w, top + plot_h), fill="#9ca3af", width=1)
    tick_step = max(1, math.ceil(len(years) / 12))
    for idx, year in enumerate(years):
        value = counts[year]
        x = left + idx * (plot_w / max(1, len(years)))
        h = int(plot_h * (value / max_value))
        draw.rectangle((int(x), top + plot_h - h, int(x) + bar_w, top + plot_h), fill="#6b8f3f")
        if idx % tick_step == 0:
            draw.text((int(x) - 8, top + plot_h + 8), str(year), fill="#374151", font=axis_font)
    draw.text((left, height - 26), "Publication year", fill="#374151", font=axis_font)
    draw.text((left + 5, top + 5), f"max={max_value}", fill="#374151", font=axis_font)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(out_path)
    return out_path


def split_counter(rows: Iterable[Dict[str, str]], field: str) -> Counter:
    counter: Counter = Counter()
    for row in rows:
        value = row.get(field, "")
        for token in re.split(r"[;,|]\s*", value):
            token = token.strip()
            if token:
                counter[token] += 1
    return counter


def api_call_rows(api_calls: List[Dict[str, Any]], max_rows: int = REPORT_API_ROWS) -> List[List[str]]:
    rows = [["No.", "API function", "Call summary", "Example URL"]]
    for idx, call in enumerate(api_calls[:max_rows], 1):
        page_summary = ""
        if call.get("page_min") or call.get("page_max"):
            page_summary = f"pages {call.get('page_min')}-{call.get('page_max')}, limit={call.get('limit')}"
        call_summary = "; ".join(part for part in [
            endpoint_purpose(stringify(call.get("tool") or call.get("path"))),
            page_summary,
            f"calls={call.get('call_count')}" if call.get("call_count") else "",
            f"examples={'; '.join(call.get('query_examples', [])[:3])}" if call.get("query_examples") else "",
        ] if part)
        rows.append([
            str(idx),
            stringify(call.get("tool") or call.get("path")),
            truncate(call_summary, 160),
            truncate(stringify(call.get("example_url") or call.get("url")), 160),
        ])
    return rows


def endpoint_error_rows(bundle: Dict[str, Any]) -> List[List[str]]:
    rows = [["Endpoint", "Error"]]
    for key, value in bundle.items():
        if key.endswith("_error") and value:
            rows.append([key[:-6], truncate(stringify(value), 240)])
    return rows if len(rows) > 1 else []


def endpoint_purpose(tool: str) -> str:
    purposes = {
        "ricemind_get_gene_profile": "gene profile and external links",
        "ricemind_get_gene_omics_sequence": "omics sequence retrieval",
        "ricemind_get_traits_by_gene": "full GTA trait landscape",
        "ricemind_get_varieties_by_gene": "gene-variety co-occurrence",
        "ricemind_search_by_gene": "gene-wide sentence evidence",
        "ricemind_search_by_trait_and_gene": "trait-specific sentence evidence",
        "ricemind_get_sentence_context": "sentence context",
    }
    return purposes.get(tool, "")


def add_table(doc: Any, rows: List[List[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for i, value in enumerate(rows[0]):
        table.rows[0].cells[i].text = stringify(value)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            add_cell_value(cells[i], value)


def is_url(value: str) -> bool:
    return bool(re.match(r"^https?://[^\s]+$", stringify(value).strip()))


def add_cell_value(cell: Any, value: Any) -> None:
    text = stringify(value)
    if is_url(text):
        paragraph = cell.paragraphs[0]
        paragraph.text = ""
        add_hyperlink(paragraph, text, text, size=FONT_SIZES["table"])
    else:
        cell.text = text


def add_hyperlink(paragraph: Any, url: str, text: str, size: float = FONT_SIZES["body"]) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE
    except Exception:
        paragraph.add_run(text)
        return

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(rpr)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_picture_if_exists(doc: Any, path: Optional[Path], caption: str, width_inches: float = 6.2) -> None:
    if not path or not path.exists():
        return
    try:
        from docx.shared import Inches

        doc.add_picture(str(path), width=Inches(width_inches))
        doc.add_paragraph(caption)
    except Exception:
        pass


def zh(enabled: bool, cn: str, en: str) -> str:
    return cn if enabled else en


def truncate(text: str, max_chars: int) -> str:
    text = " ".join(stringify(text).split())
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3] + "..."


def configure_document_styles(doc: Any) -> None:
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except Exception:
        return

    style_specs = {
        "Normal": (FONT_SIZES["body"], False),
        "Title": (FONT_SIZES["title"], True),
        "Heading 1": (FONT_SIZES["heading1"], True),
        "Heading 2": (FONT_SIZES["heading2"], True),
        "Heading 3": (FONT_SIZES["body"], True),
    }
    for style_name, (size, bold) in style_specs.items():
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_LATIN
            style.font.size = Pt(size)
            style.font.bold = bold
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        except Exception:
            continue


def apply_document_fonts(doc: Any) -> None:
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        if style_name == "Title":
            size = FONT_SIZES["title"]
            bold = True
        elif style_name == "Heading 1":
            size = FONT_SIZES["heading1"]
            bold = True
        elif style_name == "Heading 2":
            size = FONT_SIZES["heading2"]
            bold = True
        else:
            size = FONT_SIZES["body"]
            bold = None
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=FONT_SIZES["table"], bold=(row_idx == 0))


def set_run_font(run: Any, size: float, bold: Optional[bool] = None) -> None:
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except Exception:
        return
    run.font.name = FONT_LATIN
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


def flatten_profile_rows(profile: Dict[str, Any]) -> List[List[str]]:
    rows = [["Field", "Value"]]

    def walk(prefix: str, value: Any) -> None:
        if isinstance(value, dict):
            for key, subvalue in value.items():
                walk(f"{prefix}.{key}" if prefix else stringify(key), subvalue)
        elif isinstance(value, list):
            rows.append([prefix, truncate("; ".join(stringify(item) for item in value), 500)])
        elif value not in (None, ""):
            rows.append([prefix, truncate(stringify(value), 500)])

    walk("", profile)
    return rows


def sequence_api_url_rows(bundle: Dict[str, Any], profile: Dict[str, Any], gene: str) -> List[List[str]]:
    urls: List[str] = []
    for call in compact_api_calls(bundle.get("api_calls") or []):
        tool = stringify(call.get("tool") or call.get("path"))
        if tool in {"ricemind_get_gene_omics_sequence", "gene-omics-sequence/"}:
            url = stringify(call.get("example_url") or call.get("url") or call.get("last_url"))
            if url and url not in urls:
                urls.append(url)
    if not urls:
        rap_id = extract_standard_rap_id(profile) or gene
        if rap_id:
            urls.append(build_url(bundle.get("api_base", DEFAULT_API_BASE), DEFAULT_ENDPOINTS["ricemind_get_gene_omics_sequence"], {"rap_id": rap_id}))
    if not urls:
        return []
    rows = [["No.", "RiceMind sequence API URL"]]
    for idx, url in enumerate(urls, 1):
        rows.append([str(idx), url])
    return rows


HOTSPOT_TERMS = [
    "heading date",
    "heading",
    "flowering",
    "photoperiod",
    "long-day",
    "short-day",
    "circadian",
    "Ehd1",
    "Hd3a",
    "RFT1",
    "grain number",
    "grain yield",
    "yield",
    "plant height",
    "panicle",
    "tiller",
    "QTL",
    "allele",
    "natural variation",
    "domestication",
    "adaptation",
    "drought",
    "salt",
    "nitrogen",
    "ABA",
    "stress",
    "transcriptome",
    "scRNA",
    "single-cell",
    "brown planthopper",
    "BPH",
    "resistance",
    "feeding",
    "defense",
    "introgression",
    "pyramiding",
    "marker-assisted",
    "GA",
    "gibberellin",
    "SLR1",
    "DELLA",
    "lodging",
    "semi-dwarf",
]

CONFLICT_CONTEXTS = [
    ("Heading/flowering time", "抽穗期开花", ["heading", "flowering", "floral", "photoperiod", "long-day", "short-day"]),
    ("Yield and grain traits", "产量和穗粒性状", ["yield", "grain", "panicle", "tiller", "spikelet", "harvest index"]),
    ("Plant height and growth", "株高和生长", ["plant height", "height", "growth", "dwarf", "semi-dwarf", "elongation"]),
    ("Drought and water deficit", "干旱和水分亏缺", ["drought", "water deficit", "wilting", "dehydration"]),
    ("Salt/osmotic stress", "盐和渗透胁迫", ["salt", "salinity", "nacl", "osmotic"]),
    ("Pest/pathogen resistance", "虫害病害抗性", ["brown planthopper", "bph", "pest", "insect", "pathogen", "disease", "resistance"]),
    ("Expression/regulatory effect", "表达和调控方向", ["expression", "transcription", "up-regulated", "down-regulated", "repress", "activate"]),
]

POSITIVE_DIRECTION_TERMS = [
    "promote",
    "promotes",
    "enhance",
    "enhanced",
    "increase",
    "increased",
    "improve",
    "improved",
    "activate",
    "activated",
    "up-regulated",
    "upregulated",
    "positive regulator",
    "confers resistance",
    "resistant",
    "tolerant",
    "higher",
    "more",
]

NEGATIVE_DIRECTION_TERMS = [
    "inhibit",
    "inhibited",
    "repress",
    "repressed",
    "suppress",
    "suppressed",
    "decrease",
    "decreased",
    "reduce",
    "reduced",
    "down-regulated",
    "downregulated",
    "negative regulator",
    "susceptible",
    "hypersensitive",
    "delay",
    "delayed",
    "lower",
    "less",
]


def evidence_summary_rows(traits: List[Dict[str, str]], evidence: List[Dict[str, str]]) -> List[List[str]]:
    return [
        ["Metric", "Value"],
        ["Trait associations", str(len(traits))],
        ["Sentence evidence records", str(len(evidence))],
        ["Unique PMIDs", str(len({row.get("pmid", "") for row in evidence if row.get("pmid", "")}))],
        ["Ontology types in traits", str(len({row.get("ontology_type", "") for row in traits if row.get("ontology_type", "")}))],
        ["Evidence years", year_span_text(evidence)],
    ]


def year_span_text(rows: List[Dict[str, str]]) -> str:
    years = sorted({parse_int(row.get("year", "")) for row in rows if parse_int(row.get("year", "")) > 0})
    if not years:
        return "unavailable"
    return f"{years[0]}-{years[-1]}"


def ontology_distribution_rows(traits: List[Dict[str, str]], evidence: List[Dict[str, str]]) -> List[List[str]]:
    trait_counts = Counter(row.get("ontology_type", "") or "Unspecified" for row in traits)
    evidence_counts = Counter(row.get("ontology_type", "") or "Unspecified" for row in evidence)
    ontologies = sorted(set(trait_counts) | set(evidence_counts), key=lambda key: (-(trait_counts[key] + evidence_counts[key]), key))
    rows = [["Ontology", "Trait records", "Evidence records", "Top trait examples"]]
    for ontology in ontologies[:12]:
        examples = []
        for row in traits:
            if (row.get("ontology_type", "") or "Unspecified") == ontology and row.get("trait", "") not in examples:
                examples.append(row.get("trait", ""))
            if len(examples) >= 4:
                break
        rows.append([ontology, str(trait_counts[ontology]), str(evidence_counts[ontology]), truncate("; ".join(examples), 160)])
    return rows


def representative_sentence_rows(evidence: List[Dict[str, str]], per_tier: int = 2) -> List[List[str]]:
    rows = [["Tier", "Trait / Ontology", "Evidence code / Source", "PMID / Year", "Representative RiceMind sentence"]]
    tier_order = ["High", "Medium", "Low", "Unspecified"]
    seen = set()
    for tier in tier_order:
        candidates = [row for row in evidence if (row.get("confidence", "") or "Unspecified") == tier and row.get("sentence", "")]
        candidates.sort(key=lambda row: (0 if row.get("pmid", "") else 1, -len(row.get("sentence", "")), row.get("trait", "")))
        added = 0
        for row in candidates:
            key = (tier, row.get("pmid", ""), row.get("sentence_id", ""), row.get("sentence", "")[:80])
            if key in seen:
                continue
            seen.add(key)
            rows.append([
                tier,
                " ".join(x for x in [truncate(row.get("trait", ""), 70), row.get("ontology_id", "")] if x),
                " / ".join(x for x in [truncate(row.get("evidence_code", ""), 60), truncate(row.get("source_db", ""), 60)] if x),
                " / ".join(x for x in [row.get("pmid", ""), row.get("year", "")] if x),
                truncate(row.get("sentence", ""), 260),
            ])
            added += 1
            if added >= per_tier:
                break
    return rows


def phase_ranges(years: List[int]) -> List[Tuple[str, int, int]]:
    unique_years = sorted(set(years))
    if not unique_years:
        return []
    if len(unique_years) <= 3:
        return [(f"Phase {idx + 1}", year, year) for idx, year in enumerate(unique_years)]
    cut1 = unique_years[max(0, len(unique_years) // 3 - 1)]
    cut2 = unique_years[max(0, (2 * len(unique_years)) // 3 - 1)]
    return [
        ("Phase I", unique_years[0], cut1),
        ("Phase II", cut1 + 1, cut2),
        ("Phase III", cut2 + 1, unique_years[-1]),
    ]


def temporal_hotspot_rows(evidence: List[Dict[str, str]], is_zh: bool) -> List[List[str]]:
    years = [parse_int(row.get("year", "")) for row in evidence if parse_int(row.get("year", "")) > 0]
    rows = [["Phase", "Years", "Evidence records", "Unique PMIDs", "Hotspot terms", "Interpretation"]]
    for phase, start, end in phase_ranges(years):
        phase_rows = [row for row in evidence if start <= parse_int(row.get("year", "")) <= end]
        if not phase_rows:
            continue
        terms = observed_terms(phase_rows, HOTSPOT_TERMS, 8)
        top_traits = [trait for trait, _ in Counter(row.get("trait", "") for row in phase_rows if row.get("trait", "")).most_common(5)]
        interpretation = zh(
            is_zh,
            f"热点集中在 {join_cn(terms) if terms else '可追溯句证'}；相关 trait/语境包括 {join_cn(top_traits) if top_traits else '未明确'}。",
            f"Hotspots center on {join_en(terms) if terms else 'traceable evidence'}; associated traits/contexts include {join_en(top_traits) if top_traits else 'unspecified'}.",
        )
        rows.append([
            phase,
            f"{start}-{end}",
            str(len(phase_rows)),
            str(len({row.get("pmid", "") for row in phase_rows if row.get("pmid", "")})),
            truncate(join_cn(terms) if is_zh else join_en(terms), 160),
            truncate(interpretation, 240),
        ])
    return rows


def rows_with_terms(rows: List[Dict[str, str]], terms: Sequence[str]) -> List[Dict[str, str]]:
    lowered = [term.lower() for term in terms]
    return [row for row in rows if any(term in f"{row.get('trait', '')} {row.get('sentence', '')}".lower() for term in lowered)]


def conflict_rows(evidence: List[Dict[str, str]], is_zh: bool) -> List[List[str]]:
    rows = [["Context", "Positive-direction evidence", "Negative-direction evidence", "Synthesis / Caveat"]]
    for context_en, context_cn, context_terms in CONFLICT_CONTEXTS:
        context_rows = rows_with_terms(evidence, context_terms)
        if len(context_rows) < 3:
            continue
        positive_rows = rows_with_terms(context_rows, POSITIVE_DIRECTION_TERMS)
        negative_rows = rows_with_terms(context_rows, NEGATIVE_DIRECTION_TERMS)
        if not positive_rows or not negative_rows:
            continue
        positive_cite = pmid_citation(collect_pmids(positive_rows, 4))
        negative_cite = pmid_citation(collect_pmids(negative_rows, 4))
        caveat = zh(
            is_zh,
            "RiceMind 检测到同一语境下存在正向和负向方向词；这通常提示光周期、发育阶段、遗传背景、胁迫强度或证据等级差异，而不是简单的互相否定。",
            "RiceMind detected both positive and negative directional terms in the same context; this usually indicates photoperiod, developmental stage, genetic background, stress intensity or evidence-tier dependence rather than simple mutual contradiction.",
        )
        rows.append([
            context_cn if is_zh else context_en,
            f"n={len(positive_rows)} {positive_cite}",
            f"n={len(negative_rows)} {negative_cite}",
            caveat,
        ])
    return rows if len(rows) > 1 else []


def bibliometric_rows(evidence: List[Dict[str, str]]) -> Tuple[List[List[str]], List[List[str]]]:
    journal_counter = Counter(row.get("journal", "").strip() for row in evidence if row.get("journal", "").strip())
    journal_rows = [["Journal", "Evidence records"]]
    for journal, count in journal_counter.most_common(12):
        journal_rows.append([truncate(journal, 120), str(count)])

    pmid_counter = Counter(row.get("pmid", "").strip() for row in evidence if row.get("pmid", "").strip())
    title_by_pmid: Dict[str, str] = {}
    year_by_pmid: Dict[str, str] = {}
    for row in evidence:
        pmid = row.get("pmid", "").strip()
        if not pmid:
            continue
        title_by_pmid.setdefault(pmid, row.get("title", ""))
        year_by_pmid.setdefault(pmid, row.get("year", ""))
    pmid_rows = [["PMID", "Year", "Evidence records", "Title"]]
    for pmid, count in pmid_counter.most_common(12):
        pmid_rows.append([pmid, year_by_pmid.get(pmid, ""), str(count), truncate(title_by_pmid.get(pmid, ""), 180)])
    return (journal_rows if len(journal_rows) > 1 else [], pmid_rows if len(pmid_rows) > 1 else [])


def row_limit(rows: List[Dict[str, str]], max_rows: int) -> List[Dict[str, str]]:
    if max_rows and max_rows > 0:
        return rows[:max_rows]
    return rows


THEME_RULES = [
    (
        "Gibberellin and hormone metabolism",
        "赤霉素和激素代谢",
        ["gibberellin", " ga ", "ga20", "ga3", "hormone", "auxin", "cytokinin", "jasmonic", "abscisic", "salicylic", "brassinosteroid", "ethylene"],
    ),
    (
        "Plant height, dwarfism and architecture",
        "株高、矮化和株型建成",
        ["plant height", "height", "dwarf", "semi-dwarf", "semidwarf", "culm", "internode", "elongation", "architecture", "stem"],
    ),
    (
        "Yield and reproductive development",
        "产量和生殖发育",
        ["yield", "grain", "seed", "panicle", "flower", "heading", "tiller", "fertility", "reproductive", "spikelet"],
    ),
    (
        "Stress response and resistance",
        "逆境响应和抗性",
        ["stress", "drought", "salt", "cold", "heat", "pathogen", "blast", "resistance", "disease", "immunity", "defense"],
    ),
    (
        "Vegetative growth and physiology",
        "营养生长和生理性状",
        ["leaf", "root", "shoot", "chlorophyll", "photosynthesis", "green", "biomass", "seedling"],
    ),
    (
        "Molecular function and gene regulation",
        "分子功能和基因调控",
        ["biosynthetic", "metabolic", "oxidase", "enzyme", "protein", "expression", "transcription", "regulation", "mutant", "mutation", "allele"],
    ),
]

MECHANISM_TERMS = [
    "gibberellin",
    "GA20",
    "GA",
    "plant height",
    "dwarf",
    "semi-dwarf",
    "biosynthetic",
    "oxidase",
    "mutation",
    "mutant",
    "expression",
    "internode",
    "culm",
    "grain",
    "yield",
    "panicle",
    "tiller",
    "stress",
    "resistance",
    "photosynthesis",
    "chlorophyll",
    "enzyme",
    "regulation",
]

MECHANISM_SIGNAL_RULES = [
    {
        "key": "ga_biosynthesis",
        "zh": "GA/赤霉素生物合成",
        "en": "GA/gibberellin biosynthesis",
        "keywords": [
            "gibberellin biosynthesis",
            "gibberellin synthesis",
            "ga biosynthesis",
            "ga synthesis",
            "ga20 oxidase",
            "ga 20-oxidase",
            "ga20ox",
            "c20-oxidase",
            "20-oxidase",
            "bioactive ga",
            "gibberellic acid",
        ],
    },
    {
        "key": "mutation_loss",
        "zh": "功能缺失、突变或表达降低",
        "en": "loss-of-function, mutation or reduced expression",
        "keywords": [
            "loss-of-function",
            "loss of function",
            "defective",
            "deletion",
            "mutation",
            "mutant allele",
            "mutant alleles",
            "loss of expression",
            "reduced expression",
            "knockout",
            "allelic",
        ],
    },
    {
        "key": "architecture_height",
        "zh": "株高、半矮秆和节间/秆长建成",
        "en": "plant height, semi-dwarfism and internode/culm architecture",
        "keywords": [
            "plant height",
            "semi-dwarf",
            "semidwarf",
            "dwarf",
            "culm length",
            "culm",
            "internode",
            "elongation",
            "reduced height",
            "shorter",
        ],
    },
    {
        "key": "yield_lodging_breeding",
        "zh": "产量、抗倒伏和育种利用",
        "en": "yield, lodging resistance and breeding use",
        "keywords": [
            "yield",
            "lodging",
            "harvest index",
            "fertilizer",
            "green revolution",
            "breeding",
            "modern rice cultivars",
            "elite rice",
            "high-yielding",
        ],
    },
    {
        "key": "reproductive_growth",
        "zh": "生殖发育和生长阶段",
        "en": "reproductive development and growth stages",
        "keywords": [
            "heading",
            "panicle",
            "grain",
            "seed",
            "spikelet",
            "flower",
            "fertility",
            "mesocotyl",
            "seedling",
        ],
    },
    {
        "key": "stress_defense",
        "zh": "逆境、抗病和防御响应",
        "en": "stress, disease resistance and defense response",
        "keywords": [
            "stress",
            "drought",
            "salt",
            "cold",
            "heat",
            "pathogen",
            "blast",
            "disease",
            "resistance",
            "immunity",
            "defense",
        ],
    },
    {
        "key": "expression_regulation",
        "zh": "表达调控和分子功能",
        "en": "expression regulation and molecular function",
        "keywords": [
            "expression",
            "transcription",
            "regulation",
            "encoded by",
            "encoding",
            "enzyme",
            "protein",
            "homolog",
            "candidate gene",
            "qtl",
            "locus",
        ],
    },
]


def assign_theme(text: str) -> Tuple[str, str]:
    padded = " " + text.lower() + " "
    for en, cn, keywords in THEME_RULES:
        if any(keyword in padded for keyword in keywords):
            return en, cn
    return "Other RiceMind trait associations", "其他 RiceMind 性状关联"


def collect_pmids(rows: List[Dict[str, str]], limit: int = 8) -> List[str]:
    counts: Counter = Counter(row.get("pmid", "").strip() for row in rows if row.get("pmid", "").strip())
    if counts:
        return [pmid for pmid, _ in counts.most_common(limit)]
    pmids = []
    for row in rows:
        pmid = row.get("pmid", "").strip()
        if pmid and pmid not in pmids:
            pmids.append(pmid)
        if len(pmids) >= limit:
            break
    return pmids


def pmid_citation(pmids: List[str]) -> str:
    if not pmids:
        return ""
    return "[" + ", ".join(pmids) + "]"


def top_terms(rows: List[Dict[str, str]], limit: int = 6) -> List[str]:
    counter: Counter = Counter()
    for row in rows:
        haystack = f"{row.get('trait', '')} {row.get('sentence', '')}".lower()
        for term in MECHANISM_TERMS:
            if term.lower() in haystack:
                counter[term] += 1
    return [term for term, _ in counter.most_common(limit)]


def top_traits_for_rows(rows: List[Dict[str, str]], traits: List[Dict[str, str]], limit: int = 6) -> List[str]:
    support = {row["trait"]: parse_int(row.get("support")) for row in traits}
    counter: Counter = Counter()
    for row in rows:
        trait = row.get("trait", "").strip()
        if trait:
            counter[trait] += 1 + support.get(trait, 0) / 1000
    return [trait for trait, _ in counter.most_common(limit)]


def signal_matches(row: Dict[str, str], keywords: Sequence[str]) -> bool:
    haystack = f"{row.get('trait', '')} {row.get('sentence', '')}".lower()
    return any(keyword in haystack for keyword in keywords)


def summarize_signal_rows(rows: List[Dict[str, str]], limit: int = 5) -> List[Tuple[Dict[str, Any], List[Dict[str, str]]]]:
    matched: List[Tuple[Dict[str, Any], List[Dict[str, str]]]] = []
    for rule in MECHANISM_SIGNAL_RULES:
        signal_rows = [row for row in rows if signal_matches(row, rule["keywords"])]
        if signal_rows:
            matched.append((rule, signal_rows))
    matched.sort(key=lambda item: len(item[1]), reverse=True)
    return matched[:limit]


def compact_counter_text(rows: List[Dict[str, str]], field: str, limit: int = 3) -> str:
    counts: Counter = Counter()
    for row in rows:
        for token in re.split(r"[;,|]\s*", row.get(field, "")):
            token = token.strip()
            if token:
                counts[token] += 1
    return ", ".join(f"{key}={value}" for key, value in counts.most_common(limit))


def signal_clause(gene: str, rule: Dict[str, Any], rows: List[Dict[str, str]], is_zh: bool) -> str:
    citation = pmid_citation(collect_pmids(rows, 6))
    if is_zh:
        clauses = {
            "ga_biosynthesis": (
                f"句子内容反复把 {gene} 放在 GA20 oxidase/C20-oxidase、赤霉素生物合成和活性 GA 水平调控的语境中，"
                f"这使该主题下的证据核心从一般 trait 共现具体化为 GA 合成通路及其后期氧化步骤 {citation}。"
            ),
            "mutation_loss": (
                f"同一组证据还把 deletion、defective gene、loss-of-function、loss of expression 或等位突变等表述与 {gene} 连接起来，"
                f"提示 RiceMind 文本证据中的机制线索集中在功能缺失或表达降低造成的通路扰动 {citation}。"
            ),
            "architecture_height": (
                f"在表型端，句子证据将上述分子变化与 semi-dwarf/dwarf、plant height、culm length、internode 或 elongation 等表型词汇串联，"
                f"说明 {gene} 相关证据主要通过株高、半矮秆和节间/秆长建成体现出来 {citation}。"
            ),
            "yield_lodging_breeding": (
                f"育种应用相关句子进一步把 {gene} 的半矮化语境与 lodging resistance、fertilizer response、harvest index、yield 或 Green Revolution 品种利用相连，"
                f"因此该部分更适合表述为 RiceMind 文献证据支持的育种价值链条，而不是单一分子事件 {citation}。"
            ),
            "reproductive_growth": (
                f"生长发育相关句子把 {gene} 与 heading、panicle、grain、seed、spikelet、flower、mesocotyl 或 seedling 等发育阶段和器官性状共同报告，"
                f"为其在生殖发育或生长阶段中的表型延伸提供 RiceMind 句子级线索 {citation}。"
            ),
            "stress_defense": (
                f"逆境和抗性相关句子把 {gene} 与 stress、disease、resistance、immunity 或 defense 等词汇共同报告，"
                f"这些证据目前应被解释为 RiceMind 文献中的候选关联或探索性生物学语境 {citation}。"
            ),
            "expression_regulation": (
                f"表达和分子功能相关句子包含 encoding、enzyme、protein、expression、candidate gene、QTL 或 locus 等线索，"
                f"可用于把 trait 关联追溯到 RiceMind 文本中出现的分子身份、表达或候选位点描述 {citation}。"
            ),
        }
        return clauses.get(rule["key"], f"句子证据反复出现 {rule['zh']} 相关线索 {citation}。")

    clauses = {
        "ga_biosynthesis": (
            f"Sentence evidence repeatedly places {gene} in the context of GA20 oxidase/C20-oxidase, gibberellin biosynthesis and bioactive GA control, "
            f"turning the association into a pathway-level RiceMind signal for late GA biosynthetic oxidation steps {citation}."
        ),
        "mutation_loss": (
            f"The same evidence links {gene} with deletion, defective gene, loss-of-function, loss of expression or allelic mutation language, "
            f"indicating that the RiceMind textual mechanism centers on pathway disturbance caused by reduced function or expression {citation}."
        ),
        "architecture_height": (
            f"At the phenotype level, the evidence connects those molecular descriptions with semi-dwarf/dwarf, plant height, culm length, internode and elongation terms, "
            f"framing {gene} through plant-height and architecture outcomes {citation}."
        ),
        "yield_lodging_breeding": (
            f"Breeding-oriented sentences connect {gene}'s semi-dwarf context with lodging resistance, fertilizer response, harvest index, yield or Green Revolution cultivar use, "
            f"so this section is best framed as a RiceMind-supported breeding-value chain rather than a single molecular event {citation}."
        ),
        "reproductive_growth": (
            f"Developmental sentences co-report {gene} with heading, panicle, grain, seed, spikelet, flower, mesocotyl or seedling traits, "
            f"providing RiceMind sentence-level evidence for developmental extensions of the main mechanism {citation}."
        ),
        "stress_defense": (
            f"Stress and resistance sentences co-report {gene} with stress, disease, resistance, immunity or defense terms; these records should be treated as candidate or exploratory biological contexts {citation}."
        ),
        "expression_regulation": (
            f"Expression and molecular-function sentences include encoding, enzyme, protein, expression, candidate gene, QTL or locus language, "
            f"which helps trace trait associations back to molecular identity, expression and candidate-region descriptions in RiceMind text {citation}."
        ),
    }
    return clauses.get(rule["key"], f"Sentence evidence repeatedly contains {rule['en']} signals {citation}.")


MECHANISM_TOPIC_RULES = [
    {
        "key": "bph_resistance_loci",
        "zh": "褐飞虱抗性基因、QTL 和抗性位点",
        "en": "Brown planthopper resistance genes, QTLs and loci",
        "keywords": [
            "brown planthopper",
            "nilaparvata lugens",
            "bph resistance",
            "resistance to bph",
            "bph1",
            "bph2",
            "bph3",
            "bph4",
            "bph5",
            "bph6",
            "bph7",
            "bph8",
            "bph9",
            "bph10",
            "bph11",
            "bph12",
            "bph13",
            "bph14",
            "bph15",
            "bph17",
            "bph18",
            "bph26",
            "bph29",
            "bph30",
            "bph32",
            "bph38",
            "qtl",
            "locus",
            "resistance gene",
        ],
    },
    {
        "key": "bph_defense_response",
        "zh": "褐飞虱取食诱导的防御响应和信号通路",
        "en": "BPH-feeding-induced defense responses and signaling",
        "keywords": [
            "feeding",
            "sucking",
            "phloem",
            "sieve",
            "honeydew",
            "defense",
            "defence",
            "jasmonic",
            "salicylic",
            "ethylene",
            "callose",
            "ros",
            "oxidative",
            "transcriptome",
            "differentially expressed",
            "expression",
            "metabolite",
            "secondary metabol",
        ],
    },
    {
        "key": "bph_breeding_introgression",
        "zh": "抗褐飞虱资源导入、聚合育种和抗性改良",
        "en": "BPH-resistance introgression, pyramiding and breeding improvement",
        "keywords": [
            "introgression",
            "pyramiding",
            "marker-assisted",
            "mas",
            "breeding",
            "cultivar",
            "resistant variety",
            "resistant varieties",
            "oryza officinalis",
            "wild rice",
            "near-isogenic",
            "backcross",
            "donor",
            "rice improvement",
        ],
    },
    {
        "key": "bph_candidate_genes",
        "zh": "候选抗虫基因、受体/凝集素和功能验证线索",
        "en": "Candidate resistance genes, receptors/lectins and functional evidence",
        "keywords": [
            "candidate gene",
            "nbs-lrr",
            "cc-nb-lrr",
            "nb-lrr",
            "receptor",
            "lectin",
            "gna",
            "protein",
            "transgenic",
            "rna-seq",
            "map-based cloning",
            "fine mapping",
            "cloned",
            "functional",
        ],
    },
    {
        "key": "ga_core",
        "zh": "SD1/GA20ox2 介导的赤霉素合成缺陷与半矮秆机制",
        "en": "SD1/GA20ox2-mediated GA biosynthesis and semi-dwarfism",
        "keywords": [
            "osga20ox2",
            "ga20ox-2",
            "ga20ox2",
            "ga 20-oxidase",
            "ga20 oxidase",
            "c20-oxidase",
            "gibberellin",
            "bioactive ga",
            "semi-dwarf",
            "semidwarf",
            "plant height",
        ],
    },
    {
        "key": "allele_function",
        "zh": "等位突变、功能缺失与表达调控",
        "en": "Allelic mutation, loss of function and expression regulation",
        "keywords": [
            "383",
            "382",
            "deletion",
            "loss-of-function",
            "loss of function",
            "null allele",
            "defective",
            "mutation",
            "mutant allele",
            "loss of expression",
            "promoter",
            "expression",
            "transcription",
        ],
    },
    {
        "key": "architecture_lodging",
        "zh": "株高、节间/秆形态与抗倒伏农艺链条",
        "en": "Plant height, internode/culm morphology and lodging-resistance agronomy",
        "keywords": [
            "plant height",
            "semi-dwarf",
            "semidwarf",
            "dwarf",
            "culm",
            "internode",
            "elongation",
            "lodging",
            "harvest index",
            "fertilizer",
            "breaking strength",
            "culm strength",
        ],
    },
    {
        "key": "yield_breeding",
        "zh": "产量、分蘖穗粒性状与育种利用",
        "en": "Yield, tiller/panicle/grain traits and breeding use",
        "keywords": [
            "yield",
            "grain",
            "panicle",
            "tiller",
            "spikelet",
            "fertility",
            "harvest index",
            "breeding",
            "cultivar",
            "green revolution",
            "high-yielding",
            "elite rice",
            "ir8",
        ],
    },
    {
        "key": "seed_development",
        "zh": "种子萌发、休眠和早期发育中的 GA 线索",
        "en": "GA-related evidence in seed germination, dormancy and early development",
        "keywords": [
            "seed germination",
            "germination",
            "dormancy",
            "seedling",
            "mesocotyl",
            "root",
            "shoot",
            "flowering",
            "heading",
            "aba",
            "ja",
        ],
    },
    {
        "key": "stress_growth_balance",
        "zh": "盐、旱、铝和养分胁迫中的生长-胁迫平衡线索",
        "en": "Growth-stress balance under salt, drought, aluminum and nutrient stress",
        "keywords": [
            "salt stress",
            "salt in the soil",
            "salinity",
            "nacl",
            "drought",
            "drought-stress",
            "aluminum stress",
            "alkali stress",
            "nutrient-deficiency",
            "nutrient deficiency",
            "stress tolerance",
            "osmotic",
            "aba",
            "root growth",
        ],
    },
    {
        "key": "candidate_networks",
        "zh": "QTL、候选基因和跨物种同源网络",
        "en": "QTL, candidate-gene and cross-species homolog networks",
        "keywords": [
            "qtl",
            "candidate gene",
            "candidate genes",
            "gwas",
            "homolog",
            "homologous",
            "ortholog",
            "syntenic",
            "locus",
            "haplotype",
            "comparative genomic",
        ],
    },
]


def rows_matching_keywords(rows: List[Dict[str, str]], keywords: Sequence[str]) -> List[Dict[str, str]]:
    matched = []
    lowered_keywords = [keyword.lower() for keyword in keywords]
    for row in rows:
        haystack = f"{row.get('trait', '')} {row.get('sentence', '')}".lower()
        if any(keyword in haystack for keyword in lowered_keywords):
            matched.append(row)
    return matched


def group_rows_by_mechanism_topic(rows: List[Dict[str, str]]) -> List[Tuple[Dict[str, Any], List[Dict[str, str]]]]:
    grouped = []
    for topic in MECHANISM_TOPIC_RULES:
        topic_rows = rows_matching_keywords(rows, topic["keywords"])
        if topic_rows:
            grouped.append((topic, topic_rows))
    grouped.sort(key=lambda item: len(item[1]), reverse=True)
    return grouped[:REPORT_MECHANISM_THEMES]


def confidence_distribution_text(rows: List[Dict[str, str]]) -> str:
    counter = Counter(row.get("confidence", "") or "Unspecified" for row in rows)
    return ", ".join(f"{tier}={count}" for tier, count in counter.most_common())


def has_any(rows: List[Dict[str, str]], keywords: Sequence[str]) -> bool:
    return bool(rows_matching_keywords(rows, keywords))


def rows_for_signal(rows: List[Dict[str, str]], keywords: Sequence[str]) -> List[Dict[str, str]]:
    return rows_matching_keywords(rows, keywords)


def observed_terms(rows: List[Dict[str, str]], terms: Sequence[str], limit: int = 8) -> List[str]:
    counter: Counter = Counter()
    for row in rows:
        haystack = f"{row.get('trait', '')} {row.get('sentence', '')}".lower()
        for term in terms:
            if term.lower() in haystack:
                counter[term] += 1
    return [term for term, _ in counter.most_common(limit)]


def join_cn(items: Sequence[str]) -> str:
    return "、".join(items)


def join_en(items: Sequence[str]) -> str:
    return ", ".join(items)


def add_claim(paragraphs: List[str], text: str) -> None:
    cleaned = " ".join(text.split())
    if cleaned:
        paragraphs.append(cleaned)


def build_bph_review_paragraphs(gene: str, topic: Dict[str, Any], rows: List[Dict[str, str]], is_zh: bool) -> List[str]:
    paragraphs: List[str] = []
    unique_pmids = len({row.get("pmid", "") for row in rows if row.get("pmid", "")})
    citation = pmid_citation(collect_pmids(rows, 8))
    tier_text = confidence_distribution_text(rows)
    core_terms = observed_terms(
        rows,
        [
            "BPH",
            "brown planthopper",
            "Nilaparvata lugens",
            "Bph14",
            "Bph15",
            "Bph18",
            "Bph26",
            "Bph3",
            "QTL",
            "resistance gene",
            "Oryza officinalis",
            "feeding",
            "phloem",
            "honeydew",
            "NBS-LRR",
            "lectin",
            "GNA",
            "introgression",
            "pyramiding",
        ],
        10,
    )
    if is_zh:
        add_claim(
            paragraphs,
            f"RiceMind 在“{topic['zh']}”主题下提供 {len(rows)} 条句子级证据，覆盖 {unique_pmids} 个 PMID，置信层级分布为 {tier_text}。"
            f"需要特别注意，BPH 在文献句子中常同时指 brown planthopper 这个害虫/性状语境和 BPH/Bph 抗性位点；因此本节把它作为褐飞虱抗性证据集合来综述，"
            f"不把所有句子都强行解释为单一已克隆基因的功能。该证据集合中反复出现的核心词包括 {join_cn(core_terms) if core_terms else '褐飞虱、抗性、QTL 和候选基因'} {citation}。",
        )
    else:
        add_claim(
            paragraphs,
            f"RiceMind provides {len(rows)} sentence-level records for {topic['en']}, covering {unique_pmids} PMIDs with confidence distribution {tier_text}. "
            f"BPH in the text often denotes both brown planthopper as the pest/trait context and BPH/Bph resistance loci; therefore this section synthesizes a BPH-resistance evidence set rather than forcing every sentence into a single cloned-gene function. "
            f"Repeated terms include {join_en(core_terms) if core_terms else 'brown planthopper, resistance, QTL and candidate genes'} {citation}.",
        )

    loci_rows = rows_for_signal(rows, ["bph1", "bph2", "bph3", "bph4", "bph5", "bph6", "bph7", "bph8", "bph9", "bph10", "bph11", "bph12", "bph13", "bph14", "bph15", "bph17", "bph18", "bph26", "bph29", "bph30", "bph32", "bph38", "qtl", "locus", "resistance gene"])
    feeding_rows = rows_for_signal(rows, ["feeding", "sucking", "phloem", "sieve", "honeydew", "nymph", "adult", "oviposition", "survival", "antibiosis", "antixenosis", "tolerance"])
    signaling_rows = rows_for_signal(rows, ["defense", "defence", "jasmonic", "salicylic", "ethylene", "callose", "ros", "oxidative", "transcriptome", "differentially expressed", "expression", "metabolite", "secondary metabol"])
    breeding_rows = rows_for_signal(rows, ["introgression", "pyramiding", "marker-assisted", "mas", "breeding", "cultivar", "resistant variety", "resistant varieties", "oryza officinalis", "wild rice", "near-isogenic", "backcross", "donor", "rice improvement"])
    candidate_rows = rows_for_signal(rows, ["candidate gene", "nbs-lrr", "cc-nb-lrr", "nb-lrr", "receptor", "lectin", "gna", "transgenic", "map-based cloning", "fine mapping", "cloned", "functional"])

    plan = {
        "bph_resistance_loci": {"loci", "candidate", "breeding"},
        "bph_defense_response": {"feeding", "signaling", "candidate"},
        "bph_breeding_introgression": {"breeding", "loci"},
        "bph_candidate_genes": {"candidate", "signaling", "loci"},
    }.get(topic.get("key", ""), {"loci", "feeding", "signaling", "breeding", "candidate"})

    if "loci" in plan and loci_rows:
        c = pmid_citation(collect_pmids(loci_rows, 7))
        loci_terms = observed_terms(loci_rows, ["Bph1", "Bph2", "Bph3", "Bph10", "Bph11", "Bph12", "Bph14", "Bph15", "Bph18", "Bph26", "Bph29", "Bph30", "Bph32", "BPH38", "QTL", "locus", "resistance gene"], 12)
        if is_zh:
            add_claim(
                paragraphs,
                f"遗传定位层面的句子把褐飞虱抗性组织为多位点体系，而不是单一 trait 标签：{join_cn(loci_terms) if loci_terms else '多个 Bph/BPH 位点、QTL 和 resistance gene'} 在文献证据中反复出现。"
                f"这些句子提示，RiceMind 中的 BPH 抗性机制首先体现为抗性位点发现、精细定位、候选区间缩小和已报道抗性基因资源的累积；报告应将其写成抗性遗传资源网络，而不是把 BPH 简化成一个孤立基因 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"At the genetic-mapping level, the sentences organize brown planthopper resistance as a multi-locus system rather than a single trait label: {join_en(loci_terms) if loci_terms else 'multiple Bph/BPH loci, QTLs and resistance genes'} recur in the evidence. "
                f"This indicates that the RiceMind BPH-resistance mechanism starts from locus discovery, fine mapping, candidate-region narrowing and accumulation of reported resistance genes, best described as a resistance-resource network rather than one isolated gene {c}.",
            )

    if "feeding" in plan and feeding_rows:
        c = pmid_citation(collect_pmids(feeding_rows, 7))
        feeding_terms = observed_terms(feeding_rows, ["feeding", "phloem", "sieve", "honeydew", "nymph", "adult", "oviposition", "survival", "antibiosis", "antixenosis", "tolerance"], 10)
        if is_zh:
            add_claim(
                paragraphs,
                f"害虫-寄主互作层面的句子把抗性表型具体化为取食和存活过程：{join_cn(feeding_terms) if feeding_terms else '取食、蜜露、若虫/成虫表现和存活'} 等词汇把抗性从“有/无抗性”推进到褐飞虱在水稻上的取食行为、韧皮部利用、繁殖或存活结果。"
                f"因此，BPH 相关证据可被综述为寄主因素改变褐飞虱取食适合度和群体表现的文本证据链 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"At the pest-host interaction level, sentences make resistance phenotypes concrete through feeding and survival processes: {join_en(feeding_terms) if feeding_terms else 'feeding, honeydew, nymph/adult performance and survival'} connect resistance to planthopper behavior, phloem use, reproduction or survival on rice. "
                f"The BPH evidence can therefore be synthesized as a text-derived chain in which host factors alter brown-planthopper feeding fitness and population performance {c}.",
            )

    if "signaling" in plan and signaling_rows:
        c = pmid_citation(collect_pmids(signaling_rows, 7))
        signal_terms = observed_terms(signaling_rows, ["defense", "jasmonic", "salicylic", "ethylene", "callose", "ROS", "oxidative", "transcriptome", "differentially expressed", "expression", "metabolite", "secondary metabol"], 10)
        if is_zh:
            add_claim(
                paragraphs,
                f"防御响应层面的句子进一步把 BPH 抗性连接到诱导表达和激素/代谢网络：{join_cn(signal_terms) if signal_terms else 'defense、表达变化、激素和代谢'} 等词汇提示，褐飞虱取食后的水稻响应不只是结构性抗性，"
                f"还可能包括转录重编程、防御相关激素信号、氧化/ROS 过程、胼胝质或次生代谢调节。若这些句子主要来自 NLP 共现，应写成候选防御机制，而不是已验证因果通路 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"Defense-response sentences connect BPH resistance with induced expression and hormone/metabolic networks: {join_en(signal_terms) if signal_terms else 'defense, expression change, hormones and metabolism'} suggest that rice responses to planthopper feeding may include transcriptional reprogramming, defense hormone signaling, oxidative/ROS processes, callose or secondary metabolism. "
                f"When supported mainly by NLP co-occurrence, this should be written as a candidate defense mechanism rather than a validated causal pathway {c}.",
            )

    if "breeding" in plan and breeding_rows:
        c = pmid_citation(collect_pmids(breeding_rows, 7))
        breeding_terms = observed_terms(breeding_rows, ["introgression", "pyramiding", "marker-assisted", "MAS", "breeding", "Oryza officinalis", "wild rice", "near-isogenic", "backcross", "donor", "resistant variety"], 10)
        if is_zh:
            add_claim(
                paragraphs,
                f"育种利用层面的句子显示，BPH 抗性证据常与 {join_cn(breeding_terms) if breeding_terms else '抗性资源导入、回交、聚合和分子标记辅助选择'} 相连。"
                f"这说明 RiceMind 文本中的机制外延不仅是分子防御，还包括从野生稻或抗性供体导入 Bph 位点、在栽培背景中聚合多个抗性基因、并通过近等基因系或回交群体验证抗性的育种流程 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"Breeding-use sentences connect BPH resistance with {join_en(breeding_terms) if breeding_terms else 'introgression, backcrossing, pyramiding and marker-assisted selection'}. "
                f"Thus, the RiceMind mechanism extends beyond molecular defense to the breeding workflow of introgressing Bph loci from wild rice or resistant donors, pyramiding resistance genes in cultivated backgrounds and validating resistance in near-isogenic or backcross populations {c}.",
            )

    if "candidate" in plan and candidate_rows:
        c = pmid_citation(collect_pmids(candidate_rows, 7))
        candidate_terms = observed_terms(candidate_rows, ["candidate gene", "NBS-LRR", "CC-NB-LRR", "receptor", "lectin", "GNA", "transgenic", "map-based cloning", "fine mapping", "cloned", "functional"], 10)
        if is_zh:
            add_claim(
                paragraphs,
                f"候选基因和功能验证层面的句子出现 {join_cn(candidate_terms) if candidate_terms else '候选基因、受体、NBS-LRR/凝集素和转基因验证'} 等线索。"
                f"这些证据可将 BPH 抗性从遗传区间进一步推进到候选蛋白类型和实验系统：受体样蛋白、NBS-LRR 类抗性蛋白、凝集素/转基因表达或图位克隆等描述，为后续功能实验提供更具体的机制入口 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"Candidate-gene and functional-evidence sentences contain {join_en(candidate_terms) if candidate_terms else 'candidate genes, receptors, NBS-LRR/lectins and transgenic validation'} signals. "
                f"These records move BPH resistance from genetic intervals toward candidate protein classes and experimental systems, including receptor-like proteins, NBS-LRR resistance proteins, lectin/transgenic expression or map-based cloning, providing more specific entry points for functional validation {c}.",
            )

    if len(paragraphs) == 1:
        add_claim(
            paragraphs,
            zh(
                is_zh,
                "该 BPH 主题尚未形成可稳定抽取的机制链条；完整原始证据已保存在 normalized_evidence.csv 以供人工追溯。",
                "This BPH topic did not yield a stable mechanism chain; full original evidence is preserved in normalized_evidence.csv for manual tracing.",
            ),
        )
    return paragraphs


def build_mechanism_review_paragraphs(gene: str, topic: Dict[str, Any], rows: List[Dict[str, str]], is_zh: bool) -> List[str]:
    if stringify(topic.get("key", "")).startswith("bph_"):
        return build_bph_review_paragraphs(gene, topic, rows, is_zh)

    paragraphs: List[str] = []
    unique_pmids = len({row.get("pmid", "") for row in rows if row.get("pmid", "")})
    citation = pmid_citation(collect_pmids(rows, 8))
    tier_text = confidence_distribution_text(rows)
    core_terms = observed_terms(
        rows,
        [
            "SD1",
            "sd1",
            "OsGA20ox2",
            "GA20ox-2",
            "GA20 oxidase",
            "C20-oxidase",
            "gibberellin",
            "bioactive GA",
            "SLR1",
            "DELLA",
            "GID1",
            "ABA",
            "JA",
            "QTL",
            "lodging",
            "semi-dwarf",
            "plant height",
        ],
        10,
    )

    if is_zh:
        add_claim(
            paragraphs,
            f"RiceMind 在“{topic['zh']}”主题下提供 {len(rows)} 条句子级证据，覆盖 {unique_pmids} 个 PMID，置信层级分布为 {tier_text}。"
            f"这些句子的价值不在于再次证明若干 trait 与 {gene} 共现，而在于它们共同给出了可追溯的文本证据链："
            f"{join_cn(core_terms) if core_terms else '相关分子、通路和表型词汇'} 在同一证据集合中反复出现，使 {gene} 的机制描述可以从 RiceMind 原始句子中组织出来 {citation}。",
        )
    else:
        add_claim(
            paragraphs,
            f"RiceMind provides {len(rows)} sentence-level records for {topic['en']}, covering {unique_pmids} PMIDs with confidence distribution {tier_text}. "
            f"The value of these sentences is not another trait list, but a traceable text-derived chain in which "
            f"{join_en(core_terms) if core_terms else 'molecular, pathway and phenotype terms'} repeatedly co-occur and support a RiceMind-grounded mechanism narrative for {gene} {citation}.",
        )

    ga_rows = rows_for_signal(rows, ["gibberellin", "ga20", "ga 20-oxidase", "c20-oxidase", "ga biosynthesis", "ga synthesis", "bioactive ga", "ga pathway", "ga homeostasis"])
    loss_rows = rows_for_signal(rows, ["383", "382", "deletion", "loss-of-function", "loss of function", "null allele", "defective", "loss of expression", "reduced expression", "mutation", "mutant allele"])
    della_rows = rows_for_signal(rows, ["slr1", "della", "gid1", "ga signal", "ga signaling", "ga-triggered degradation"])
    architecture_rows = rows_for_signal(rows, ["plant height", "semi-dwarf", "semidwarf", "dwarf", "culm", "internode", "elongation", "cell elongation", "shorter culm"])
    agronomy_rows = rows_for_signal(rows, ["lodging", "yield", "harvest index", "fertilizer", "high-yielding", "green revolution", "breeding", "cultivar", "culm strength", "breaking strength", "compromised yield"])
    development_rows = rows_for_signal(rows, ["seed germination", "germination", "dormancy", "mesocotyl", "seedling", "panicle", "grain", "spikelet", "heading", "flower", "tiller"])
    stress_rows = rows_for_signal(rows, ["salt stress", "salt in the soil", "salinity", "nacl", "drought", "drought-stress", "aluminum stress", "alkali stress", "nutrient-deficiency", "nutrient deficiency", "root growth", "stress tolerance", "osmotic"])
    regulation_rows = rows_for_signal(rows, ["promoter", "expression", "transcription", "qtl", "candidate gene", "gwas", "homolog", "ortholog", "syntenic", "locus", "haplotype", "oseil1a", "osphq1", "bst", "bzip72", "zfp207", "osstp28"])
    claim_plan = {
        "ga_core": {"ga", "loss", "della", "architecture"},
        "allele_function": {"loss", "ga", "regulation"},
        "architecture_lodging": {"architecture", "agronomy", "ga", "loss"},
        "yield_breeding": {"agronomy", "development", "architecture"},
        "seed_development": {"development", "ga", "regulation"},
        "stress_growth_balance": {"stress", "ga", "della"},
        "candidate_networks": {"regulation", "ga", "loss"},
    }.get(topic.get("key", ""), {"ga", "loss", "della", "architecture", "agronomy", "development", "stress", "regulation"})

    if "ga" in claim_plan and ga_rows:
        c = pmid_citation(collect_pmids(ga_rows, 7))
        if is_zh:
            late_step = "、并多次具体到 GA 通路的 late step / second-to-last step" if has_any(ga_rows, ["late stage", "late steps", "second-to-last"]) else ""
            add_claim(
                paragraphs,
                f"在分子层面，证据句将 {gene} 具体定位到 SD1/OsGA20ox2/GA20ox-2、GA20 oxidase 或 C20-oxidase 等身份，而不是泛泛的激素关联；"
                f"这些句子把它放入 gibberellin/GA biosynthesis、GA synthesis、bioactive GA 或 GA homeostasis 的语境{late_step}。"
                f"因此，RiceMind 文本支持的核心机制可以概括为：{gene} 通过 GA20-oxidase 相关活性影响活性赤霉素供应，从而为后续株高和发育表型提供上游通路解释 {c}。",
            )
        else:
            late_step = ", often specifying late or second-to-last GA biosynthetic steps" if has_any(ga_rows, ["late stage", "late steps", "second-to-last"]) else ""
            add_claim(
                paragraphs,
                f"At the molecular level, the sentences identify {gene} as SD1/OsGA20ox2/GA20ox-2, GA20 oxidase or C20-oxidase rather than a generic hormone association. "
                f"They place it in gibberellin/GA biosynthesis, GA synthesis, bioactive GA or GA homeostasis contexts{late_step}. "
                f"Thus, the RiceMind-supported core mechanism is that {gene} affects bioactive gibberellin supply through GA20-oxidase-related activity, providing the upstream pathway basis for height and developmental phenotypes {c}.",
            )

    if "loss" in claim_plan and loss_rows:
        c = pmid_citation(collect_pmids(loss_rows, 7))
        deletion_note = "383/382-bp deletion、" if has_any(loss_rows, ["383", "382"]) else ""
        expression_note = "loss of expression / reduced expression、" if has_any(loss_rows, ["loss of expression", "reduced expression"]) else ""
        if is_zh:
            add_claim(
                paragraphs,
                f"在等位变异层面，句子证据进一步把 {deletion_note}{expression_note}loss-of-function、null allele、defective enzyme 或 mutant allele 与 {gene} 联系起来。"
                f"这些描述共同指向同一个文本机制：等位突变或表达降低削弱 GA20-oxidase 酶功能，导致 GA 合成受阻或 bioactive/cellular GA 水平下降；"
                f"这比单纯列出“dwarf”“height”这样的 trait 更接近 RiceMind 句子证据中的因果链条 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"At the allelic level, the evidence links {gene} to {deletion_note}{expression_note}loss-of-function, null alleles, defective enzymes or mutant alleles. "
                f"Together these sentences support a text-derived mechanism in which allelic mutation or reduced expression weakens GA20-oxidase function, blocks GA synthesis or lowers bioactive/cellular GA, which is closer to a causal chain than merely listing dwarf or height traits {c}.",
            )

    if "della" in claim_plan and della_rows:
        c = pmid_citation(collect_pmids(della_rows, 7))
        if is_zh:
            add_claim(
                paragraphs,
                f"部分证据句把这一 GA 合成模块继续连接到 GA signal transduction：其中出现了 DELLA、SLR1、GID1 或 GA-triggered degradation 等词汇。"
                f"这些文本线索支持这样一种 RiceMind 内部可追溯的解释：当 {gene}/OsGA20ox2 相关 GA 合成下降时，GA 信号输出和 DELLA/SLR1 负调控模块随之被牵动，"
                f"从而把“酶功能缺失”连接到“生长受限”的激素信号层 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"Some evidence extends the GA biosynthetic module into GA signal transduction through DELLA, SLR1, GID1 or GA-triggered degradation language. "
                f"These RiceMind text signals support an interpretation in which reduced {gene}/OsGA20ox2-dependent GA biosynthesis perturbs GA output and the DELLA/SLR1 growth-repression module, connecting enzyme disruption to constrained growth {c}.",
            )

    if "architecture" in claim_plan and architecture_rows:
        c = pmid_citation(collect_pmids(architecture_rows, 7))
        if is_zh:
            add_claim(
                paragraphs,
                f"在表型层面，句子证据把上述分子扰动与 plant height、semi-dwarf/dwarf、culm length、internode、cell elongation 等词汇相连。"
                f"因此，{gene} 的表型机制不是孤立的“矮化”标签，而是“GA20-oxidase 功能变化-活性 GA 供应变化-细胞或节间伸长受限-株高下降”的连续链条；"
                f"在 High 证据或含 curated evidence code 的 trait 中，这一链条尤其适合用作较保守的机制表述 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"At the phenotype level, the evidence connects the molecular disturbance with plant height, semi-dwarf/dwarf, culm length, internode and cell-elongation language. "
                f"The mechanism is therefore not an isolated dwarf label, but a chain from altered GA20-oxidase function to changed active GA supply, constrained cell or internode elongation and reduced plant height; curated or High evidence can support this conservative wording {c}.",
            )

    if "agronomy" in claim_plan and agronomy_rows:
        c = pmid_citation(collect_pmids(agronomy_rows, 7))
        caution = "同时，部分句子也提到 culm strength 降低、fine/weak culm 或 excessive dwarfing/compromised yield，提示 sd1 等位型并非越强越好，育种利用需要在抗倒伏、株型和产量之间平衡。" if has_any(agronomy_rows, ["culm strength", "weak", "fine culm", "excessive dwarfing", "compromised yield"]) else ""
        if is_zh:
            add_claim(
                paragraphs,
                f"在农艺价值层面，RiceMind 句子将半矮化机制与 lodging resistance、fertilizer response、harvest index、yield、Green Revolution 以及现代高产品种利用相连。"
                f"这些证据把 {gene} 从单一生理机制延伸为育种链条：降低株高有助于降低倒伏风险，使高肥条件下的收获指数和产量潜力更容易被利用。"
                f"{caution} {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"At the agronomic level, RiceMind sentences connect semi-dwarfism with lodging resistance, fertilizer response, harvest index, yield, the Green Revolution and modern high-yielding cultivar use. "
                f"This extends {gene} from a physiological mechanism into a breeding chain: reduced height lowers lodging risk and enables harvest-index and yield gains under fertilized conditions. "
                f"{caution} {c}.",
            )

    if "development" in claim_plan and development_rows:
        c = pmid_citation(collect_pmids(development_rows, 7))
        if is_zh:
            add_claim(
                paragraphs,
                f"在发育阶段层面，证据句还把 {gene}/OsGA20ox2 相关 GA 调控与 seed germination、dormancy、mesocotyl、seedling、panicle、grain、spikelet 或 tiller 等语境连接。"
                f"这些句子提示 sd1 的机制边界可从株高扩展到 GA/ABA/JA 相关的萌发、休眠和早期生长调控，但这类扩展常由 NLP 共现或候选基因证据支持，"
                f"应表述为 RiceMind 数据中的候选发育机制线索 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"At developmental stages, evidence links {gene}/OsGA20ox2-related GA regulation with seed germination, dormancy, mesocotyl, seedling, panicle, grain, spikelet or tiller contexts. "
                f"These sentences extend the sd1 mechanism beyond plant height into GA/ABA/JA-related germination, dormancy and early growth regulation, but this extension is often NLP or candidate-gene supported and should be phrased as a RiceMind candidate mechanism {c}.",
            )

    if "stress" in claim_plan and stress_rows:
        c = pmid_citation(collect_pmids(stress_rows, 7))
        stress_terms = observed_terms(stress_rows, ["salt stress", "salinity", "drought", "aluminum stress", "alkali stress", "root growth", "ABA", "GA20ox", "SD1-OE", "stress tolerance", "nutrient-deficiency"], 8)
        if is_zh:
            add_claim(
                paragraphs,
                f"逆境相关证据更适合写成“生长-胁迫平衡”的候选线索，而不是直接断言 {gene} 是抗盐或抗旱因果基因。"
                f"在这些句子中，{join_cn(stress_terms) if stress_terms else '盐、旱、铝、养分和根系生长'} 等语境与 GA20ox/SD1 或 GA 信号共同出现；"
                f"例如有证据把盐胁迫转录组差异基因连接到 plant hormone metabolism/signaling 和 GA20-oxidase，也有证据在铝胁迫或根生长背景下讨论 SD1/GA biosynthesis。"
                f"因此，该部分应总结为：RiceMind 文本提示 sd1 介导的 GA 生长模块可能参与逆境下的生长分配或激素响应，但低置信度/NLP 证据不宜被写成已验证的抗逆机制 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"Stress-related evidence is best written as a candidate growth-stress balance signal rather than a direct claim that {gene} is a causal salt- or drought-resistance gene. "
                f"Terms such as {join_en(stress_terms) if stress_terms else 'salt, drought, aluminum, nutrient stress and root growth'} co-occur with GA20ox/SD1 or GA signaling; for example, some sentences connect salt-stress transcriptomic DEGs with plant hormone metabolism/signaling and GA20-oxidase, while others discuss SD1/GA biosynthesis under aluminum stress or root-growth contexts. "
                f"Thus, RiceMind text suggests that the sd1-mediated GA growth module may participate in stress-time growth allocation or hormonal response, but Low/NLP evidence should not be written as a validated stress-resistance mechanism {c}.",
            )

    if "regulation" in claim_plan and regulation_rows:
        c = pmid_citation(collect_pmids(regulation_rows, 7))
        regulators = observed_terms(regulation_rows, ["OsEIL1a", "OsPHQ1", "GID1", "SLR1", "bZIP72", "ZFP207", "OsSTP28", "QTL", "GWAS", "candidate gene", "homolog", "promoter"], 10)
        if is_zh:
            add_claim(
                paragraphs,
                f"表达调控和候选位点证据提供了机制网络的另一层：句子中出现 {join_cn(regulators) if regulators else 'promoter、QTL、candidate gene、homolog 或 expression'} 等线索。"
                f"这些证据可以把 {gene} 与上游转录调控、QTL/候选区间、同源基因或跨物种比较联系起来，说明 RiceMind 中的 sd1 机制并不只是一条单基因线性路径，"
                f"还包括可用于后续实验验证的调控节点和遗传背景差异 {c}。",
            )
        else:
            add_claim(
                paragraphs,
                f"Expression and candidate-locus evidence adds another network layer: sentences contain {join_en(regulators) if regulators else 'promoter, QTL, candidate gene, homolog or expression'} terms. "
                f"These records connect {gene} with upstream transcriptional regulation, QTL/candidate regions, homologs or cross-species comparisons, indicating that the RiceMind sd1 mechanism is not only a single-gene linear path but also a set of regulatory nodes and genetic-background differences for follow-up validation {c}.",
            )

    if len(paragraphs) == 1:
        if is_zh:
            add_claim(
                paragraphs,
                f"该主题的证据句尚未形成更具体的机制链条；报告保留所有原始句子、PMID、trait 和 sentence ID 到 normalized_evidence.csv，以便人工追溯和二次综述。",
            )
        else:
            add_claim(
                paragraphs,
                "This topic did not yield a more specific mechanism chain; all original sentences, PMIDs, traits and sentence IDs remain in normalized_evidence.csv for manual tracing and secondary synthesis.",
            )
    return paragraphs


def add_mechanism_synthesis(doc: Any, gene: str, traits: List[Dict[str, str]], evidence: List[Dict[str, str]], is_zh: bool) -> None:
    if not evidence:
        doc.add_paragraph(
            zh(
                is_zh,
                "输入数据未提供句子级证据；不能生成 PMID 支撑的机制综述。完整句子证据应通过 normalized_evidence.csv 保存。",
                "No sentence-level evidence was available; PMID-supported mechanism synthesis cannot be generated. Full sentence evidence should be preserved in normalized_evidence.csv.",
            )
        )
        return

    ranked_themes = group_rows_by_mechanism_topic(evidence)
    intro = zh(
        is_zh,
        "以下综述段落以全量 RiceMind sentence evidence 为核心证据，不再逐层罗列 trait 名称。生成逻辑是先把句子按机制主题聚合，再从原句中抽取“分子身份-等位突变/表达变化-通路扰动-激素信号-表型/农艺结果”的证据链，并在每个机制判断后保留 PMID。完整原句、trait、PMID 和 sentence ID 保存在 normalized_evidence.csv。",
        "The following synthesis uses the full RiceMind sentence evidence as the evidence substrate rather than listing trait names tier by tier. Sentences are grouped by mechanism topic, then converted into a text-derived chain from molecular identity, allelic/expression change, pathway perturbation and hormonal signaling to phenotype or agronomic outcome, with PMID traceability. Full original sentences, traits, PMIDs and sentence IDs are preserved in normalized_evidence.csv.",
    )
    doc.add_paragraph(intro)

    for topic, rows in ranked_themes:
        doc.add_heading(topic["zh"] if is_zh else topic["en"], 2)
        for paragraph in build_mechanism_review_paragraphs(gene, topic, rows, is_zh):
            doc.add_paragraph(paragraph)


def build_docx(
    gene: str,
    bundle: Dict[str, Any],
    traits: List[Dict[str, str]],
    evidence: List[Dict[str, str]],
    varieties: List[Dict[str, str]],
    out_path: Path,
    fig_paths: Dict[str, Optional[Path]],
    language: str,
    max_evidence_rows: int,
    max_trait_rows: int,
) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx") from exc

    is_zh = language.lower().startswith("zh")
    profile = choose_profile(bundle)
    doc = Document()
    configure_document_styles(doc)
    doc.add_heading(zh(is_zh, f"{gene} 的 RiceMind 基因机制证据报告", f"RiceMind Gene Mechanism Evidence Report: {gene}"), 0)
    doc.add_paragraph(zh(is_zh, f"生成日期：{date.today().isoformat()}", f"Generated: {date.today().isoformat()}"))
    doc.add_paragraph(
        zh(
            is_zh,
            "数据来源：RiceMind REST API / MCP payload。本文档采用固定模板生成；所有机制性描述均限制在 RiceMind 返回的基因名片、GTA 元信息、证据代码、PMID、句子证据和上下文范围内。",
            "Data source: RiceMind REST API / MCP payload. This document follows the fixed report template; all mechanistic statements are restricted to RiceMind gene profile data, GTA metadata, evidence codes, PMIDs, sentence evidence and retrieved context.",
        )
    )

    doc.add_heading(zh(is_zh, "1. 数据来源、检索范围与完整性", "1. Data Source, Retrieval Scope and Completeness"), 1)
    doc.add_paragraph(
        zh(
            is_zh,
            f"API base URL：{bundle.get('api_base', 'input JSON bundle')}。本次规范化得到 {len(traits)} 条 gene-trait 关联、{len(evidence)} 条句子级证据、{len(varieties)} 个共现品种记录。",
            f"API base URL: {bundle.get('api_base', 'input JSON bundle')}. Normalization yielded {len(traits)} gene-trait associations, {len(evidence)} sentence-level evidence records and {len(varieties)} co-reported variety records.",
        )
    )
    api_calls = compact_api_calls(bundle.get("api_calls") or [])
    if api_calls:
        add_table(doc, api_call_rows(api_calls, REPORT_API_ROWS))
        doc.add_paragraph(
            zh(
                is_zh,
                f"上表仅展示最多 {REPORT_API_ROWS} 类功能不同的 API 调用摘要；分页访问已按 endpoint 和查询功能聚合，不逐页罗列。完整聚合后的调用摘要与原始返回数据保存在 payload JSON sidecar 中。",
                f"The table shows at most {REPORT_API_ROWS} functionally distinct API call summaries. Paginated calls are aggregated by endpoint and query function rather than listed page by page. The compact call summary and raw returned data are preserved in the payload JSON sidecar.",
            )
        )
    error_rows = endpoint_error_rows(bundle)
    if error_rows:
        doc.add_paragraph(
            zh(
                is_zh,
                "以下 endpoint 在本次检索中返回错误或不可用；报告继续使用其他可用 RiceMind 数据生成，并在 payload JSON 中保留错误信息。",
                "The following endpoints returned errors or were unavailable in this retrieval. The report continues with other available RiceMind data and preserves error details in the payload JSON.",
            )
        )
        add_table(doc, error_rows)

    doc.add_heading(zh(is_zh, "2. 基因概貌与外部链接", "2. Gene Overview and External Links"), 1)
    if profile:
        add_table(doc, flatten_profile_rows(profile))
    else:
        doc.add_paragraph(zh(is_zh, "未在输入数据中找到 gene-profile 记录。", "No gene-profile record was found in the input data."))

    doc.add_heading(zh(is_zh, "3. 全量句证输出与本体分布", "3. Evidence Distribution and Sentence Provenance"), 1)
    add_table(doc, evidence_summary_rows(traits, evidence))
    if traits or evidence:
        doc.add_heading(zh(is_zh, "本体维度分布", "Ontology Mapping Distribution"), 2)
        add_table(doc, ontology_distribution_rows(traits, evidence))
    if evidence:
        doc.add_heading(zh(is_zh, "代表性句子证据", "Representative Sentence Evidence"), 2)
        doc.add_paragraph(
            zh(
                is_zh,
                "下表按置信层级抽取代表性 RiceMind 句子证据，仅用于展示证据形态；全量句证保存在 normalized_evidence.csv。",
                "The table samples representative RiceMind sentence evidence by confidence tier only to show evidence shape; full sentence evidence is preserved in normalized_evidence.csv.",
            )
        )
        add_table(doc, representative_sentence_rows(evidence))

    doc.add_heading(zh(is_zh, "4. 全量 GTA 概貌", "4. Full GTA Landscape"), 1)
    add_picture_if_exists(doc, fig_paths.get("confidence"), zh(is_zh, "图 1. RiceMind confidence tier 分布。", "Figure 1. RiceMind confidence-tier distribution."))
    add_picture_if_exists(doc, fig_paths.get("ontology"), zh(is_zh, "图 2. Ontology 类型分布。", "Figure 2. Ontology type distribution."))
    add_picture_if_exists(doc, fig_paths.get("top_traits"), zh(is_zh, "图 3. 按文献支持数排序的主要 trait。", "Figure 3. Top traits ranked by supporting article count."))
    sorted_traits = sorted(
        traits,
        key=lambda r: ({"High": 0, "Medium": 1, "Low": 2}.get(r["confidence"], 3), -parse_int(r["support"]), r["trait"]),
    )
    trait_rows = [["Trait", "Ontology", "Confidence", "Evidence code", "Source", "Support", "Earliest year"]]
    trait_display_limit = max_trait_rows if max_trait_rows and max_trait_rows > 0 else REPORT_TRAIT_ROWS
    for row in row_limit(sorted_traits, trait_display_limit):
        trait_rows.append([
            truncate(row["trait"], 90),
            " ".join(x for x in [row["ontology_type"], row["ontology_id"]] if x),
            row["confidence"],
            truncate(row["evidence_code"], 80),
            truncate(row["source_db"], 80),
            row["support"],
            row["year"],
        ])
    add_table(doc, trait_rows)
    if len(sorted_traits) > trait_display_limit:
        doc.add_paragraph(
            zh(
                is_zh,
                f"报告正文仅展示支持度和置信层级靠前的 {trait_display_limit} 条 trait 记录；全部 {len(sorted_traits)} 条 trait 关联已保存到 normalized_traits.csv。",
                f"The report body shows only the top {trait_display_limit} trait records by confidence/support; all {len(sorted_traits)} trait associations are saved in normalized_traits.csv.",
            )
        )

    doc.add_heading(zh(is_zh, "5. 证据代码、来源与置信层级统计", "5. Evidence Codes, Sources and Confidence Statistics"), 1)
    add_picture_if_exists(doc, fig_paths.get("evidence_code"), zh(is_zh, "图 4. Evidence code 分布。", "Figure 4. Evidence-code distribution."))
    add_picture_if_exists(doc, fig_paths.get("source"), zh(is_zh, "图 5. RiceMind 来源数据库分布。", "Figure 5. RiceMind source database distribution."))
    summary_rows = [["Statistic", "Value"]]
    summary_rows.extend([
        ["High traits", str(sum(1 for row in traits if row["confidence"] == "High"))],
        ["Medium traits", str(sum(1 for row in traits if row["confidence"] == "Medium"))],
        ["Low traits", str(sum(1 for row in traits if row["confidence"] == "Low"))],
        ["Unique PMIDs", str(len({row["pmid"] for row in evidence if row["pmid"]}))],
        ["Evidence sentences", str(len(evidence))],
    ])
    add_table(doc, summary_rows)

    doc.add_heading(zh(is_zh, "6. RiceMind 句子证据驱动的机制综述", "6. RiceMind Sentence-Evidence-Driven Mechanism Synthesis"), 1)
    add_mechanism_synthesis(doc, gene, traits, evidence, is_zh)

    doc.add_heading(zh(is_zh, "7. 研究热点与年代变迁分析", "7. Temporal Analysis and Hotspot Evolution"), 1)
    temporal_rows = temporal_hotspot_rows(evidence, is_zh)
    if len(temporal_rows) > 1:
        add_table(doc, temporal_rows)
        doc.add_paragraph(
            zh(
                is_zh,
                "阶段划分由 RiceMind 句证年份自动分箱得到；热点词来自各阶段句子文本，因此反映研究关注度变化，不等同于生物学重要性排序。",
                "Phase boundaries are computed from RiceMind sentence years; hotspot terms come from phase-specific sentence text and therefore reflect attention shifts rather than biological-importance ranking.",
            )
        )
    else:
        doc.add_paragraph(zh(is_zh, "输入证据缺少年份信息，无法生成热点阶段分析。", "No usable publication years were available for hotspot phase analysis."))

    doc.add_heading(zh(is_zh, "8. 不一致或条件依赖的机制线索", "8. Conflicting or Context-Dependent Mechanistic Signals"), 1)
    detected_conflicts = conflict_rows(evidence, is_zh)
    if detected_conflicts:
        add_table(doc, detected_conflicts)
    else:
        doc.add_paragraph(
            zh(
                is_zh,
                "未在当前 RiceMind 句子证据中检测到足够强的正负方向词共存模式。没有检测到冲突并不代表文献完全一致，只表示本次规则化扫描未发现可稳定报告的条件依赖线索。",
                "No sufficiently strong positive/negative directional co-occurrence pattern was detected in the current RiceMind sentence evidence. Absence of detected conflict does not prove full literature agreement; it only means this rule-based scan did not find a stable reportable context-dependence signal.",
            )
        )

    doc.add_heading(zh(is_zh, "9. 二级文献计量与 PMID 可追溯性", "9. Secondary Bibliometrics and PMID Traceability"), 1)
    add_picture_if_exists(doc, fig_paths.get("years"), zh(is_zh, "图 6. 句子级证据的发表年份分布。", "Figure 6. Publication-year distribution of sentence-level evidence."))
    if not fig_paths.get("years"):
        doc.add_paragraph(zh(is_zh, "输入证据缺少可用发表年份，未生成年度趋势图。", "No usable publication years were available for a trend chart."))
    journal_rows, pmid_rows = bibliometric_rows(evidence)
    if journal_rows:
        doc.add_heading(zh(is_zh, "Top 发表刊物", "Top Journals"), 2)
        add_table(doc, journal_rows)
    else:
        doc.add_paragraph(
            zh(
                is_zh,
                "当前 RiceMind 句证未提供可稳定统计的 journal 字段，因此不输出 Top 期刊排名。",
                "The current RiceMind sentence evidence does not provide stable journal metadata, so top-journal ranking is not reported.",
            )
        )
    if pmid_rows:
        doc.add_heading(zh(is_zh, "高频 PMID / 文献簇", "High-Frequency PMID / Article Clusters"), 2)
        add_table(doc, pmid_rows)
    doc.add_paragraph(
        zh(
            is_zh,
            "机构地理热点只有在 RiceMind 返回通讯作者单位或 affiliation 元数据时才可报告；若 payload 中没有该字段，报告不得凭常识或外部记忆补写机构地点。",
            "Institutional geography is reported only when RiceMind returns corresponding-author or affiliation metadata; if absent from the payload, the report must not fill locations from general knowledge or memory.",
        )
    )

    doc.add_heading(zh(is_zh, "10. 品种共现与组学序列信息", "10. Variety Co-Occurrence and Omics Sequence Information"), 1)
    if varieties:
        variety_rows = [["Variety"]]
        for row in varieties:
            variety_rows.append([row["variety"]])
        add_table(doc, variety_rows)
    else:
        doc.add_paragraph(zh(is_zh, "未检索到或未提供基因-品种共现记录。", "No gene-variety co-occurrence records were retrieved or provided."))
    sequence_rows = sequence_api_url_rows(bundle, profile, gene)
    if sequence_rows:
        doc.add_paragraph(
            zh(
                is_zh,
                "序列相关信息不在正文中展开为长表格；请直接追溯下列 RiceMind gene-omics-sequence API 超链接。完整接口返回仍保存在 payload JSON sidecar 中。",
                "Sequence information is not expanded into a long table in the report body. Use the following RiceMind gene-omics-sequence API hyperlink for direct tracing. The full endpoint response remains in the payload JSON sidecar.",
            )
        )
        add_table(doc, sequence_rows)
    elif bundle.get("omics_sequence_error"):
        doc.add_paragraph(zh(is_zh, f"组学序列接口未返回可用结果：{bundle['omics_sequence_error']}", f"Omics sequence endpoint did not return usable data: {bundle['omics_sequence_error']}"))

    doc.add_heading(zh(is_zh, "11. 证据边界与解释限制", "11. Evidence Boundaries and Interpretation Limits"), 1)
    doc.add_paragraph(
        zh(
            is_zh,
            "High/Tier 1 适合较保守的机制验证；Medium/Tier 2 更适合候选机制发现；Low/Tier 3 和 NLP-only 共现只能作为探索性线索。句子共现本身不能证明因果关系；报告中的机制表述必须回溯到本报告列出的 PMID、sentence ID 和 RiceMind evidence code。",
            "High/Tier 1 evidence is appropriate for conservative mechanism validation; Medium/Tier 2 is better treated as candidate discovery; Low/Tier 3 and NLP-only co-occurrence are exploratory. Sentence co-occurrence alone does not establish causality; mechanistic statements must trace back to the PMIDs, sentence IDs and RiceMind evidence codes listed in this report.",
        )
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    apply_document_fonts(doc)
    doc.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--gene", required=True, help="Target rice gene symbol, alias, or RAP ID")
    parser.add_argument("--input-json", type=Path, help="RiceMind API/MCP payload bundle")
    parser.add_argument("--api-base", default=DEFAULT_API_BASE, help=f"RiceMind REST API base URL, default: {DEFAULT_API_BASE}")
    parser.add_argument("--no-api", action="store_true", help="Do not call the live API; use only --input-json")
    parser.add_argument("--endpoint-map", type=Path, help="Optional JSON map overriding REST endpoint paths")
    parser.add_argument("--out", type=Path, help="Output DOCX path")
    parser.add_argument("--fig-dir", type=Path, help="Directory for generated figures")
    parser.add_argument("--language", default="zh", choices=["en", "zh"], help="Report language")
    parser.add_argument("--page-limit", type=int, default=DEFAULT_PAGE_LIMIT, help="API page size used for full pagination")
    parser.add_argument("--max-pages", type=int, default=DEFAULT_MAX_PAGES, help="Safety cap for paginated API calls")
    parser.add_argument("--skip-trait-evidence", action="store_true", help="Skip per-trait search-by-trait-and-gene calls")
    parser.add_argument("--max-evidence-rows", type=int, default=0, help="Retained for backward compatibility; full evidence is written to CSV, not the DOCX body")
    parser.add_argument("--max-trait-rows", type=int, default=REPORT_TRAIT_ROWS, help=f"Max trait rows in DOCX summary table; default {REPORT_TRAIT_ROWS}")
    args = parser.parse_args()

    if args.no_api and not args.input_json:
        parser.error("Provide --input-json when using --no-api")

    bundle: Dict[str, Any] = {"gene": args.gene}
    if args.input_json:
        bundle.update(read_json(args.input_json))
    if not args.no_api:
        fetched = fetch_from_api(
            args.api_base,
            args.gene,
            args.endpoint_map,
            args.page_limit,
            args.max_pages,
            fetch_trait_evidence=not args.skip_trait_evidence,
        )
        bundle.update(fetched)

    out_path = args.out or Path(f"{safe_filename(args.gene)}_RiceMind_gene_mechanism_report.docx")
    fig_dir = args.fig_dir or out_path.with_suffix("").parent / f"{out_path.stem}_figures"

    traits = normalize_traits(bundle)
    evidence = normalize_evidence(bundle, traits)
    varieties = normalize_varieties(bundle)
    bundle["api_calls"] = compact_api_calls(bundle.get("api_calls") or [])

    write_csv(
        out_path.with_name(f"{out_path.stem}_normalized_traits.csv"),
        traits,
        ["gene", "rap_id", "trait", "ontology_type", "ontology_id", "evidence_code", "source_db", "confidence", "support", "year"],
    )
    write_csv(
        out_path.with_name(f"{out_path.stem}_normalized_evidence.csv"),
        evidence,
        ["trait", "ontology_type", "ontology_id", "evidence_code", "source_db", "confidence", "pmid", "sentence_id", "year", "title", "journal", "doi", "sentence"],
    )
    write_csv(out_path.with_name(f"{out_path.stem}_normalized_varieties.csv"), varieties, ["variety"])
    write_json(out_path.with_name(f"{out_path.stem}_payload.json"), bundle)

    fig_paths = {
        "confidence": plot_counter(Counter(row["confidence"] for row in traits if row["confidence"]), "Confidence-tier distribution", fig_dir / "confidence_distribution.png"),
        "ontology": plot_counter(Counter(row["ontology_type"] for row in traits if row["ontology_type"]), "Ontology distribution", fig_dir / "ontology_distribution.png"),
        "top_traits": plot_top_traits(traits, fig_dir / "top_traits_by_support.png"),
        "evidence_code": plot_counter(split_counter(traits + evidence, "evidence_code"), "Evidence-code distribution", fig_dir / "evidence_code_distribution.png"),
        "source": plot_counter(split_counter(traits + evidence, "source_db"), "Source database distribution", fig_dir / "source_distribution.png"),
        "years": plot_years((row["year"] for row in evidence), fig_dir / "publication_year_trend.png"),
    }

    try:
        build_docx(args.gene, bundle, traits, evidence, varieties, out_path, fig_paths, args.language, args.max_evidence_rows, args.max_trait_rows)
    except RuntimeError as exc:
        print(str(exc), file=sys.stderr)
        return 2

    print(f"Wrote {out_path}")
    print(f"Wrote {out_path.with_name(f'{out_path.stem}_payload.json')}")
    if traits:
        print(f"Wrote {out_path.with_name(f'{out_path.stem}_normalized_traits.csv')}")
    if evidence:
        print(f"Wrote {out_path.with_name(f'{out_path.stem}_normalized_evidence.csv')}")
    if varieties:
        print(f"Wrote {out_path.with_name(f'{out_path.stem}_normalized_varieties.csv')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
