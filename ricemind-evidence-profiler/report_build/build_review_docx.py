#!/usr/bin/env python
"""Render the planthopper review Markdown into a formatted DOCX."""

from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path
from typing import Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


CITATION_RE = re.compile(r"\[((?:\d{8}(?:,\s*)?)+)\]")
IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
TABLE_RE = re.compile(r"^\|.*\|$")
SEPARATOR_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+$")


def read_articles(path: Path) -> Dict[str, Dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return {row["PMID"]: row for row in csv.DictReader(handle) if row.get("PMID")}


def cited_pmids(text: str) -> List[str]:
    seen = set()
    out = []
    for match in CITATION_RE.finditer(text):
        for pmid in re.findall(r"\d{8}", match.group(1)):
            if pmid not in seen:
                seen.add(pmid)
                out.append(pmid)
    return out


def reference_lines(pmids: Sequence[str], articles: Dict[str, Dict[str, str]]) -> List[str]:
    refs = []
    for idx, pmid in enumerate(pmids, 1):
        row = articles.get(pmid, {})
        if row:
            refs.append(
                f"{idx}. {row.get('title', '').rstrip('.')} "
                f"{row.get('journal', '')}. {row.get('year', '')}. "
                f"PMID: {pmid}. https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
            )
        else:
            refs.append(f"{idx}. PMID: {pmid}. https://pubmed.ncbi.nlm.nih.gov/{pmid}/")
    return refs


def finalize_markdown(text: str, refs: Sequence[str]) -> str:
    return text.replace("{{REFERENCES}}", "\n\n".join(refs))


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "SimSun")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str, size: float = 9.5):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F5F8F")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "SimSun")
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    rpr.extend([fonts, color, underline, size_el])
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, size: float = 10.5) -> None:
    pos = 0
    token_re = re.compile(r"(\*\*.*?\*\*|\[((?:\d{8}(?:,\s*)?)+)\])")
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size, bold=True)
        else:
            pmids = re.findall(r"\d{8}", token)
            run = paragraph.add_run("[")
            set_run_font(run, size)
            for idx, pmid in enumerate(pmids):
                if idx:
                    comma = paragraph.add_run(", ")
                    set_run_font(comma, size)
                add_hyperlink(paragraph, pmid, f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/", size)
            run = paragraph.add_run("]")
            set_run_font(run, size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, source_row in enumerate(rows):
        target_row = table.rows[r_idx]
        prevent_row_split(target_row)
        for c_idx, value in enumerate(source_row):
            cell = target_row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "DCE6F1")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value.strip(), 8.5 if r_idx else 9)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc: Document, path: Path, caption: str) -> None:
    if not path.is_file():
        paragraph = doc.add_paragraph()
        add_inline(paragraph, f"[缺失图件：{path.name}]", 9.5)
        return
    with Image.open(path) as image:
        width_px, height_px = image.size
    max_width = 6.55
    max_height = 5.35
    ratio = min(max_width / width_px, max_height / height_px) * width_px
    width_in = max(2.0, min(max_width, ratio))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    paragraph.paragraph_format.keep_with_next = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    run = cap.add_run(caption)
    set_run_font(run, 9, bold=False, color="444444")


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "打开文档后右键更新目录。"
    fld_sep.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, fld_end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, 9)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    run = paragraph.add_run(" 页")
    set_run_font(run, 9)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in (
        ("Title", 18, "1F3A56"),
        ("Heading 1", 14, "1F3A56"),
        ("Heading 2", 12, "2F5F70"),
        ("Heading 3", 11, "365F4B"),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    if "CaptionCustom" not in styles:
        style = styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(68, 68, 68)


def parse_table(lines: Sequence[str], start: int) -> Tuple[List[List[str]], int]:
    rows = []
    index = start
    while index < len(lines) and TABLE_RE.match(lines[index].strip()):
        line = lines[index].strip()
        if not SEPARATOR_RE.match(line):
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        index += 1
    return rows, index


def render(markdown: str, markdown_path: Path, out_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("RiceMind 水稻飞虱抗性基因分子机理综述")
    set_run_font(run, 8.5, color="666666")
    add_page_number(section.footer.paragraphs[0])

    lines = markdown.splitlines()
    index = 0
    title_seen = False
    section_count = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, line[2:].strip(), 18)
            title_seen = True
            index += 1
            continue
        if title_seen and line.startswith("**") and line.endswith("**"):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, line, 10)
            index += 1
            continue
        if line.startswith("## "):
            section_count += 1
            if section_count == 1:
                toc_title = doc.add_paragraph()
                toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = toc_title.add_run("目录")
                set_run_font(run, 14, bold=True, color="1F3A56")
                add_toc(doc.add_paragraph())
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline(paragraph, line[3:].strip(), 14)
            index += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, line[4:].strip(), 12)
            index += 1
            continue
        image_match = IMAGE_RE.fullmatch(line)
        if image_match:
            caption, rel_path = image_match.groups()
            add_image(doc, markdown_path.parent / rel_path, caption)
            index += 1
            continue
        if TABLE_RE.match(line):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", line), 10.5)
            index += 1
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, line[2:], 10.5)
            index += 1
            continue

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(21)
        paragraph.paragraph_format.widow_control = True
        add_inline(paragraph, line, 10.5)
        index += 1

    core = doc.core_properties
    core.title = "基于 RiceMind 证据的水稻飞虱抗性基因分子机理综述"
    core.subject = "RiceMind sentence-evidence-driven review of rice planthopper resistance genes"
    core.author = "RiceMind Evidence Profiler"
    core.keywords = "RiceMind, rice, brown planthopper, BPH, resistance genes, breeding"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--article-index", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--finalize-markdown", action="store_true")
    args = parser.parse_args()

    text = args.markdown.read_text(encoding="utf-8")
    articles = read_articles(args.article_index)
    pmids = cited_pmids(text)
    missing = [pmid for pmid in pmids if pmid not in articles]
    if missing:
        raise RuntimeError(f"Cited PMIDs missing from RiceMind article index: {', '.join(missing)}")
    refs = reference_lines(pmids, articles)
    finalized = finalize_markdown(text, refs)
    if args.finalize_markdown:
        args.markdown.write_text(finalized, encoding="utf-8")
    render(finalized, args.markdown, args.out)
    print(f"citations={len(pmids)}")
    print(f"docx={args.out}")


if __name__ == "__main__":
    main()
