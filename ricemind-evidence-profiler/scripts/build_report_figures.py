#!/usr/bin/env python
"""Render and place data-supported figures in RiceMind reports.

Automatic mode creates only task-neutral summaries. For a personalized report,
provide a JSON figure plan that selects the data fields, chart type, title,
caption, and report section for each figure.
"""

from __future__ import annotations

import argparse
import csv
import html
import json
import math
import re
import textwrap
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple


LEGACY_BLOCK_START = "<!-- RICEMIND_FIGURES_START -->"
LEGACY_BLOCK_END = "<!-- RICEMIND_FIGURES_END -->"
BLOCK_START = "<!-- RICEMIND_FIGURE_BLOCK:{block_id}:START -->"
BLOCK_END = "<!-- RICEMIND_FIGURE_BLOCK:{block_id}:END -->"

ENTITY_FIELDS = ("target", "candidate_gene", "gene", "trait", "variety", "candidate", "symbol")
YEAR_FIELDS = ("year", "publication_year", "earliest_year")
JOURNAL_FIELDS = ("journal",)
CONFIDENCE_FIELDS = ("confidence", "confidence_tier", "tier")
ONTOLOGY_FIELDS = ("ontology_type", "onto_type")
EVIDENCE_CODE_FIELDS = ("evidence_code", "evidence_codes", "Evidence_Code")
SOURCE_FIELDS = ("source_db", "sources", "Source_DB")
TRAIT_FIELDS = ("trait", "trait_name", "trait_description", "Trait_Description")

DEFAULT_COLORS = ("#4f7cac", "#59a14f", "#f28e2b", "#b07aa1", "#76b7b2", "#e15759")


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def first_field(rows: Sequence[Dict[str, str]], candidates: Sequence[str]) -> str:
    if not rows:
        return ""
    fields = set(rows[0])
    return next((field for field in candidates if field in fields), "")


def number(value: object) -> Optional[float]:
    try:
        parsed = float(str(value).strip())
    except (TypeError, ValueError):
        return None
    return parsed if math.isfinite(parsed) else None


def split_values(value: str) -> Iterable[str]:
    for item in re.split(r"[;,|]", value or ""):
        cleaned = item.strip().strip("[]'\"")
        if cleaned:
            yield cleaned


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("_") or "figure"


def humanize(field: str) -> str:
    text = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", field or "")
    text = text.replace("_", " ").replace("-", " ")
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\btier\s*([123])\b", r"Tier \1", text, flags=re.I)
    text = re.sub(r"\bpmids\b", "PMIDs", text, flags=re.I)
    text = re.sub(r"\bpmid\b", "PMID", text, flags=re.I)
    return text[:1].upper() + text[1:]


def detect_language(text: str) -> str:
    return "zh" if len(re.findall(r"[\u4e00-\u9fff]", text or "")) >= 8 else "en"


def localized(language: str, zh: str, en: str) -> str:
    return zh if language.lower().startswith("zh") else en


def wrap_label(value: object, width: int = 34) -> str:
    text = str(value or "").strip()
    if len(text) <= width:
        return text
    return "\n".join(textwrap.wrap(text, width=width, break_long_words=False, break_on_hyphens=False))


def configure_matplotlib() -> None:
    import matplotlib

    matplotlib.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": ["Microsoft YaHei", "SimHei", "Arial", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "savefig.facecolor": "white",
        }
    )


def aggregate_by(
    rows: Sequence[Dict[str, str]],
    category: str,
    value: Optional[str] = None,
    aggregate: str = "count",
) -> List[Tuple[str, float]]:
    grouped: Dict[str, List[float]] = defaultdict(list)
    for row in rows:
        label = str(row.get(category, "")).strip()
        if not label:
            continue
        if value:
            parsed = number(row.get(value))
            if parsed is None:
                continue
            grouped[label].append(parsed)
        else:
            grouped[label].append(1.0)

    def reduce(values: List[float]) -> float:
        if aggregate == "sum":
            return sum(values)
        if aggregate == "max":
            return max(values)
        if aggregate == "mean":
            return sum(values) / len(values)
        if aggregate == "min":
            return min(values)
        if aggregate == "unique":
            return float(len(set(values)))
        return float(len(values))

    return [(label, reduce(values)) for label, values in grouped.items()]


def save_figure(fig: Any, out_path: Path) -> Path:
    import matplotlib.pyplot as plt

    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=200, bbox_inches="tight", pad_inches=0.12)
    plt.close(fig)
    return out_path


def render_ranked_bar(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    category = spec.get("category") or first_field(rows, ENTITY_FIELDS)
    value = spec.get("value")
    if not category or not value:
        return None
    pairs = aggregate_by(rows, category, value, spec.get("aggregate", "max"))
    pairs.sort(key=lambda item: item[1], reverse=spec.get("sort", "desc") != "asc")
    pairs = pairs[: int(spec.get("top_n", 12))]
    pairs = [(label, value) for label, value in pairs if value != 0]
    if not pairs:
        return None

    import matplotlib.pyplot as plt

    configure_matplotlib()
    labels = [wrap_label(item[0]) for item in pairs][::-1]
    values = [item[1] for item in pairs][::-1]
    fig_height = min(6.6, max(3.2, 0.38 * len(labels) + 1.25))
    fig, ax = plt.subplots(figsize=(7.1, fig_height), constrained_layout=True)
    ax.barh(labels, values, color=spec.get("color", DEFAULT_COLORS[0]))
    ax.set_xlabel(spec.get("xlabel") or humanize(value))
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    ax.grid(axis="x", alpha=0.2)
    return save_figure(fig, out_path)


def render_category_bar(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    category = spec.get("category")
    if not category:
        return None
    expanded: List[Dict[str, str]] = []
    if spec.get("split_values", True):
        for row in rows:
            expanded.extend({category: value} for value in split_values(row.get(category, "")))
    else:
        expanded = list(rows)
    pairs = aggregate_by(expanded, category)
    pairs.sort(key=lambda item: item[1], reverse=True)
    pairs = pairs[: int(spec.get("top_n", 12))]
    if not pairs:
        return None

    import matplotlib.pyplot as plt

    configure_matplotlib()
    labels = [wrap_label(item[0]) for item in pairs][::-1]
    values = [item[1] for item in pairs][::-1]
    fig_height = min(6.6, max(3.2, 0.38 * len(labels) + 1.25))
    fig, ax = plt.subplots(figsize=(7.1, fig_height), constrained_layout=True)
    ax.barh(labels, values, color=spec.get("color", DEFAULT_COLORS[2]))
    ax.set_xlabel(spec.get("xlabel") or localized(spec.get("language", "en"), "记录数", "Record count"))
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    ax.grid(axis="x", alpha=0.2)
    return save_figure(fig, out_path)


def render_grouped_bar(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    category = spec.get("category") or first_field(rows, ENTITY_FIELDS)
    values = [field for field in spec.get("values", []) if field]
    if not category or len(values) < 2:
        return None
    aggregate = spec.get("aggregate", "max")
    data_by_field = {field: dict(aggregate_by(rows, category, field, aggregate)) for field in values}
    labels = sorted(
        set().union(*(mapping.keys() for mapping in data_by_field.values())),
        key=lambda label: sum(data_by_field[field].get(label, 0) for field in values),
        reverse=True,
    )[: int(spec.get("top_n", 12))]
    if not labels:
        return None

    import matplotlib.pyplot as plt
    import numpy as np

    configure_matplotlib()
    y = np.arange(len(labels))
    bar_height = min(0.8 / len(values), 0.32)
    fig_height = min(6.6, max(3.4, 0.43 * len(labels) + 1.4))
    fig, ax = plt.subplots(figsize=(7.1, fig_height), constrained_layout=True)
    display_labels = spec.get("series_labels", {})
    for index, field in enumerate(values):
        offset = (index - (len(values) - 1) / 2) * bar_height
        ax.barh(
            y + offset,
            [data_by_field[field].get(label, 0) for label in labels],
            bar_height,
            label=display_labels.get(field, humanize(field)),
            color=spec.get("colors", DEFAULT_COLORS)[index % len(spec.get("colors", DEFAULT_COLORS))],
        )
    ax.set_yticks(y)
    ax.set_yticklabels([wrap_label(label) for label in labels])
    ax.invert_yaxis()
    ax.set_xlabel(spec.get("xlabel") or localized(spec.get("language", "en"), "数值", "Value"))
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    ax.legend(loc="best", frameon=False)
    ax.grid(axis="x", alpha=0.2)
    return save_figure(fig, out_path)


def render_timeline(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    year_field = spec.get("year") or first_field(rows, YEAR_FIELDS)
    if not year_field:
        return None
    counts = Counter()
    for row in rows:
        value = str(row.get(year_field, "")).strip()
        if re.fullmatch(r"(?:19|20)\d{2}", value):
            counts[int(value)] += 1
    if not counts:
        return None

    import matplotlib.pyplot as plt

    configure_matplotlib()
    years = sorted(counts)
    fig, ax = plt.subplots(figsize=(7.1, 3.6), constrained_layout=True)
    ax.plot(years, [counts[year] for year in years], marker="o", linewidth=1.8, color=spec.get("color", DEFAULT_COLORS[1]))
    ax.fill_between(years, [counts[year] for year in years], alpha=0.15, color=spec.get("color", DEFAULT_COLORS[1]))
    ax.set_xlabel(spec.get("xlabel") or localized(spec.get("language", "en"), "发表年份", "Publication year"))
    ax.set_ylabel(spec.get("ylabel") or localized(spec.get("language", "en"), "证据记录数", "Evidence records"))
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    ax.grid(alpha=0.2)
    return save_figure(fig, out_path)


def render_histogram(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    value = spec.get("value")
    values = [parsed for row in rows if (parsed := number(row.get(value))) is not None] if value else []
    if not values:
        return None

    import matplotlib.pyplot as plt

    configure_matplotlib()
    fig, ax = plt.subplots(figsize=(7.1, 3.8), constrained_layout=True)
    ax.hist(values, bins=int(spec.get("bins", min(20, max(6, round(math.sqrt(len(values))))))), color=spec.get("color", DEFAULT_COLORS[0]), alpha=0.9)
    ax.set_xlabel(spec.get("xlabel") or humanize(value))
    ax.set_ylabel(spec.get("ylabel") or localized(spec.get("language", "en"), "记录数", "Record count"))
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    ax.grid(axis="y", alpha=0.2)
    return save_figure(fig, out_path)


def render_scatter(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    x_field, y_field = spec.get("x"), spec.get("y")
    label_field = spec.get("label") or first_field(rows, ENTITY_FIELDS)
    points = []
    for row in rows:
        x_value, y_value = number(row.get(x_field)), number(row.get(y_field))
        if x_value is not None and y_value is not None:
            points.append((x_value, y_value, row.get(label_field, "") if label_field else ""))
    if not points:
        return None

    import matplotlib.pyplot as plt

    configure_matplotlib()
    fig, ax = plt.subplots(figsize=(6.8, 5.0), constrained_layout=True)
    ax.scatter([point[0] for point in points], [point[1] for point in points], color=spec.get("color", DEFAULT_COLORS[0]), alpha=0.78)
    for x_value, y_value, label in sorted(points, key=lambda item: item[0] + item[1], reverse=True)[: int(spec.get("label_top_n", 10))]:
        if label:
            ax.annotate(str(label), (x_value, y_value), xytext=(4, 4), textcoords="offset points", fontsize=8)
    ax.set_xlabel(spec.get("xlabel") or humanize(x_field))
    ax.set_ylabel(spec.get("ylabel") or humanize(y_field))
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    ax.grid(alpha=0.2)
    return save_figure(fig, out_path)


def render_heatmap(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    row_field, column_field, value_field = spec.get("row"), spec.get("column"), spec.get("value")
    if not row_field or not column_field:
        return None
    row_labels = Counter(row.get(row_field, "") for row in rows if row.get(row_field, "")).most_common(int(spec.get("top_rows", 12)))
    column_labels = Counter(row.get(column_field, "") for row in rows if row.get(column_field, "")).most_common(int(spec.get("top_columns", 10)))
    row_names = [item[0] for item in row_labels]
    column_names = [item[0] for item in column_labels]
    if not row_names or not column_names:
        return None
    row_index, column_index = {name: i for i, name in enumerate(row_names)}, {name: i for i, name in enumerate(column_names)}

    import matplotlib.pyplot as plt
    import numpy as np

    matrix = np.zeros((len(row_names), len(column_names)))
    for row in rows:
        if row.get(row_field) not in row_index or row.get(column_field) not in column_index:
            continue
        amount = number(row.get(value_field)) if value_field else 1.0
        matrix[row_index[row[row_field]], column_index[row[column_field]]] += amount if amount is not None else 0
    if not matrix.any():
        return None
    configure_matplotlib()
    fig, ax = plt.subplots(figsize=(7.1, min(6.4, max(3.8, len(row_names) * 0.38 + 1.5))), constrained_layout=True)
    image = ax.imshow(matrix, aspect="auto", cmap=spec.get("cmap", "YlGnBu"))
    ax.set_xticks(range(len(column_names)))
    ax.set_xticklabels([wrap_label(name, 18) for name in column_names], rotation=40, ha="right")
    ax.set_yticks(range(len(row_names)))
    ax.set_yticklabels([wrap_label(name, 28) for name in row_names])
    fig.colorbar(image, ax=ax, shrink=0.8)
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    return save_figure(fig, out_path)


def render_network(rows: Sequence[Dict[str, str]], spec: Dict[str, Any], out_path: Path) -> Optional[Path]:
    source, target = spec.get("source_field", "source"), spec.get("target_field", "target")
    weight = spec.get("weight")
    valid_rows = [row for row in rows if row.get(source) and row.get(target)]
    if not valid_rows:
        return None
    try:
        import matplotlib.pyplot as plt
        import networkx as nx
    except ImportError:
        return None
    graph = nx.Graph()
    for row in valid_rows:
        edge_weight = number(row.get(weight)) if weight else 1.0
        graph.add_edge(row[source], row[target], weight=edge_weight or 1.0)
    max_nodes = int(spec.get("max_nodes", 35))
    if graph.number_of_nodes() > max_nodes:
        keep = {node for node, _ in sorted(graph.degree, key=lambda item: item[1], reverse=True)[:max_nodes]}
        graph = graph.subgraph(keep).copy()
    if not graph.number_of_edges():
        return None
    configure_matplotlib()
    fig, ax = plt.subplots(figsize=(7.1, 6.0), constrained_layout=True)
    positions = nx.spring_layout(graph, seed=17, weight="weight")
    degrees = dict(graph.degree)
    nx.draw_networkx_edges(graph, positions, ax=ax, alpha=0.32, width=0.8)
    nx.draw_networkx_nodes(
        graph,
        positions,
        ax=ax,
        node_size=[80 + degrees[node] * 35 for node in graph],
        node_color=spec.get("color", DEFAULT_COLORS[0]),
        alpha=0.85,
    )
    nx.draw_networkx_labels(graph, positions, ax=ax, font_size=7)
    ax.axis("off")
    if spec.get("plot_title"):
        ax.set_title("\n".join(textwrap.wrap(str(spec["plot_title"]), 68)))
    return save_figure(fig, out_path)


RENDERERS = {
    "ranked_bar": render_ranked_bar,
    "category_bar": render_category_bar,
    "grouped_bar": render_grouped_bar,
    "timeline": render_timeline,
    "histogram": render_histogram,
    "scatter": render_scatter,
    "heatmap": render_heatmap,
    "network": render_network,
}


def resolve_source(data_dir: Path, source: str) -> Optional[Path]:
    if not source:
        return None
    candidate = Path(source)
    if not candidate.is_absolute():
        candidate = data_dir / candidate
    if candidate.is_file():
        return candidate
    matches = [path for path in data_dir.glob(source) if path.is_file()]
    return max(matches, key=lambda path: path.stat().st_size) if matches else None


def render_plan(data_dir: Path, fig_dir: Path, plan: Dict[str, Any]) -> List[Dict[str, Any]]:
    figures = []
    language = plan.get("language", "en")
    for index, raw_spec in enumerate(plan.get("figures", []), 1):
        spec = dict(raw_spec)
        spec.setdefault("language", language)
        figure_type = spec.get("type")
        source_path = resolve_source(data_dir, spec.get("source", ""))
        if figure_type not in RENDERERS or not source_path:
            continue
        rows = read_csv(source_path)
        if not rows:
            continue
        figure_id = safe_name(spec.get("id") or f"figure_{index}")
        out_path = fig_dir / f"{figure_id}.png"
        rendered = RENDERERS[figure_type](rows, spec, out_path)
        if not rendered:
            if out_path.is_file():
                out_path.unlink()
            continue
        figures.append(
            {
                "id": figure_id,
                "path": str(rendered),
                "type": figure_type,
                "title": spec.get("title") or spec.get("plot_title") or humanize(figure_id),
                "caption": spec.get("caption") or spec.get("title") or humanize(figure_id),
                "section": spec.get("section", ""),
                "section_keywords": spec.get("section_keywords", []),
                "subsection": spec.get("subsection") or spec.get("title") or "",
                "placement": spec.get("placement", "end_of_section"),
                "width_percent": max(55, min(94, int(spec.get("width_percent", 88)))),
                "max_height_inches": max(3.2, min(7.2, float(spec.get("max_height_inches", 6.4)))),
            }
        )
    return figures


def find_csv(data_dir: Path, patterns: Sequence[str]) -> Optional[Path]:
    matches = [path for pattern in patterns for path in data_dir.glob(pattern) if path.is_file()]
    return max(matches, key=lambda path: path.stat().st_size) if matches else None


def numeric_columns(rows: Sequence[Dict[str, str]]) -> List[str]:
    if not rows:
        return []
    columns = []
    for field in rows[0]:
        values = [row.get(field, "") for row in rows[:100] if str(row.get(field, "")).strip()]
        if values and sum(number(value) is not None for value in values) / len(values) >= 0.8:
            columns.append(field)
    return columns


def auto_plan(data_dir: Path, language: str = "en") -> Dict[str, Any]:
    """Create a task-neutral fallback plan without objective-specific biology."""
    evidence_path = find_csv(data_dir, ("*sentence_evidence.csv", "*normalized_evidence.csv", "*sentences.csv"))
    ranking_path = find_csv(data_dir, ("*candidate_ranking.csv", "*candidate_genes.csv", "*recommendations.csv"))
    trait_path = find_csv(data_dir, ("*traits.csv", "*tier1_traits.csv", "*normalized_traits.csv"))
    edge_path = find_csv(data_dir, ("*edges.csv", "*network_edges.csv"))
    specs: List[Dict[str, Any]] = []

    evidence_rows = read_csv(evidence_path) if evidence_path else []
    if evidence_path and first_field(evidence_rows, YEAR_FIELDS):
        specs.append(
            {
                "id": "publication_year_distribution",
                "type": "timeline",
                "source": evidence_path.name,
                "year": first_field(evidence_rows, YEAR_FIELDS),
                "title": localized(language, "证据的时间分布", "Temporal distribution of evidence"),
                "caption": localized(language, "RiceMind 句子证据按发表年份的分布。", "Publication-year distribution of RiceMind Sentence Evidence."),
                "section_keywords": ["证据分布", "检索", "文献计量", "evidence distribution", "retrieval", "bibliometric"],
                "subsection": localized(language, "证据的时间分布", "Temporal evidence distribution"),
                "max_height_inches": 4.2,
            }
        )
    for fields, figure_id, zh_title, en_title in (
        (JOURNAL_FIELDS, "journal_distribution", "证据来源期刊分布", "Journal distribution"),
        (EVIDENCE_CODE_FIELDS, "evidence_code_distribution", "证据代码构成", "Evidence-code composition"),
        (SOURCE_FIELDS, "source_distribution", "数据来源构成", "Source-database composition"),
        (CONFIDENCE_FIELDS, "confidence_distribution", "置信层级构成", "Confidence-tier composition"),
    ):
        field = first_field(evidence_rows, fields)
        if evidence_path and field:
            specs.append(
                {
                    "id": figure_id,
                    "type": "category_bar",
                    "source": evidence_path.name,
                    "category": field,
                    "title": localized(language, zh_title, en_title),
                    "caption": localized(language, f"当前 RiceMind 证据集中的{zh_title}。", f"{en_title} in the current RiceMind evidence set."),
                    "section_keywords": ["证据分布", "检索", "文献计量", "evidence distribution", "retrieval", "bibliometric"],
                    "subsection": localized(language, "证据构成", "Evidence composition"),
                }
            )

    ranking_rows = read_csv(ranking_path) if ranking_path else []
    entity_field = first_field(ranking_rows, ENTITY_FIELDS)
    preferred = [
        field
        for field in ("score", "unique_pmids", "pmid_count", "sentence_count", "article_count", "support")
        if field in numeric_columns(ranking_rows)
    ]
    if not preferred:
        preferred = [
            field
            for field in numeric_columns(ranking_rows)
            if field.lower() not in {"rank", "year", "earliest_year"} and not field.lower().endswith("_flag")
        ]
    for value_field in preferred[:2]:
        specs.append(
            {
                "id": f"candidate_{safe_name(value_field)}",
                "type": "ranked_bar",
                "source": ranking_path.name if ranking_path else "",
                "category": entity_field,
                "value": value_field,
                "title": localized(language, f"候选对象的{humanize(value_field)}比较", f"Candidate comparison by {humanize(value_field).lower()}"),
                "caption": localized(language, f"按 {humanize(value_field)} 展示证据支持度最高的候选对象。", f"Top candidates ranked by {humanize(value_field).lower()}."),
                "section_keywords": ["候选", "优先靶点", "排序", "candidate", "prioritized", "ranking"],
                "subsection": localized(language, "候选对象的证据比较", "Candidate evidence comparison"),
                "top_n": 12,
            }
        )

    trait_rows = read_csv(trait_path) if trait_path else []
    trait_field = first_field(trait_rows, TRAIT_FIELDS)
    if trait_path and trait_field:
        specs.append(
            {
                "id": "trait_distribution",
                "type": "category_bar",
                "source": trait_path.name,
                "category": trait_field,
                "title": localized(language, "关联性状构成", "Associated-trait composition"),
                "caption": localized(language, "当前候选或目标对象所关联性状的分布。", "Distribution of traits associated with the current candidates or targets."),
                "section_keywords": ["性状", "关联", "gta", "trait", "association"],
                "subsection": localized(language, "关联性状构成", "Associated-trait composition"),
                "top_n": 12,
            }
        )

    if edge_path:
        specs.append(
            {
                "id": "evidence_network",
                "type": "network",
                "source": edge_path.name,
                "title": localized(language, "证据关系网络", "Evidence relationship network"),
                "caption": localized(language, "依据当前标准化边表构建的证据关系网络。", "Evidence relationship network built from the normalized edge table."),
                "section_keywords": ["网络", "关系", "network", "relationship"],
                "subsection": localized(language, "证据关系网络", "Evidence relationship network"),
            }
        )
    return {
        "language": language,
        "fallback_section_title": localized(language, "数据支持的可视化摘要", "Data-supported visual summary"),
        "figures": specs,
    }


def build_trait_figures(
    sentences: Sequence[Dict[str, str]],
    candidates: Sequence[Dict[str, str]],
    fig_dir: Path,
) -> List[Dict[str, Any]]:
    """Generate task-neutral figures for the lightweight trait builder."""
    temp_dir = fig_dir.parent / ".figure_input"
    temp_dir.mkdir(parents=True, exist_ok=True)
    sentence_path, candidate_path = temp_dir / "sentence_evidence.csv", temp_dir / "candidate_genes.csv"
    write_rows(sentence_path, sentences)
    write_rows(candidate_path, candidates)
    plan = auto_plan(temp_dir, "en")
    figures = render_plan(temp_dir, fig_dir, plan)
    for path in (sentence_path, candidate_path):
        if path.is_file():
            path.unlink()
    try:
        temp_dir.rmdir()
    except OSError:
        pass
    return figures


def write_rows(path: Path, rows: Sequence[Dict[str, str]]) -> None:
    if not rows:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def build_figures_from_data_dir(
    data_dir: Path,
    fig_dir: Optional[Path] = None,
    plan: Optional[Dict[str, Any]] = None,
    language: str = "en",
) -> List[Dict[str, Any]]:
    fig_dir = fig_dir or data_dir / "figures"
    return render_plan(data_dir, fig_dir, plan or auto_plan(data_dir, language))


def strip_generated_blocks(text: str) -> str:
    if LEGACY_BLOCK_START in text and LEGACY_BLOCK_END in text:
        text = re.sub(
            re.escape(LEGACY_BLOCK_START) + r".*?" + re.escape(LEGACY_BLOCK_END),
            "",
            text,
            flags=re.S,
        )
    text = re.sub(
        r"<!-- RICEMIND_FIGURE_BLOCK:[^:>]+:START -->.*?<!-- RICEMIND_FIGURE_BLOCK:[^:>]+:END -->",
        "",
        text,
        flags=re.S,
    )
    return re.sub(r"\n{3,}", "\n\n", text).rstrip() + "\n"


def markdown_headings(lines: Sequence[str]) -> List[Tuple[int, int, str]]:
    headings = []
    for index, line in enumerate(lines):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            headings.append((index, len(match.group(1)), re.sub(r"[`*_]", "", match.group(2)).strip()))
    return headings


def normalized_heading(value: str) -> str:
    return re.sub(r"[\s:：、，,。.·\-—_`*#]+", "", value or "").lower()


def locate_section(lines: Sequence[str], figure: Dict[str, Any]) -> Optional[Tuple[int, int, int]]:
    headings = markdown_headings(lines)
    section = normalized_heading(figure.get("section", ""))
    keywords = [normalized_heading(item) for item in figure.get("section_keywords", []) if item]
    selected = None
    for position, level, title in headings:
        normalized = normalized_heading(title)
        if section and (section == normalized or section in normalized or normalized in section):
            selected = (position, level)
            break
        if keywords and any(keyword in normalized for keyword in keywords):
            selected = (position, level)
            break
    if not selected:
        return None
    start, level = selected
    end = len(lines)
    for position, next_level, _ in headings:
        if position > start and next_level <= level:
            end = position
            break
    return start, level, end


def figure_html(path: Path, report_path: Path, caption: str, number_index: int, width_percent: int, max_height: float) -> str:
    try:
        relative = path.relative_to(report_path.parent).as_posix()
    except ValueError:
        relative = path.as_posix()
    label = "图" if detect_language(report_path.read_text(encoding="utf-8")) == "zh" else "Figure"
    escaped_path = html.escape(relative, quote=True)
    escaped_caption = html.escape(caption, quote=True)
    return "\n".join(
        [
            '<figure class="ricemind-figure" style="break-inside: avoid; page-break-inside: avoid; margin: 1.1em auto 1.4em; text-align: center;">',
            f'  <img src="{escaped_path}" alt="{escaped_caption}" style="display: block; width: {width_percent}%; max-width: 7.0in; max-height: {max_height:.1f}in; height: auto; object-fit: contain; margin: 0 auto;">',
            f'  <figcaption style="font-size: 0.9em; line-height: 1.45; margin: 0.55em auto 0; max-width: 90%;"><strong>{label} {number_index}.</strong> {escaped_caption}</figcaption>',
            "</figure>",
        ]
    )


def update_markdown(
    path: Path,
    figures: Sequence[Dict[str, Any]],
    fallback_section_title: Optional[str] = None,
) -> None:
    if not figures:
        return
    original = path.read_text(encoding="utf-8") if path.is_file() else ""
    text = strip_generated_blocks(original)
    lines = text.rstrip().splitlines()
    language = detect_language(text)
    fallback_title = fallback_section_title or localized(language, "数据支持的可视化摘要", "Data-supported visual summary")

    prepared_figures = []
    for index, raw_figure in enumerate(figures, 1):
        figure = dict(raw_figure)
        figure_path = Path(figure["path"])
        figure.setdefault("id", safe_name(figure_path.stem or f"figure_{index}"))
        figure.setdefault("title", humanize(figure["id"]))
        figure.setdefault("caption", figure["title"])
        figure.setdefault("subsection", figure["title"])
        figure.setdefault("section", "")
        figure.setdefault("section_keywords", [])
        figure.setdefault("placement", "end_of_section")
        figure.setdefault("width_percent", 88)
        figure.setdefault("max_height_inches", 6.4)
        prepared_figures.append(figure)

    assignments: List[Tuple[int, int, int, Dict[str, Any]]] = []
    unmatched: List[Dict[str, Any]] = []
    for figure in prepared_figures:
        location = locate_section(lines, figure)
        if location:
            assignments.append((*location, figure))
        else:
            unmatched.append(figure)

    # Figure numbers follow report order, not generation order.
    boundary_keywords = ["证据边界", "局限", "limitations", "evidence boundaries", "interpretation limits"]
    boundary_index = next(
        (
            index
            for index, _, title in markdown_headings(lines)
            if any(normalized_heading(keyword) in normalized_heading(title) for keyword in boundary_keywords)
        ),
        len(lines),
    )
    for figure in unmatched:
        assignments.append((boundary_index, 1, boundary_index, figure))
    assignments.sort(key=lambda item: (item[2], item[0]))

    numbered = {id(item[3]): index for index, item in enumerate(assignments, 1)}
    grouped: Dict[Tuple[int, int, str], List[Dict[str, Any]]] = defaultdict(list)
    for start, level, end, figure in assignments:
        placement = figure.get("placement", "end_of_section")
        insert_at = start + 1 if placement == "after_heading" else end
        grouped[(insert_at, level, "fallback" if figure in unmatched else "section")].append(figure)

    insertions = []
    for (insert_at, level, group_type), group_figures in grouped.items():
        block_id = safe_name("_".join(figure["id"] for figure in group_figures))
        block = [BLOCK_START.format(block_id=block_id), ""]
        if group_type == "fallback":
            block.extend([f"## {fallback_title}", ""])
            subsection_level = 3
        else:
            subsection_level = min(6, level + 1)
        current_subsection = None
        for figure in group_figures:
            subsection = figure.get("subsection") or figure.get("title")
            heading_html = ""
            if subsection and subsection != current_subsection:
                escaped_subsection = html.escape(str(subsection))
                heading_html = (
                    f'<h{subsection_level} class="ricemind-figure-heading" '
                    'style="break-after: avoid; page-break-after: avoid;">'
                    f"{escaped_subsection}</h{subsection_level}>"
                )
                current_subsection = subsection
            group_style = "break-inside: avoid; page-break-inside: avoid;"
            if figure.get("placement") == "after_heading":
                group_style += " break-before: avoid; page-break-before: avoid;"
            block.extend(
                [
                    f'<div class="ricemind-figure-group" style="{group_style}">',
                    heading_html,
                    figure_html(
                        Path(figure["path"]),
                        path,
                        figure["caption"],
                        numbered[id(figure)],
                        max(55, min(94, int(figure["width_percent"]))),
                        max(3.2, min(7.2, float(figure["max_height_inches"]))),
                    ),
                    "</div>",
                    "",
                ]
            )
        block.append(BLOCK_END.format(block_id=block_id))
        insertions.append((insert_at, "\n".join(block)))

    for insert_at, block in sorted(insertions, key=lambda item: item[0], reverse=True):
        lines[insert_at:insert_at] = ["", block, ""]
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip() + "\n", encoding="utf-8")


def heading_level(paragraph: Any) -> Optional[int]:
    name = paragraph.style.name if paragraph.style else ""
    match = re.match(r"Heading\s+([1-6])$", name)
    return int(match.group(1)) if match else None


def fit_docx_size(image_path: Path, max_width: float, max_height: float) -> Tuple[float, float]:
    try:
        from PIL import Image

        with Image.open(image_path) as image:
            ratio = image.width / max(1, image.height)
    except Exception:
        ratio = 1.4
    width = min(max_width, max_height * ratio)
    height = width / ratio
    return width, min(height, max_height)


def update_docx(path: Path, figures: Sequence[Dict[str, Any]], fallback_section_title: Optional[str] = None) -> None:
    if not figures or not path.is_file():
        return
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX figure placement") from exc

    document = Document(path)
    existing_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    language = detect_language(existing_text)
    fallback_title = fallback_section_title or localized(language, "数据支持的可视化摘要", "Data-supported visual summary")
    figure_label = "图" if language == "zh" else "Figure"
    for index, raw_figure in enumerate(figures, 1):
        figure = dict(raw_figure)
        figure_path = Path(figure["path"])
        figure.setdefault("id", safe_name(figure_path.stem or f"figure_{index}"))
        figure.setdefault("title", humanize(figure["id"]))
        figure.setdefault("caption", figure["title"])
        figure.setdefault("subsection", figure["title"])
        figure.setdefault("section", "")
        figure.setdefault("section_keywords", [])
        figure.setdefault("placement", "end_of_section")
        figure.setdefault("max_height_inches", 6.4)
        caption_text = f"{figure_label} {index}. {figure['caption']}"
        if caption_text in existing_text:
            continue
        section = normalized_heading(figure.get("section", ""))
        keywords = [normalized_heading(item) for item in figure.get("section_keywords", [])]
        anchor = None
        anchor_level = 1
        for paragraph in document.paragraphs:
            level = heading_level(paragraph)
            normalized = normalized_heading(paragraph.text)
            if level and (
                (section and (section in normalized or normalized in section))
                or (keywords and any(keyword in normalized for keyword in keywords))
            ):
                anchor, anchor_level = paragraph, level
                break

        if anchor is None:
            anchor = next(
                (
                    paragraph
                    for paragraph in document.paragraphs
                    if heading_level(paragraph) == 1
                    and normalized_heading(paragraph.text) == normalized_heading(fallback_title)
                ),
                None,
            )
            if anchor is None:
                anchor = document.add_heading(fallback_title, level=1)
            anchor_level = 1

        insertion_before = None
        if figure.get("placement", "end_of_section") == "end_of_section":
            found_anchor = False
            for paragraph in document.paragraphs:
                if paragraph._p is anchor._p:
                    found_anchor = True
                    continue
                if found_anchor and (heading_level(paragraph) or 99) <= anchor_level:
                    insertion_before = paragraph._p
                    break

        created = []
        subsection = figure.get("subsection") or figure.get("title")
        if subsection:
            created.append(document.add_heading(subsection, level=min(6, anchor_level + 1))._p)
        width, height = fit_docx_size(figure_path, 6.2, max(3.2, min(7.2, float(figure["max_height_inches"]))))
        picture_paragraph = document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(figure["path"]), width=Inches(width), height=Inches(height))
        created.append(picture_paragraph._p)
        caption = document.add_paragraph(caption_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.size = Pt(9)
        created.append(caption._p)

        if insertion_before is not None:
            for element in created:
                insertion_before.addprevious(element)
        else:
            cursor = anchor._p
            for element in created:
                cursor.addnext(element)
                cursor = element
    document.save(path)


def load_plan(path: Path) -> Dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict) or not isinstance(data.get("figures"), list):
        raise ValueError("Figure plan must be a JSON object containing a 'figures' list")
    return data


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--data-dir", type=Path, required=True)
    parser.add_argument("--fig-dir", type=Path)
    parser.add_argument("--plan", type=Path, help="Personalized JSON figure plan; omit only for task-neutral automatic summaries")
    parser.add_argument("--write-auto-plan", type=Path, help="Write the task-neutral detected plan for inspection")
    parser.add_argument("--language", choices=["en", "zh"])
    parser.add_argument("--markdown", type=Path, help="Place figures in relevant Markdown sections")
    parser.add_argument("--docx", type=Path, help="Place figures in relevant DOCX sections")
    parser.add_argument("--manifest", type=Path, help="Optional JSON manifest path")
    args = parser.parse_args()

    report_text = args.markdown.read_text(encoding="utf-8") if args.markdown and args.markdown.is_file() else ""
    language = args.language or detect_language(report_text)
    plan = load_plan(args.plan) if args.plan else auto_plan(args.data_dir, language)
    plan.setdefault("language", language)
    if args.write_auto_plan:
        args.write_auto_plan.parent.mkdir(parents=True, exist_ok=True)
        args.write_auto_plan.write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")

    fig_dir = args.fig_dir or args.data_dir / "figures"
    figures = build_figures_from_data_dir(args.data_dir, fig_dir, plan, language)
    if args.markdown:
        update_markdown(args.markdown, figures, plan.get("fallback_section_title"))
    if args.docx:
        update_docx(args.docx, figures, plan.get("fallback_section_title"))
    manifest = args.manifest or args.data_dir / "figure_manifest.json"
    if figures:
        manifest.parent.mkdir(parents=True, exist_ok=True)
        manifest.write_text(json.dumps({"plan": plan, "figures": figures}, ensure_ascii=False, indent=2), encoding="utf-8")
    elif manifest.is_file():
        manifest.unlink()
    if fig_dir.is_dir():
        try:
            fig_dir.rmdir()
        except OSError:
            pass
    print(f"Generated {len(figures)} figure(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
