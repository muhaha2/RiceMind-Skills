#!/usr/bin/env python3
"""Build a fixed-template RiceMind gene mechanism DOCX report from full RiceMind API data."""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import sys
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from urllib.parse import parse_qs, quote, urlencode, urljoin, urlparse
from urllib.request import Request, urlopen


DEFAULT_API_BASE = "http://lit-evi.hzau.edu.cn/ricemind-api/"
DEFAULT_PAGE_LIMIT = 500
DEFAULT_MAX_PAGES = 10000
REPORT_TRAIT_ROWS = 20
REPORT_API_ROWS = 6
REPORT_MECHANISM_THEMES = 8

FONT_EAST_ASIA = "SimSun"
FONT_LATIN = "Times New Roman"
FONT_SIZES = {
    "title": 16,
    "heading1": 14,
    "heading2": 12,
    "body": 10.5,
    "table": 9,
    "caption": 9,
}

DEFAULT_ENDPOINTS: Dict[str, Dict[str, Any]] = {
    "ricemind_get_gene_profile": {
        "method": "GET",
        "path": "gene-profile/",
        "params": {"gene": "{gene}"},
    },
    "ricemind_get_traits_by_gene": {
        "method": "GET",
        "path": "traits-by-gene/",
        "params": {
            "gene": "{gene}",
            "confidence": "{confidence}",
            "onto_type": "{onto_type}",
            "page": "{page}",
            "limit": "{limit}",
        },
        "record_keys": ["associated_traits"],
    },
    "ricemind_get_varieties_by_gene": {
        "method": "GET",
        "path": "varieties-by-gene/",
        "params": {"gene": "{gene}", "page": "{page}", "limit": "{limit}"},
        "record_keys": ["varieties"],
    },
    "ricemind_search_by_gene": {
        "method": "GET",
        "path": "search-by-gene/",
        "params": {"gene": "{gene}", "page": "{page}", "limit": "{limit}"},
        "record_keys": ["results"],
    },
    "ricemind_search_by_trait_and_gene": {
        "method": "GET",
        "path": "search-by-trait-and-gene/",
        "params": {"gene": "{gene}", "trait": "{trait}", "page": "{page}", "limit": "{limit}"},
        "record_keys": ["sentence_evidence"],
    },
    "ricemind_get_sentence_context": {
        "method": "GET",
        "path": "sentence-context/",
        "params": {"pmid": "{pmid}", "sent_id": "{sent_id}", "window": "{window}"},
    },
    "ricemind_get_gene_omics_sequence": {
        "method": "GET",
        "path": "gene-omics-sequence/",
        "params": {"rap_id": "{rap_id}"},
    },
}

FIELD_ALIASES = {
    "gene": ["Gene_Symbol", "gene_symbol", "geneSymbol", "gene", "symbol", "Original_Gene_ID", "canonical_symbol"],
    "rap_id": ["RAP_ID", "rap_id", "rapId", "RAP", "locus", "locus_id", "standardized_id", "standard_rap_id", "primary_id"],
    "trait": ["Trait_Description", "trait_description", "traitDescription", "trait", "trait_name", "name", "trait_query"],
    "ontology_type": ["Ontology_Type", "ontology_type", "ontologyType", "ontology"],
    "ontology_id": ["Ontology_ID", "ontology_id", "ontologyId", "term_id", "id"],
    "evidence_code": ["Evidence_Code", "evidence_code", "evidenceCode", "evidence_codes"],
    "source_db": ["Source_DB", "source_db", "sourceDb", "source", "sources", "source_databases", "database"],
    "confidence": ["confidence", "Confidence", "confidence_tier", "evidence_tier", "tier", "category"],
    "support": [
        "supporting_article_count",
        "supportingArticleCount",
        "support_count",
        "supportCount",
        "article_count",
        "literature_support_count",
        "pmid_count",
        "row_count",
    ],
    "year": ["year", "Year", "publication_year", "publicationYear", "earliest_year", "first_mention_year"],
    "pmid": ["PMID", "pmid", "pubmed_id", "pubmedId"],
    "sentence_id": ["Sentence_ID", "sentence_id", "sentenceId", "sent_id", "sentId"],
    "sentence": ["text", "sentence", "evidence_sentence", "evidenceText", "Evidence_Text"],
    "title": ["title", "Title", "article_title"],
    "journal": ["journal", "Journal"],
    "doi": ["doi", "DOI"],
}

def stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "; ".join(stringify(item) for item in value if item not in (None, ""))
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def get_value(record: Any, canonical: str, default: str = "") -> str:
    if not isinstance(record, dict):
        return default
    lower_map = {str(k).lower(): k for k in record}
    for alias in FIELD_ALIASES.get(canonical, []):
        key = lower_map.get(alias.lower())
        if key is not None and record.get(key) not in (None, ""):
            return stringify(record[key])
    return default


def deep_find(obj: Any, candidate_keys: Sequence[str]) -> str:
    candidates = {key.lower() for key in candidate_keys}
    if isinstance(obj, dict):
        for key, value in obj.items():
            if str(key).lower() in candidates and value not in (None, ""):
                return stringify(value)
        for value in obj.values():
            found = deep_find(value, candidate_keys)
            if found:
                return found
    elif isinstance(obj, list):
        for item in obj:
            found = deep_find(item, candidate_keys)
            if found:
                return found
    return ""


def flatten_records(obj: Any) -> List[Dict[str, Any]]:
    if obj is None:
        return []
    if isinstance(obj, list):
        rows: List[Dict[str, Any]] = []
        for item in obj:
            rows.extend(flatten_records(item))
        return rows
    if isinstance(obj, dict):
        candidate_keys = [
            "records",
            "associated_traits",
            "sentence_evidence",
            "gta_metadata",
            "results",
            "items",
            "data",
            "associations",
            "traits",
            "evidence",
            "sentences",
            "rows",
            "varieties",
            "transcripts",
        ]
        for key in candidate_keys:
            if key in obj and isinstance(obj[key], list):
                return flatten_records(obj[key])
        return [obj]
    return []


def read_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    if not isinstance(data, dict):
        return {"data": data}
    return data


def write_json(path: Path, payload: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def render_params(params: Dict[str, Any], values: Dict[str, Any]) -> Dict[str, str]:
    rendered: Dict[str, str] = {}
    for key, value in params.items():
        if value is None or value == "":
            continue
        text = stringify(value)
        for name, replacement in values.items():
            text = text.replace("{" + name + "}", stringify(replacement))
        if text != "":
            rendered[key] = text
    return rendered


def build_url(base: str, endpoint: Dict[str, Any], values: Dict[str, Any]) -> str:
    path = stringify(endpoint["path"])
    for name, replacement in values.items():
        path = path.replace("{" + name + "}", quote(stringify(replacement)))
    url = urljoin(base.rstrip("/") + "/", path.lstrip("/"))
    params = render_params(endpoint.get("params", {}), values)
    if params:
        url += "?" + urlencode(params)
    return url


def call_get(base: str, endpoint: Dict[str, Any], values: Dict[str, Any], api_calls: List[Dict[str, Any]]) -> Any:
    method = stringify(endpoint.get("method", "GET")).upper()
    if method != "GET":
        raise ValueError(f"Only GET endpoints are supported by this helper, got {method}")
    url = build_url(base, endpoint, values)
    record_api_call(api_calls, endpoint, values, url)
    request = Request(url, headers={"Accept": "application/json"})
    with urlopen(request, timeout=60) as resp:
        return json.loads(resp.read().decode("utf-8"))


def record_api_call(api_calls: List[Dict[str, Any]], endpoint: Dict[str, Any], values: Dict[str, Any], url: str) -> None:
    tool = stringify(endpoint.get("tool", ""))
    path = stringify(endpoint.get("path", ""))
    group_values = {k: stringify(v) for k, v in values.items() if k not in {"page", "limit"}}
    query_value = group_values.get("trait") if tool == "ricemind_search_by_trait_and_gene" else ""
    if tool == "ricemind_search_by_trait_and_gene":
        group_values["trait"] = "<multiple traits>"
    key = json.dumps({"tool": tool, "path": path, "values": group_values}, sort_keys=True, ensure_ascii=False)
    page = parse_int(values.get("page"))
    limit = parse_int(values.get("limit"))
    for entry in api_calls:
        if entry.get("key") == key:
            entry["call_count"] = parse_int(entry.get("call_count")) + 1
            if page:
                entry["page_min"] = min(parse_int(entry.get("page_min")) or page, page)
                entry["page_max"] = max(parse_int(entry.get("page_max")) or page, page)
            if limit:
                entry["limit"] = limit
            entry["last_url"] = url
            if query_value:
                examples = entry.setdefault("query_examples", [])
                if query_value not in examples and len(examples) < 8:
                    examples.append(query_value)
                entry["query_count"] = parse_int(entry.get("query_count")) + (0 if query_value in examples[:-1] else 1)
            return
    entry = {
        "key": key,
        "tool": tool,
        "path": path,
        "parameters": group_values,
        "example_url": url,
        "call_count": 1,
        "page_min": page or "",
        "page_max": page or "",
        "limit": limit or "",
    }
    if query_value:
        entry["query_examples"] = [query_value]
        entry["query_count"] = 1
    api_calls.append(entry)


def compact_api_calls(api_calls: Any) -> List[Dict[str, Any]]:
    if not isinstance(api_calls, list):
        return []
    compacted: List[Dict[str, Any]] = []
    for raw in api_calls:
        if not isinstance(raw, dict):
            continue
        if "example_url" in raw and "path" in raw:
            item = {k: v for k, v in raw.items() if k != "key"}
            compacted.append(item)
            continue
        url = raw.get("url") or raw.get("example_url") or ""
        if not url:
            continue
        parsed = urlparse(url)
        query = {k: values[0] if values else "" for k, values in parse_qs(parsed.query).items()}
        path = parsed.path.strip("/").split("/")[-1] + "/"
        tool = raw.get("tool") or tool_from_path(path)
        group_params = {k: v for k, v in query.items() if k not in {"page", "limit"}}
        if tool == "ricemind_search_by_trait_and_gene":
            group_params["trait"] = "<multiple traits>"
        key = json.dumps({"tool": tool, "path": path, "values": group_params}, sort_keys=True, ensure_ascii=False)
        page = parse_int(query.get("page"))
        limit = parse_int(query.get("limit"))
        for entry in compacted:
            if entry.get("key") == key:
                entry["call_count"] = parse_int(entry.get("call_count")) + 1
                if page:
                    entry["page_min"] = min(parse_int(entry.get("page_min")) or page, page)
                    entry["page_max"] = max(parse_int(entry.get("page_max")) or page, page)
                if limit:
                    entry["limit"] = limit
                entry["last_url"] = url
                trait = query.get("trait", "")
                if trait:
                    examples = entry.setdefault("query_examples", [])
                    if trait not in examples and len(examples) < 8:
                        examples.append(trait)
                break
        else:
            entry = {
                "key": key,
                "tool": tool,
                "path": path,
                "parameters": group_params,
                "example_url": url,
                "call_count": 1,
                "page_min": page or "",
                "page_max": page or "",
                "limit": limit or "",
            }
            trait = query.get("trait", "")
            if trait:
                entry["query_examples"] = [trait]
            compacted.append(entry)
    return [{k: v for k, v in entry.items() if k != "key"} for entry in compacted]


def tool_from_path(path: str) -> str:
    mapping = {
        "gene-profile/": "ricemind_get_gene_profile",
        "gene-omics-sequence/": "ricemind_get_gene_omics_sequence",
        "traits-by-gene/": "ricemind_get_traits_by_gene",
        "varieties-by-gene/": "ricemind_get_varieties_by_gene",
        "search-by-gene/": "ricemind_search_by_gene",
        "search-by-trait-and-gene/": "ricemind_search_by_trait_and_gene",
        "sentence-context/": "ricemind_get_sentence_context",
    }
    return mapping.get(path, path.rstrip("/"))


def extract_records(payload: Any, record_keys: Sequence[str]) -> List[Dict[str, Any]]:
    if not isinstance(payload, dict):
        return flatten_records(payload)
    for key in record_keys:
        if isinstance(payload.get(key), list):
            return flatten_records(payload[key])
    if isinstance(payload.get("sentence_pagination"), dict) and isinstance(payload.get("sentence_evidence"), list):
        return flatten_records(payload["sentence_evidence"])
    return flatten_records(payload)


def find_total_pages(payload: Any) -> int:
    if not isinstance(payload, dict):
        return 0
    for key in ("total_pages", "totalPages"):
        if key in payload:
            return parse_int(payload[key])
    if isinstance(payload.get("sentence_pagination"), dict):
        for key in ("total_pages", "totalPages"):
            if key in payload["sentence_pagination"]:
                return parse_int(payload["sentence_pagination"][key])
    return 0


def find_total_count(payload: Any) -> int:
    if not isinstance(payload, dict):
        return 0
    for key in ("total_count", "total_associated_traits", "total_associated_varieties", "total_sentences"):
        if key in payload:
            return parse_int(payload[key])
    if isinstance(payload.get("sentence_pagination"), dict):
        for key in ("total_sentences", "total_count"):
            if key in payload["sentence_pagination"]:
                return parse_int(payload["sentence_pagination"][key])
    return 0


def collect_paginated(
    base: str,
    endpoint: Dict[str, Any],
    values: Dict[str, Any],
    record_keys: Sequence[str],
    page_limit: int,
    max_pages: int,
    api_calls: List[Dict[str, str]],
) -> Dict[str, Any]:
    records: List[Dict[str, Any]] = []
    pages: List[Dict[str, Any]] = []
    page = 1
    total_count = 0
    total_pages = 0
    while page <= max_pages:
        page_values = dict(values)
        page_values.update({"page": page, "limit": page_limit})
        payload = call_get(base, endpoint, page_values, api_calls)
        if isinstance(payload, dict):
            pages.append(payload)
        page_records = extract_records(payload, record_keys)
        records.extend(page_records)
        total_count = total_count or find_total_count(payload)
        total_pages = total_pages or find_total_pages(payload)
        if total_pages and page >= total_pages:
            break
        if not total_pages:
            if not page_records:
                break
            if total_count and len(records) >= total_count:
                break
            if len(page_records) < page_limit:
                break
        page += 1
    else:
        raise RuntimeError(f"Exceeded max_pages={max_pages} for endpoint {endpoint.get('path')}")
    return {
        "records": records,
        "pages": pages,
        "total_count": total_count or len(records),
        "total_pages": total_pages or page,
    }


def fetch_from_api(
    base: str,
    gene: str,
    endpoint_map_path: Optional[Path],
    page_limit: int,
    max_pages: int,
    fetch_trait_evidence: bool,
) -> Dict[str, Any]:
    endpoints = dict(DEFAULT_ENDPOINTS)
    if endpoint_map_path:
        endpoints.update(read_json(endpoint_map_path))
    for name, endpoint in endpoints.items():
        endpoint.setdefault("tool", name)

    api_calls: List[Dict[str, str]] = []
    bundle: Dict[str, Any] = {"api_base": base, "api_calls": api_calls}
    bundle["gene_profile"] = call_get(base, endpoints["ricemind_get_gene_profile"], {"gene": gene}, api_calls)

    profile = choose_profile(bundle)
    rap_id = extract_standard_rap_id(profile) or gene
    if rap_id:
        try:
            bundle["omics_sequence"] = call_get(
                base,
                endpoints["ricemind_get_gene_omics_sequence"],
                {"rap_id": rap_id},
                api_calls,
            )
        except Exception as exc:
            bundle["omics_sequence_error"] = str(exc)

    try:
        bundle["traits_by_gene"] = collect_paginated(
            base,
            endpoints["ricemind_get_traits_by_gene"],
            {"gene": gene, "confidence": "All", "onto_type": "ALL"},
            endpoints["ricemind_get_traits_by_gene"].get("record_keys", ["associated_traits"]),
            page_limit,
            max_pages,
            api_calls,
        )
    except Exception as exc:
        bundle["traits_by_gene"] = {"records": [], "total_count": 0, "total_pages": 0, "error": str(exc)}
        bundle["traits_by_gene_error"] = str(exc)
    try:
        bundle["varieties_by_gene"] = collect_paginated(
            base,
            endpoints["ricemind_get_varieties_by_gene"],
            {"gene": gene},
            endpoints["ricemind_get_varieties_by_gene"].get("record_keys", ["varieties"]),
            page_limit,
            max_pages,
            api_calls,
        )
    except Exception as exc:
        bundle["varieties_by_gene_error"] = str(exc)
    try:
        bundle["gene_sentences"] = collect_paginated(
            base,
            endpoints["ricemind_search_by_gene"],
            {"gene": gene},
            endpoints["ricemind_search_by_gene"].get("record_keys", ["results"]),
            page_limit,
            max_pages,
            api_calls,
        )
    except Exception as exc:
        bundle["gene_sentences_error"] = str(exc)

    if fetch_trait_evidence:
        trait_evidence: Dict[str, Any] = {}
        for trait_record in normalize_traits(bundle):
            trait = trait_record["trait"]
            if not trait or trait in trait_evidence:
                continue
            try:
                trait_evidence[trait] = collect_paginated(
                    base,
                    endpoints["ricemind_search_by_trait_and_gene"],
                    {"gene": gene, "trait": trait},
                    endpoints["ricemind_search_by_trait_and_gene"].get("record_keys", ["sentence_evidence"]),
                    page_limit,
                    max_pages,
                    api_calls,
                )
            except Exception as exc:
                trait_evidence[trait] = {"records": [], "error": str(exc)}
        bundle["trait_evidence"] = trait_evidence
    return bundle


def normalize_api_confidence(record: Dict[str, Any]) -> str:
    """Normalize an API-provided tier label without inferring a tier from metadata."""
    explicit = get_value(record, "confidence")
    if not explicit:
        return "Unspecified"
    normalized = explicit.upper().replace("-", "_")
    if normalized == "HIGH" or re.match(r"^TIER[ _]?1(?:\b|_)", normalized):
        return "High"
    if normalized == "MEDIUM" or re.match(r"^TIER[ _]?2(?:\b|_)", normalized):
        return "Medium"
    if normalized == "LOW" or re.match(r"^TIER[ _]?3(?:\b|_)", normalized):
        return "Low"
    return explicit


def parse_int(value: Any) -> int:
    try:
        return int(float(stringify(value)))
    except Exception:
        return 0


def normalize_traits(bundle: Dict[str, Any]) -> List[Dict[str, str]]:
    source = (
        bundle.get("traits_by_gene")
        or bundle.get("ricemind_get_traits_by_gene")
        or bundle.get("traits")
        or bundle.get("associations")
    )
    rows = []
    for rec in flatten_records(source):
        trait = get_value(rec, "trait")
        if not trait:
            continue
        ontology_id = get_value(rec, "ontology_id")
        rows.append(
            {
                "gene": get_value(rec, "gene") or stringify(bundle.get("gene", "")),
                "rap_id": get_value(rec, "rap_id"),
                "trait": trait,
                "ontology_type": get_value(rec, "ontology_type") or infer_ontology_type(ontology_id),
                "ontology_id": ontology_id,
                "evidence_code": get_value(rec, "evidence_code"),
                "source_db": get_value(rec, "source_db"),
                "confidence": normalize_api_confidence(rec),
                "support": stringify(parse_int(get_value(rec, "support"))),
                "year": get_value(rec, "year"),
            }
        )
    return dedupe_rows(rows, ["trait", "ontology_id", "confidence", "evidence_code", "source_db"])


def infer_ontology_type(ontology_id: str) -> str:
    if ":" not in ontology_id:
        return ""
    prefix = ontology_id.split(":", 1)[0].strip().upper()
    return prefix or ""


def normalize_varieties(bundle: Dict[str, Any]) -> List[Dict[str, str]]:
    rows = []
    source = bundle.get("varieties_by_gene")
    if isinstance(source, dict) and isinstance(source.get("records"), list):
        raw_records = source["records"]
    elif isinstance(source, dict) and isinstance(source.get("varieties"), list):
        raw_records = source["varieties"]
    elif isinstance(source, list):
        raw_records = source
    else:
        raw_records = flatten_records(source)
    for rec in raw_records:
        if isinstance(rec, dict):
            variety = get_value(rec, "trait") or get_value(rec, "gene") or stringify(rec.get("normalization", ""))
        else:
            variety = stringify(rec)
        if variety:
            rows.append({"variety": variety})
    return dedupe_rows(rows, ["variety"])


def trait_lookup(traits: List[Dict[str, str]]) -> Dict[str, Dict[str, str]]:
    lookup: Dict[str, Dict[str, str]] = {}
    for row in traits:
        key = row["trait"].strip().lower()
        if key and key not in lookup:
            lookup[key] = row
    return lookup


def normalize_evidence(bundle: Dict[str, Any], traits: Optional[List[Dict[str, str]]] = None) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    evidence_sources: List[Tuple[str, Any]] = []
    for key in ("trait_evidence", "evidence_by_trait"):
        if isinstance(bundle.get(key), dict):
            evidence_sources.extend(bundle[key].items())
    for key in ("gene_sentences", "evidence", "sentences", "ricemind_search_by_gene"):
        if key in bundle:
            evidence_sources.append(("", bundle[key]))

    lookup = trait_lookup(traits or [])
    for trait_name, source in evidence_sources:
        trait_meta = lookup.get(trait_name.strip().lower(), {})
        for rec in flatten_records(source):
            sentence = get_value(rec, "sentence")
            pmid = get_value(rec, "pmid")
            if not sentence and not pmid:
                continue
            row_trait = get_value(rec, "trait", trait_name) or trait_name
            meta = lookup.get(row_trait.strip().lower(), trait_meta)
            rows.append(
                {
                    "trait": row_trait,
                    "ontology_type": get_value(rec, "ontology_type") or meta.get("ontology_type", ""),
                    "ontology_id": get_value(rec, "ontology_id") or meta.get("ontology_id", ""),
                    "evidence_code": get_value(rec, "evidence_code") or meta.get("evidence_code", ""),
                    "source_db": get_value(rec, "source_db") or meta.get("source_db", ""),
                    "confidence": normalize_api_confidence(rec) if get_value(rec, "confidence") else meta.get("confidence", "Unspecified"),
                    "pmid": pmid,
                    "sentence_id": get_value(rec, "sentence_id"),
                    "year": get_value(rec, "year"),
                    "title": get_value(rec, "title"),
                    "journal": get_value(rec, "journal"),
                    "doi": get_value(rec, "doi"),
                    "sentence": sentence,
                }
            )
    return dedupe_rows(rows, ["trait", "pmid", "sentence_id", "sentence"])


def dedupe_rows(rows: List[Dict[str, str]], keys: Sequence[str]) -> List[Dict[str, str]]:
    seen = set()
    deduped = []
    for row in rows:
        key = tuple(row.get(k, "") for k in keys)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def choose_profile(bundle: Dict[str, Any]) -> Dict[str, Any]:
    source = bundle.get("gene_profile") or bundle.get("ricemind_get_gene_profile") or bundle.get("profile") or {}
    if isinstance(source, dict) and isinstance(source.get("data"), dict):
        return source["data"]
    if isinstance(source, dict):
        return source
    records = flatten_records(source)
    return records[0] if records else {}


def extract_standard_rap_id(profile: Dict[str, Any]) -> str:
    return deep_find(profile, ["standard_rap_id", "rap_id", "primary_id", "gene_id"])


def safe_filename(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", text).strip("_") or "gene"


def remove_output_file(path: Path) -> None:
    if path.is_file():
        path.unlink()


def remove_empty_directory(path: Path) -> None:
    if path.is_dir():
        try:
            path.rmdir()
        except OSError:
            pass


def write_csv(path: Path, rows: List[Dict[str, str]], fieldnames: Optional[Sequence[str]] = None) -> bool:
    if not rows:
        remove_output_file(path)
        return False
    path.parent.mkdir(parents=True, exist_ok=True)
    columns = list(fieldnames or rows[0].keys())
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=columns)
        writer.writeheader()
        writer.writerows(rows)
    return True


PAYLOAD_METADATA_KEYS = {
    "api_calls",
    "current_page",
    "endpoint",
    "error",
    "errors",
    "limit",
    "max_pages",
    "message",
    "next",
    "page",
    "pages_retrieved",
    "params",
    "previous",
    "query",
    "request",
    "status",
    "status_code",
    "total_count",
    "total_pages",
}


def has_usable_payload_data(value: Any) -> bool:
    if value in (None, "", False):
        return False
    if isinstance(value, dict):
        return any(
            key.lower() not in PAYLOAD_METADATA_KEYS and has_usable_payload_data(item)
            for key, item in value.items()
        )
    if isinstance(value, (list, tuple, set)):
        return any(has_usable_payload_data(item) for item in value)
    if isinstance(value, (int, float)):
        return value != 0
    return True


def bundle_has_usable_data(bundle: Dict[str, Any]) -> bool:
    return any(
        key not in {"gene", "api_base", "api_calls"}
        and not key.endswith("_error")
        and has_usable_payload_data(value)
        for key, value in bundle.items()
    )


def plot_counter(counter: Counter, title: str, out_path: Path, xlabel: str = "Count", top_n: int = 20) -> Optional[Path]:
    if not counter:
        return None
    try:
        import matplotlib.pyplot as plt
    except Exception:
        return plot_counter_pillow(counter, title, out_path, xlabel=xlabel, top_n=top_n)
    labels, values = zip(*counter.most_common(top_n))
    fig_height = max(3.2, 0.32 * len(labels) + 1.4)
    fig, ax = plt.subplots(figsize=(7.4, fig_height))
    ax.barh(range(len(labels)), values, color="#3f6f8f")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels)
    ax.invert_yaxis()
    ax.set_xlabel(xlabel)
    ax.set_title(title)
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=220)
    plt.close(fig)
    return out_path


def plot_years(years: Iterable[str], out_path: Path) -> Optional[Path]:
    clean = [parse_int(y) for y in years if parse_int(y) > 0]
    if not clean:
        return None
    try:
        import matplotlib.pyplot as plt
    except Exception:
        return plot_years_pillow(Counter(clean), out_path)
    counts = Counter(clean)
    xs = sorted(counts)
    ys = [counts[x] for x in xs]
    fig, ax = plt.subplots(figsize=(7.4, 3.8))
    ax.bar(xs, ys, color="#6b8f3f")
    ax.set_xlabel("Publication year")
    ax.set_ylabel("Evidence records")
    ax.set_title("RiceMind evidence by publication year")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=220)
    plt.close(fig)
    return out_path


def plot_top_traits(traits: List[Dict[str, str]], out_path: Path) -> Optional[Path]:
    counter = Counter({row["trait"]: parse_int(row["support"]) for row in traits if row["trait"]})
    return plot_counter(counter, "Top traits by supporting articles", out_path, xlabel="Supporting articles", top_n=25)


def load_font(size: int, bold: bool = False) -> Any:
    try:
        from PIL import ImageFont

        candidates = [
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
            "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except Exception:
                continue
        return ImageFont.load_default()
    except Exception:
        return None


def text_width(draw: Any, text: str, font: Any) -> int:
    try:
        box = draw.textbbox((0, 0), text, font=font)
        return box[2] - box[0]
    except Exception:
        return len(text) * 8


def ellipsize(draw: Any, text: str, font: Any, max_width: int) -> str:
    text = stringify(text)
    if text_width(draw, text, font) <= max_width:
        return text
    while text and text_width(draw, text + "...", font) > max_width:
        text = text[:-1]
    return text + "..." if text else "..."


def plot_counter_pillow(counter: Counter, title: str, out_path: Path, xlabel: str = "Count", top_n: int = 20) -> Optional[Path]:
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None
    items = counter.most_common(top_n)
    if not items:
        return None
    label_font = load_font(14)
    title_font = load_font(18, bold=True)
    axis_font = load_font(12)
    left = 250
    right = 70
    top = 70
    row_h = 28
    bottom = 55
    width = 900
    height = top + len(items) * row_h + bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 18), title, fill="#1f2933", font=title_font)
    max_value = max(value for _, value in items) or 1
    bar_area = width - left - right
    for idx, (label, value) in enumerate(items):
        y = top + idx * row_h
        draw.text((20, y + 4), ellipsize(draw, label, label_font, left - 35), fill="#111827", font=label_font)
        bar_w = max(2, int(bar_area * (value / max_value)))
        draw.rectangle((left, y + 5, left + bar_w, y + row_h - 6), fill="#3f6f8f")
        draw.text((left + bar_w + 6, y + 3), str(value), fill="#111827", font=axis_font)
    draw.line((left, top - 8, left, top + len(items) * row_h), fill="#9ca3af", width=1)
    draw.text((left, height - 32), xlabel, fill="#374151", font=axis_font)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(out_path)
    return out_path


def plot_years_pillow(counts: Counter, out_path: Path) -> Optional[Path]:
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None
    if not counts:
        return None
    title_font = load_font(18, bold=True)
    axis_font = load_font(12)
    width = 950
    height = 420
    left = 60
    right = 30
    top = 60
    bottom = 70
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 18), "RiceMind evidence by publication year", fill="#1f2933", font=title_font)
    years = sorted(counts)
    max_value = max(counts.values()) or 1
    plot_w = width - left - right
    plot_h = height - top - bottom
    bar_gap = 2
    bar_w = max(3, int(plot_w / max(1, len(years))) - bar_gap)
    draw.line((left, top, left, top + plot_h), fill="#9ca3af", width=1)
    draw.line((left, top + plot_h, left + plot_w, top + plot_h), fill="#9ca3af", width=1)
    tick_step = max(1, math.ceil(len(years) / 12))
    for idx, year in enumerate(years):
        value = counts[year]
        x = left + idx * (plot_w / max(1, len(years)))
        h = int(plot_h * (value / max_value))
        draw.rectangle((int(x), top + plot_h - h, int(x) + bar_w, top + plot_h), fill="#6b8f3f")
        if idx % tick_step == 0:
            draw.text((int(x) - 8, top + plot_h + 8), str(year), fill="#374151", font=axis_font)
    draw.text((left, height - 26), "Publication year", fill="#374151", font=axis_font)
    draw.text((left + 5, top + 5), f"max={max_value}", fill="#374151", font=axis_font)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(out_path)
    return out_path


def split_counter(rows: Iterable[Dict[str, str]], field: str) -> Counter:
    counter: Counter = Counter()
    for row in rows:
        value = row.get(field, "")
        for token in re.split(r"[;,|]\s*", value):
            token = token.strip()
            if token:
                counter[token] += 1
    return counter


def api_call_rows(api_calls: List[Dict[str, Any]], max_rows: int = REPORT_API_ROWS) -> List[List[str]]:
    rows = [["No.", "API function", "Call summary", "Example URL"]]
    for idx, call in enumerate(api_calls[:max_rows], 1):
        page_summary = ""
        if call.get("page_min") or call.get("page_max"):
            page_summary = f"pages {call.get('page_min')}-{call.get('page_max')}, limit={call.get('limit')}"
        call_summary = "; ".join(part for part in [
            endpoint_purpose(stringify(call.get("tool") or call.get("path"))),
            page_summary,
            f"calls={call.get('call_count')}" if call.get("call_count") else "",
            f"examples={'; '.join(call.get('query_examples', [])[:3])}" if call.get("query_examples") else "",
        ] if part)
        rows.append([
            str(idx),
            stringify(call.get("tool") or call.get("path")),
            truncate(call_summary, 160),
            truncate(stringify(call.get("example_url") or call.get("url")), 160),
        ])
    return rows


def endpoint_error_rows(bundle: Dict[str, Any]) -> List[List[str]]:
    rows = [["Endpoint", "Error"]]
    for key, value in bundle.items():
        if key.endswith("_error") and value:
            rows.append([key[:-6], truncate(stringify(value), 240)])
    return rows if len(rows) > 1 else []


def endpoint_purpose(tool: str) -> str:
    purposes = {
        "ricemind_get_gene_profile": "gene profile and external links",
        "ricemind_get_gene_omics_sequence": "omics sequence retrieval",
        "ricemind_get_traits_by_gene": "full GTA trait landscape",
        "ricemind_get_varieties_by_gene": "gene-variety co-occurrence",
        "ricemind_search_by_gene": "gene-wide sentence evidence",
        "ricemind_search_by_trait_and_gene": "trait-specific sentence evidence",
        "ricemind_get_sentence_context": "sentence context",
    }
    return purposes.get(tool, "")


def add_table(doc: Any, rows: List[List[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for i, value in enumerate(rows[0]):
        table.rows[0].cells[i].text = stringify(value)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            add_cell_value(cells[i], value)


def is_url(value: str) -> bool:
    return bool(re.match(r"^https?://[^\s]+$", stringify(value).strip()))


def add_cell_value(cell: Any, value: Any) -> None:
    text = stringify(value)
    if is_url(text):
        paragraph = cell.paragraphs[0]
        paragraph.text = ""
        add_hyperlink(paragraph, text, text, size=FONT_SIZES["table"])
    else:
        cell.text = text


def add_hyperlink(paragraph: Any, url: str, text: str, size: float = FONT_SIZES["body"]) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE
    except Exception:
        paragraph.add_run(text)
        return

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(rpr)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_picture_if_exists(doc: Any, path: Optional[Path], caption: str, width_inches: float = 6.2) -> None:
    if not path or not path.exists():
        return
    try:
        from docx.shared import Inches

        doc.add_picture(str(path), width=Inches(width_inches))
        doc.add_paragraph(caption)
    except Exception:
        pass


def zh(enabled: bool, cn: str, en: str) -> str:
    return cn if enabled else en


def truncate(text: str, max_chars: int) -> str:
    text = " ".join(stringify(text).split())
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3] + "..."


def configure_document_styles(doc: Any) -> None:
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except Exception:
        return

    style_specs = {
        "Normal": (FONT_SIZES["body"], False),
        "Title": (FONT_SIZES["title"], True),
        "Heading 1": (FONT_SIZES["heading1"], True),
        "Heading 2": (FONT_SIZES["heading2"], True),
        "Heading 3": (FONT_SIZES["body"], True),
    }
    for style_name, (size, bold) in style_specs.items():
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_LATIN
            style.font.size = Pt(size)
            style.font.bold = bold
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        except Exception:
            continue


def apply_document_fonts(doc: Any) -> None:
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        if style_name == "Title":
            size = FONT_SIZES["title"]
            bold = True
        elif style_name == "Heading 1":
            size = FONT_SIZES["heading1"]
            bold = True
        elif style_name == "Heading 2":
            size = FONT_SIZES["heading2"]
            bold = True
        else:
            size = FONT_SIZES["body"]
            bold = None
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=FONT_SIZES["table"], bold=(row_idx == 0))


def set_run_font(run: Any, size: float, bold: Optional[bool] = None) -> None:
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except Exception:
        return
    run.font.name = FONT_LATIN
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


def flatten_profile_rows(profile: Dict[str, Any]) -> List[List[str]]:
    rows = [["Field", "Value"]]

    def walk(prefix: str, value: Any) -> None:
        if isinstance(value, dict):
            for key, subvalue in value.items():
                walk(f"{prefix}.{key}" if prefix else stringify(key), subvalue)
        elif isinstance(value, list):
            rows.append([prefix, truncate("; ".join(stringify(item) for item in value), 500)])
        elif value not in (None, ""):
            rows.append([prefix, truncate(stringify(value), 500)])

    walk("", profile)
    return rows


def sequence_api_url_rows(bundle: Dict[str, Any], profile: Dict[str, Any], gene: str) -> List[List[str]]:
    urls: List[str] = []
    for call in compact_api_calls(bundle.get("api_calls") or []):
        tool = stringify(call.get("tool") or call.get("path"))
        if tool in {"ricemind_get_gene_omics_sequence", "gene-omics-sequence/"}:
            url = stringify(call.get("example_url") or call.get("url") or call.get("last_url"))
            if url and url not in urls:
                urls.append(url)
    if not urls:
        rap_id = extract_standard_rap_id(profile) or gene
        if rap_id:
            urls.append(build_url(bundle.get("api_base", DEFAULT_API_BASE), DEFAULT_ENDPOINTS["ricemind_get_gene_omics_sequence"], {"rap_id": rap_id}))
    if not urls:
        return []
    rows = [["No.", "RiceMind sequence API URL"]]
    for idx, url in enumerate(urls, 1):
        rows.append([str(idx), url])
    return rows


GENERIC_KEYWORD_STOPWORDS = {
    "about",
    "above",
    "across",
    "after",
    "against",
    "also",
    "among",
    "analysis",
    "article",
    "because",
    "between",
    "both",
    "could",
    "database",
    "detected",
    "different",
    "during",
    "effect",
    "evidence",
    "family",
    "gene",
    "genes",
    "genetic",
    "genome",
    "identified",
    "including",
    "indicated",
    "indicates",
    "lines",
    "major",
    "mapped",
    "method",
    "more",
    "novel",
    "only",
    "other",
    "plant",
    "plants",
    "reported",
    "results",
    "rice",
    "sativa",
    "showed",
    "shows",
    "study",
    "than",
    "that",
    "their",
    "these",
    "this",
    "those",
    "through",
    "trait",
    "traits",
    "under",
    "using",
    "were",
    "with",
}

TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z0-9_-]{2,}")


def evidence_summary_rows(traits: List[Dict[str, str]], evidence: List[Dict[str, str]]) -> List[List[str]]:
    return [
        ["Metric", "Value"],
        ["Trait associations", str(len(traits))],
        ["Sentence evidence records", str(len(evidence))],
        ["Unique PMIDs", str(len({row.get("pmid", "") for row in evidence if row.get("pmid", "")}))],
        ["Ontology types in traits", str(len({row.get("ontology_type", "") for row in traits if row.get("ontology_type", "")}))],
        ["Evidence years", year_span_text(evidence)],
    ]


def year_span_text(rows: List[Dict[str, str]]) -> str:
    years = sorted({parse_int(row.get("year", "")) for row in rows if parse_int(row.get("year", "")) > 0})
    if not years:
        return "unavailable"
    return f"{years[0]}-{years[-1]}"


def normalize_keyword(token: str) -> str:
    token = token.strip(".,;:()[]{}'\"").replace("_", "-")
    if token.isupper() and len(token) <= 12:
        return token
    return token.lower()


def extract_keywords_from_text(text: str, limit: int = 12, exclude: Sequence[str] = ()) -> List[str]:
    excluded = {normalize_keyword(item) for item in exclude if item}
    counter: Counter = Counter()
    for raw in TOKEN_RE.findall(text or ""):
        token = normalize_keyword(raw)
        if len(token) < 3:
            continue
        if token in GENERIC_KEYWORD_STOPWORDS or token in excluded:
            continue
        if token.isdigit():
            continue
        counter[token] += 1
    return [term for term, _ in counter.most_common(limit)]


def extract_keywords_from_rows(
    rows: List[Dict[str, str]],
    limit: int = 12,
    fields: Sequence[str] = ("trait", "sentence", "title"),
    exclude: Sequence[str] = (),
) -> List[str]:
    text = " ".join(row.get(field, "") for row in rows for field in fields)
    return extract_keywords_from_text(text, limit=limit, exclude=exclude)


def join_items(items: Sequence[str], is_zh: bool = False) -> str:
    return ("、" if is_zh else ", ").join(item for item in items if item)


def ontology_distribution_rows(traits: List[Dict[str, str]], evidence: List[Dict[str, str]]) -> List[List[str]]:
    trait_counts = Counter(row.get("ontology_type", "") or "Unspecified" for row in traits)
    evidence_counts = Counter(row.get("ontology_type", "") or "Unspecified" for row in evidence)
    ontologies = sorted(set(trait_counts) | set(evidence_counts), key=lambda key: (-(trait_counts[key] + evidence_counts[key]), key))
    rows = [["Ontology", "Trait records", "Evidence records", "Top trait examples"]]
    for ontology in ontologies[:12]:
        examples = []
        for row in traits:
            if (row.get("ontology_type", "") or "Unspecified") == ontology and row.get("trait", "") not in examples:
                examples.append(row.get("trait", ""))
            if len(examples) >= 4:
                break
        rows.append([ontology, str(trait_counts[ontology]), str(evidence_counts[ontology]), truncate("; ".join(examples), 160)])
    return rows


def representative_sentence_rows(evidence: List[Dict[str, str]], per_tier: int = 2) -> List[List[str]]:
    rows = [["Tier", "Trait / Ontology", "Evidence code / Source", "PMID / Year", "Representative RiceMind sentence"]]
    tier_order = ["High", "Medium", "Low", "Unspecified"]
    seen = set()
    for tier in tier_order:
        candidates = [row for row in evidence if (row.get("confidence", "") or "Unspecified") == tier and row.get("sentence", "")]
        candidates.sort(key=lambda row: (0 if row.get("pmid", "") else 1, -len(row.get("sentence", "")), row.get("trait", "")))
        added = 0
        for row in candidates:
            key = (tier, row.get("pmid", ""), row.get("sentence_id", ""), row.get("sentence", "")[:80])
            if key in seen:
                continue
            seen.add(key)
            rows.append([
                tier,
                " ".join(x for x in [truncate(row.get("trait", ""), 70), row.get("ontology_id", "")] if x),
                " / ".join(x for x in [truncate(row.get("evidence_code", ""), 60), truncate(row.get("source_db", ""), 60)] if x),
                " / ".join(x for x in [row.get("pmid", ""), row.get("year", "")] if x),
                truncate(row.get("sentence", ""), 260),
            ])
            added += 1
            if added >= per_tier:
                break
    return rows


def phase_ranges(years: List[int]) -> List[Tuple[str, int, int]]:
    unique_years = sorted(set(years))
    if not unique_years:
        return []
    if len(unique_years) <= 3:
        return [(f"Phase {idx + 1}", year, year) for idx, year in enumerate(unique_years)]
    cut1 = unique_years[max(0, len(unique_years) // 3 - 1)]
    cut2 = unique_years[max(0, (2 * len(unique_years)) // 3 - 1)]
    return [
        ("Phase I", unique_years[0], cut1),
        ("Phase II", cut1 + 1, cut2),
        ("Phase III", cut2 + 1, unique_years[-1]),
    ]


def temporal_hotspot_rows(evidence: List[Dict[str, str]], is_zh: bool) -> List[List[str]]:
    years = [parse_int(row.get("year", "")) for row in evidence if parse_int(row.get("year", "")) > 0]
    rows = [["Phase", "Years", "Evidence records", "Unique PMIDs", "Hotspot terms", "Interpretation"]]
    for phase, start, end in phase_ranges(years):
        phase_rows = [row for row in evidence if start <= parse_int(row.get("year", "")) <= end]
        if not phase_rows:
            continue
        terms = extract_keywords_from_rows(phase_rows, 8)
        top_traits = [trait for trait, _ in Counter(row.get("trait", "") for row in phase_rows if row.get("trait", "")).most_common(5)]
        interpretation = zh(
            is_zh,
            f"热点词由本阶段 RiceMind 句子文本自动抽取：{join_items(terms, True) if terms else '可追溯句证'}；相关 trait/语境包括 {join_items(top_traits, True) if top_traits else '未明确'}。",
            f"Hotspot terms are extracted from phase-specific RiceMind sentence text: {join_items(terms) if terms else 'traceable evidence'}; associated traits/contexts include {join_items(top_traits) if top_traits else 'unspecified'}.",
        )
        rows.append([
            phase,
            f"{start}-{end}",
            str(len(phase_rows)),
            str(len({row.get("pmid", "") for row in phase_rows if row.get("pmid", "")})),
            truncate(join_items(terms, is_zh), 160),
            truncate(interpretation, 240),
        ])
    return rows


def bibliometric_rows(evidence: List[Dict[str, str]]) -> Tuple[List[List[str]], List[List[str]]]:
    journal_counter = Counter(row.get("journal", "").strip() for row in evidence if row.get("journal", "").strip())
    journal_rows = [["Journal", "Evidence records"]]
    for journal, count in journal_counter.most_common(12):
        journal_rows.append([truncate(journal, 120), str(count)])

    pmid_counter = Counter(row.get("pmid", "").strip() for row in evidence if row.get("pmid", "").strip())
    title_by_pmid: Dict[str, str] = {}
    year_by_pmid: Dict[str, str] = {}
    for row in evidence:
        pmid = row.get("pmid", "").strip()
        if not pmid:
            continue
        title_by_pmid.setdefault(pmid, row.get("title", ""))
        year_by_pmid.setdefault(pmid, row.get("year", ""))
    pmid_rows = [["PMID", "Year", "Evidence records", "Title"]]
    for pmid, count in pmid_counter.most_common(12):
        pmid_rows.append([pmid, year_by_pmid.get(pmid, ""), str(count), truncate(title_by_pmid.get(pmid, ""), 180)])
    return (journal_rows if len(journal_rows) > 1 else [], pmid_rows if len(pmid_rows) > 1 else [])


def row_limit(rows: List[Dict[str, str]], max_rows: int) -> List[Dict[str, str]]:
    if max_rows and max_rows > 0:
        return rows[:max_rows]
    return rows


def collect_pmids(rows: List[Dict[str, str]], limit: int = 8) -> List[str]:
    counts: Counter = Counter(row.get("pmid", "").strip() for row in rows if row.get("pmid", "").strip())
    if counts:
        return [pmid for pmid, _ in counts.most_common(limit)]
    pmids = []
    for row in rows:
        pmid = row.get("pmid", "").strip()
        if pmid and pmid not in pmids:
            pmids.append(pmid)
        if len(pmids) >= limit:
            break
    return pmids


def pmid_citation(pmids: List[str]) -> str:
    if not pmids:
        return ""
    return "[" + ", ".join(pmids) + "]"


def top_traits_for_rows(rows: List[Dict[str, str]], traits: List[Dict[str, str]], limit: int = 6) -> List[str]:
    support = {row.get("trait", ""): parse_int(row.get("support")) for row in traits}
    counter: Counter = Counter()
    for row in rows:
        trait = row.get("trait", "").strip()
        if trait:
            counter[trait] += 1 + support.get(trait, 0) / 1000
    return [trait for trait, _ in counter.most_common(limit)]


def confidence_distribution_text(rows: List[Dict[str, str]]) -> str:
    counter = Counter(row.get("confidence", "") or "Unspecified" for row in rows)
    return ", ".join(f"{tier}={count}" for tier, count in counter.most_common())


def evidence_priority(row: Dict[str, str]) -> Tuple[int, int, int, str]:
    tier_rank = {"High": 0, "Medium": 1, "Low": 2, "Unspecified": 3}.get(row.get("confidence", "") or "Unspecified", 3)
    has_pmid = 0 if row.get("pmid", "") else 1
    sentence_len = -len(row.get("sentence", ""))
    return (tier_rank, has_pmid, sentence_len, row.get("trait", ""))


def representative_evidence_records(rows: List[Dict[str, str]], limit: int = 8) -> List[Dict[str, str]]:
    selected: List[Dict[str, str]] = []
    seen = set()
    for row in sorted(rows, key=evidence_priority):
        key = (row.get("pmid", ""), row.get("sentence_id", ""), row.get("sentence", "")[:100])
        if key in seen:
            continue
        seen.add(key)
        selected.append({
            "trait": row.get("trait", ""),
            "confidence": row.get("confidence", "") or "Unspecified",
            "evidence_code": row.get("evidence_code", ""),
            "source_db": row.get("source_db", ""),
            "pmid": row.get("pmid", ""),
            "year": row.get("year", ""),
            "journal": row.get("journal", ""),
            "sentence_id": row.get("sentence_id", ""),
            "sentence": row.get("sentence", ""),
        })
        if len(selected) >= limit:
            break
    return selected


def mechanism_topic_terms(row: Dict[str, str], gene: str = "", limit: int = 16) -> List[str]:
    text = " ".join(part for part in (row.get("sentence", ""), row.get("title", "")) if part)
    return extract_keywords_from_text(text, limit=limit, exclude=[gene])


def data_driven_topic_groups(
    evidence: List[Dict[str, str]],
    gene: str = "",
    limit: int = REPORT_MECHANISM_THEMES,
) -> List[Tuple[str, List[Dict[str, str]]]]:
    """Cluster evidence by sentence/title terminology, never by trait labels."""
    if not evidence or limit <= 0:
        return []

    clusters: List[Dict[str, Any]] = []
    for row in sorted(evidence, key=evidence_priority):
        terms = mechanism_topic_terms(row, gene)
        term_set = set(terms)
        best_index = -1
        best_score = 0.0
        for index, cluster in enumerate(clusters):
            signature = {term for term, _ in cluster["terms"].most_common(12)}
            if not term_set or not signature:
                continue
            overlap = len(term_set & signature)
            score = overlap / max(1, min(len(term_set), len(signature), 6))
            if score > best_score:
                best_index = index
                best_score = score
        if best_index >= 0 and best_score >= 0.34:
            cluster = clusters[best_index]
        else:
            cluster = {"rows": [], "terms": Counter()}
            clusters.append(cluster)
        cluster["rows"].append(row)
        cluster["terms"].update(terms)

    clusters.sort(
        key=lambda cluster: (
            -len(cluster["rows"]),
            -len({row.get("pmid", "") for row in cluster["rows"] if row.get("pmid", "")}),
        )
    )
    if len(clusters) > limit:
        retained = clusters[: max(0, limit - 1)]
        merged = {"rows": [], "terms": Counter()}
        for cluster in clusters[max(0, limit - 1) :]:
            merged["rows"].extend(cluster["rows"])
            merged["terms"].update(cluster["terms"])
        clusters = retained + [merged]

    groups: List[Tuple[str, List[Dict[str, str]]]] = []
    used_labels = set()
    for index, cluster in enumerate(clusters, 1):
        label_terms = [term for term, _ in cluster["terms"].most_common(3)]
        label = " / ".join(label_terms) if label_terms else f"Sentence-evidence cluster {index}"
        if label in used_labels:
            label = f"{label} ({index})"
        used_labels.add(label)
        groups.append((label, cluster["rows"]))
    return groups


def counter_dict(counter: Counter, limit: int = 12) -> Dict[str, int]:
    return {str(key): int(value) for key, value in counter.most_common(limit)}


def build_mechanism_evidence_bundle(
    gene: str,
    traits: List[Dict[str, str]],
    evidence: List[Dict[str, str]],
    varieties: List[Dict[str, str]],
) -> Dict[str, Any]:
    topics = []
    for label, rows in data_driven_topic_groups(evidence, gene=gene):
        topics.append({
            "topic_label": label,
            "evidence_records": len(rows),
            "unique_pmids": len({row.get("pmid", "") for row in rows if row.get("pmid", "")}),
            "confidence_distribution": counter_dict(Counter(row.get("confidence", "") or "Unspecified" for row in rows)),
            "evidence_code_distribution": counter_dict(split_counter(rows, "evidence_code"), 8),
            "source_distribution": counter_dict(split_counter(rows, "source_db"), 8),
            "top_text_terms": extract_keywords_from_rows(rows, 12, exclude=[gene]),
            "top_traits": top_traits_for_rows(rows, traits, 8),
            "top_pmids": collect_pmids(rows, 10),
            "top_journals": counter_dict(Counter(row.get("journal", "").strip() for row in rows if row.get("journal", "").strip()), 8),
            "representative_sentences": representative_evidence_records(rows, 8),
        })

    return {
        "gene": gene,
        "generated": date.today().isoformat(),
        "record_counts": {
            "traits": len(traits),
            "sentence_evidence": len(evidence),
            "unique_pmids": len({row.get("pmid", "") for row in evidence if row.get("pmid", "")}),
            "varieties": len(varieties),
        },
        "global_distributions": {
            "confidence": counter_dict(Counter(row.get("confidence", "") or "Unspecified" for row in evidence)),
            "ontology": counter_dict(Counter(row.get("ontology_type", "") or "Unspecified" for row in evidence)),
            "evidence_code": counter_dict(split_counter(evidence, "evidence_code")),
            "source": counter_dict(split_counter(evidence, "source_db")),
            "journal": counter_dict(Counter(row.get("journal", "").strip() for row in evidence if row.get("journal", "").strip())),
        },
        "global_text_terms": extract_keywords_from_rows(evidence, 30, exclude=[gene]),
        "topics": topics,
        "instruction": (
            "Use this bundle together with normalized_evidence.csv to write Section 6. "
            "Topic labels and terms are data-derived from the current gene only. "
            "Inspect the complete relevant sentence-evidence table rather than relying only on representative sentences. "
            "Induce mechanism themes from the current evidence, synthesize across PMIDs in review-style prose, "
            "distinguish direct evidence from integrative inference, and discuss contradictions, context dependence, and missing links. "
            "Do not reuse fixed BPH, sd1, GA, or any other gene-specific mechanism templates unless those mechanisms are explicitly supported by the returned RiceMind sentences."
        ),
    }


def build_mechanism_prompt_markdown(gene: str, bundle: Dict[str, Any], evidence_csv_name: str, language: str) -> str:
    is_zh = language.lower().startswith("zh")
    lines = [
        f"# RiceMind mechanism synthesis brief: {gene}",
        "",
        zh(
            is_zh,
            f"请基于 `{evidence_csv_name}` 的全量 RiceMind sentence evidence 和本文件中的压缩证据包，撰写报告第 6 节的个性化机制综述 Markdown。",
            f"Write the personalized Section 6 mechanism synthesis Markdown from the full RiceMind sentence evidence in `{evidence_csv_name}` and the compact evidence bundle below.",
        ),
        "",
        zh(
            is_zh,
            "硬性要求：审阅完整的相关 evidence CSV，不要使用预设的 BPH、sd1、GA 或其他基因机制模板；标题和机制主题必须从当前基因的句子证据中归纳；用连贯的综述式段落综合多个 PMID 的证据；区分直接证据与综合推断；说明矛盾证据、情境依赖和缺失环节；每个机制判断后立即用 `[PMID1, PMID2]` 标注支持该判断的 RiceMind PMID。",
            "Hard requirements: inspect the full relevant evidence CSV; do not use preset BPH, sd1, GA, or other gene-specific mechanism templates; induce headings and mechanism topics from this gene's sentence evidence; write connected review-style paragraphs that synthesize across PMIDs; distinguish direct evidence from integrative inference; report contradictions, context dependence, and missing links; cite supporting RiceMind PMIDs immediately after each mechanistic claim using `[PMID1, PMID2]`.",
        ),
        "",
        "## Compact evidence topics",
        "",
    ]
    for idx, topic in enumerate(bundle.get("topics", []), 1):
        lines.extend([
            f"### Topic {idx}: {topic.get('topic_label', '')}",
            f"- records: {topic.get('evidence_records', 0)}",
            f"- unique_pmids: {topic.get('unique_pmids', 0)}",
            f"- confidence: {json.dumps(topic.get('confidence_distribution', {}), ensure_ascii=False)}",
            f"- top_terms: {', '.join(topic.get('top_text_terms', []))}",
            f"- top_pmids: {', '.join(topic.get('top_pmids', []))}",
            "- representative_sentences:",
        ])
        for item in topic.get("representative_sentences", [])[:5]:
            source = ", ".join(part for part in [item.get("pmid", ""), item.get("year", ""), item.get("confidence", "")] if part)
            lines.append(f"  - ({source}) {truncate(item.get('sentence', ''), 360)}")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def add_markdown_paragraph(doc: Any, text: str) -> None:
    cleaned = " ".join(text.split())
    if cleaned:
        doc.add_paragraph(cleaned)


def insert_mechanism_markdown(doc: Any, mechanism_md: Path) -> bool:
    if not mechanism_md.exists():
        return False
    text = mechanism_md.read_text(encoding="utf-8")
    buffer: List[str] = []

    def flush() -> None:
        if buffer:
            add_markdown_paragraph(doc, " ".join(buffer))
            buffer.clear()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        if line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), 3)
        elif line.startswith("## "):
            flush()
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith("# "):
            flush()
            title = line[2:].strip()
            if title:
                doc.add_heading(title, 2)
        elif line.startswith("- "):
            flush()
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            buffer.append(line)
    flush()
    return True


def add_mechanism_synthesis(
    doc: Any,
    gene: str,
    traits: List[Dict[str, str]],
    evidence: List[Dict[str, str]],
    is_zh: bool,
    mechanism_md: Optional[Path] = None,
) -> None:
    if not evidence:
        doc.add_paragraph(
            zh(
                is_zh,
                "输入数据未提供句子级证据；不能生成 PMID 支撑的机制综述。完整句子证据应通过 normalized_evidence.csv 保存。",
                "No sentence-level evidence was available; PMID-supported mechanism synthesis cannot be generated. Full sentence evidence should be preserved in normalized_evidence.csv.",
            )
        )
        return

    if not mechanism_md or not insert_mechanism_markdown(doc, mechanism_md):
        raise RuntimeError(
            "Sentence evidence is available, but a readable --mechanism-md was not supplied. "
            "Generate the personalized PMID-backed mechanism synthesis from the normalized evidence first."
        )

def build_docx(
    gene: str,
    bundle: Dict[str, Any],
    traits: List[Dict[str, str]],
    evidence: List[Dict[str, str]],
    varieties: List[Dict[str, str]],
    out_path: Path,
    fig_paths: Dict[str, Optional[Path]],
    language: str,
    max_evidence_rows: int,
    max_trait_rows: int,
    mechanism_md: Optional[Path] = None,
) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx") from exc

    is_zh = language.lower().startswith("zh")
    profile = choose_profile(bundle)
    doc = Document()
    configure_document_styles(doc)
    doc.add_heading(zh(is_zh, f"{gene} 的 RiceMind 基因机制证据报告", f"RiceMind Gene Mechanism Evidence Report: {gene}"), 0)
    doc.add_paragraph(zh(is_zh, f"生成日期：{date.today().isoformat()}", f"Generated: {date.today().isoformat()}"))
    doc.add_paragraph(
        zh(
            is_zh,
            "数据来源：RiceMind REST API / MCP payload。本文档采用固定模板生成；所有机制性描述均限制在 RiceMind 返回的基因名片、GTA 元信息、证据代码、PMID、句子证据和上下文范围内。",
            "Data source: RiceMind REST API / MCP payload. This document follows the fixed report template; all mechanistic statements are restricted to RiceMind gene profile data, GTA metadata, evidence codes, PMIDs, sentence evidence and retrieved context.",
        )
    )

    doc.add_heading(zh(is_zh, "1. 数据来源、检索范围与完整性", "1. Data Source, Retrieval Scope and Completeness"), 1)
    doc.add_paragraph(
        zh(
            is_zh,
            f"API base URL：{bundle.get('api_base', 'input JSON bundle')}。本次规范化得到 {len(traits)} 条 gene-trait 关联、{len(evidence)} 条句子级证据、{len(varieties)} 个共现品种记录。",
            f"API base URL: {bundle.get('api_base', 'input JSON bundle')}. Normalization yielded {len(traits)} gene-trait associations, {len(evidence)} sentence-level evidence records and {len(varieties)} co-reported variety records.",
        )
    )
    api_calls = compact_api_calls(bundle.get("api_calls") or [])
    if api_calls:
        add_table(doc, api_call_rows(api_calls, REPORT_API_ROWS))
        doc.add_paragraph(
            zh(
                is_zh,
                f"上表仅展示最多 {REPORT_API_ROWS} 类功能不同的 API 调用摘要；分页访问已按 endpoint 和查询功能聚合，不逐页罗列。完整聚合后的调用摘要与原始返回数据保存在 payload JSON sidecar 中。",
                f"The table shows at most {REPORT_API_ROWS} functionally distinct API call summaries. Paginated calls are aggregated by endpoint and query function rather than listed page by page. The compact call summary and raw returned data are preserved in the payload JSON sidecar.",
            )
        )
    error_rows = endpoint_error_rows(bundle)
    if error_rows:
        doc.add_paragraph(
            zh(
                is_zh,
                "以下 endpoint 在本次检索中返回错误或不可用；报告继续使用其他可用 RiceMind 数据生成，并在 payload JSON 中保留错误信息。",
                "The following endpoints returned errors or were unavailable in this retrieval. The report continues with other available RiceMind data and preserves error details in the payload JSON.",
            )
        )
        add_table(doc, error_rows)

    doc.add_heading(zh(is_zh, "2. 基因概貌与外部链接", "2. Gene Overview and External Links"), 1)
    if profile:
        add_table(doc, flatten_profile_rows(profile))
    else:
        doc.add_paragraph(zh(is_zh, "未在输入数据中找到 gene-profile 记录。", "No gene-profile record was found in the input data."))

    doc.add_heading(zh(is_zh, "3. 全量句证输出与本体分布", "3. Evidence Distribution and Sentence Provenance"), 1)
    add_table(doc, evidence_summary_rows(traits, evidence))
    if traits or evidence:
        doc.add_heading(zh(is_zh, "本体维度分布", "Ontology Mapping Distribution"), 2)
        add_table(doc, ontology_distribution_rows(traits, evidence))
    if evidence:
        doc.add_heading(zh(is_zh, "代表性句子证据", "Representative Sentence Evidence"), 2)
        doc.add_paragraph(
            zh(
                is_zh,
                "下表按置信层级抽取代表性 RiceMind 句子证据，仅用于展示证据形态；全量句证保存在 normalized_evidence.csv。",
                "The table samples representative RiceMind sentence evidence by confidence tier only to show evidence shape; full sentence evidence is preserved in normalized_evidence.csv.",
            )
        )
        add_table(doc, representative_sentence_rows(evidence))

    doc.add_heading(zh(is_zh, "4. 全量 GTA 概貌", "4. Full GTA Landscape"), 1)
    add_picture_if_exists(doc, fig_paths.get("confidence"), zh(is_zh, "图 1. RiceMind confidence tier 分布。", "Figure 1. RiceMind confidence-tier distribution."))
    add_picture_if_exists(doc, fig_paths.get("ontology"), zh(is_zh, "图 2. Ontology 类型分布。", "Figure 2. Ontology type distribution."))
    add_picture_if_exists(doc, fig_paths.get("top_traits"), zh(is_zh, "图 3. 按文献支持数排序的主要 trait。", "Figure 3. Top traits ranked by supporting article count."))
    sorted_traits = sorted(
        traits,
        key=lambda r: ({"High": 0, "Medium": 1, "Low": 2}.get(r["confidence"], 3), -parse_int(r["support"]), r["trait"]),
    )
    trait_rows = [["Trait", "Ontology", "Confidence", "Evidence code", "Source", "Support", "Earliest year"]]
    trait_display_limit = max_trait_rows if max_trait_rows and max_trait_rows > 0 else REPORT_TRAIT_ROWS
    for row in row_limit(sorted_traits, trait_display_limit):
        trait_rows.append([
            truncate(row["trait"], 90),
            " ".join(x for x in [row["ontology_type"], row["ontology_id"]] if x),
            row["confidence"],
            truncate(row["evidence_code"], 80),
            truncate(row["source_db"], 80),
            row["support"],
            row["year"],
        ])
    add_table(doc, trait_rows)
    if len(sorted_traits) > trait_display_limit:
        doc.add_paragraph(
            zh(
                is_zh,
                f"报告正文仅展示支持度和置信层级靠前的 {trait_display_limit} 条 trait 记录；全部 {len(sorted_traits)} 条 trait 关联已保存到 normalized_traits.csv。",
                f"The report body shows only the top {trait_display_limit} trait records by confidence/support; all {len(sorted_traits)} trait associations are saved in normalized_traits.csv.",
            )
        )

    doc.add_heading(zh(is_zh, "5. 证据代码、来源与置信层级统计", "5. Evidence Codes, Sources and Confidence Statistics"), 1)
    add_picture_if_exists(doc, fig_paths.get("evidence_code"), zh(is_zh, "图 4. Evidence code 分布。", "Figure 4. Evidence-code distribution."))
    add_picture_if_exists(doc, fig_paths.get("source"), zh(is_zh, "图 5. RiceMind 来源数据库分布。", "Figure 5. RiceMind source database distribution."))
    summary_rows = [["Statistic", "Value"]]
    summary_rows.extend([
        ["High traits", str(sum(1 for row in traits if row["confidence"] == "High"))],
        ["Medium traits", str(sum(1 for row in traits if row["confidence"] == "Medium"))],
        ["Low traits", str(sum(1 for row in traits if row["confidence"] == "Low"))],
        ["Unique PMIDs", str(len({row["pmid"] for row in evidence if row["pmid"]}))],
        ["Evidence sentences", str(len(evidence))],
    ])
    add_table(doc, summary_rows)

    doc.add_heading(zh(is_zh, "6. RiceMind 句子证据驱动的机制综述", "6. RiceMind Sentence-Evidence-Driven Mechanism Synthesis"), 1)
    add_mechanism_synthesis(doc, gene, traits, evidence, is_zh, mechanism_md)

    doc.add_heading(zh(is_zh, "7. 研究热点与年代变迁分析", "7. Temporal Analysis and Hotspot Evolution"), 1)
    temporal_rows = temporal_hotspot_rows(evidence, is_zh)
    if len(temporal_rows) > 1:
        add_table(doc, temporal_rows)
        doc.add_paragraph(
            zh(
                is_zh,
                "阶段划分由 RiceMind 句证年份自动分箱得到；热点词来自各阶段句子文本，因此反映研究关注度变化，不等同于生物学重要性排序。",
                "Phase boundaries are computed from RiceMind sentence years; hotspot terms come from phase-specific sentence text and therefore reflect attention shifts rather than biological-importance ranking.",
            )
        )
    else:
        doc.add_paragraph(zh(is_zh, "输入证据缺少年份信息，无法生成热点阶段分析。", "No usable publication years were available for hotspot phase analysis."))

    doc.add_heading(zh(is_zh, "8. 二级文献计量与 PMID 可追溯性", "8. Secondary Bibliometrics and PMID Traceability"), 1)
    add_picture_if_exists(doc, fig_paths.get("years"), zh(is_zh, "图 6. 句子级证据的发表年份分布。", "Figure 6. Publication-year distribution of sentence-level evidence."))
    if not fig_paths.get("years"):
        doc.add_paragraph(zh(is_zh, "输入证据缺少可用发表年份，未生成年度趋势图。", "No usable publication years were available for a trend chart."))
    journal_rows, pmid_rows = bibliometric_rows(evidence)
    if journal_rows:
        doc.add_heading(zh(is_zh, "Top 发表刊物", "Top Journals"), 2)
        add_table(doc, journal_rows)
        add_picture_if_exists(doc, fig_paths.get("journal"), zh(is_zh, "图 7. 句子级证据的期刊分布。", "Figure 7. Journal distribution of sentence-level evidence."))
    else:
        doc.add_paragraph(
            zh(
                is_zh,
                "当前 RiceMind 句证未提供可稳定统计的 journal 字段，因此不输出 Top 期刊排名。",
                "The current RiceMind sentence evidence does not provide stable journal metadata, so top-journal ranking is not reported.",
            )
        )
    if pmid_rows:
        doc.add_heading(zh(is_zh, "高频 PMID / 文献簇", "High-Frequency PMID / Article Clusters"), 2)
        add_table(doc, pmid_rows)
    doc.add_paragraph(
        zh(
            is_zh,
            "机构地理热点只有在 RiceMind 返回通讯作者单位或 affiliation 元数据时才可报告；若 payload 中没有该字段，报告不得凭常识或外部记忆补写机构地点。",
            "Institutional geography is reported only when RiceMind returns corresponding-author or affiliation metadata; if absent from the payload, the report must not fill locations from general knowledge or memory.",
        )
    )

    doc.add_heading(zh(is_zh, "9. 品种共现与组学序列信息", "9. Variety Co-Occurrence and Omics Sequence Information"), 1)
    if varieties:
        variety_rows = [["Variety"]]
        for row in varieties:
            variety_rows.append([row["variety"]])
        add_table(doc, variety_rows)
    else:
        doc.add_paragraph(zh(is_zh, "未检索到或未提供基因-品种共现记录。", "No gene-variety co-occurrence records were retrieved or provided."))
    sequence_rows = sequence_api_url_rows(bundle, profile, gene)
    if sequence_rows:
        doc.add_paragraph(
            zh(
                is_zh,
                "序列相关信息不在正文中展开为长表格；请直接追溯下列 RiceMind gene-omics-sequence API 超链接。完整接口返回仍保存在 payload JSON sidecar 中。",
                "Sequence information is not expanded into a long table in the report body. Use the following RiceMind gene-omics-sequence API hyperlink for direct tracing. The full endpoint response remains in the payload JSON sidecar.",
            )
        )
        add_table(doc, sequence_rows)
    elif bundle.get("omics_sequence_error"):
        doc.add_paragraph(zh(is_zh, f"组学序列接口未返回可用结果：{bundle['omics_sequence_error']}", f"Omics sequence endpoint did not return usable data: {bundle['omics_sequence_error']}"))

    doc.add_heading(zh(is_zh, "10. 证据边界与解释限制", "10. Evidence Boundaries and Interpretation Limits"), 1)
    doc.add_paragraph(
        zh(
            is_zh,
            "High/Tier 1 适合较保守的机制验证；Medium/Tier 2 更适合候选机制发现；Low/Tier 3 和 NLP-only 共现只能作为探索性线索。句子共现本身不能证明因果关系；报告中的机制表述必须回溯到本报告列出的 PMID、sentence ID 和 RiceMind evidence code。",
            "High/Tier 1 evidence is appropriate for conservative mechanism validation; Medium/Tier 2 is better treated as candidate discovery; Low/Tier 3 and NLP-only co-occurrence are exploratory. Sentence co-occurrence alone does not establish causality; mechanistic statements must trace back to the PMIDs, sentence IDs and RiceMind evidence codes listed in this report.",
        )
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    apply_document_fonts(doc)
    doc.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--gene", required=True, help="Target rice gene symbol, alias, or RAP ID")
    parser.add_argument("--input-json", type=Path, help="RiceMind API/MCP payload bundle")
    parser.add_argument("--api-base", default=DEFAULT_API_BASE, help=f"RiceMind REST API base URL, default: {DEFAULT_API_BASE}")
    parser.add_argument("--no-api", action="store_true", help="Do not call the live API; use only --input-json")
    parser.add_argument("--endpoint-map", type=Path, help="Optional JSON map overriding REST endpoint paths")
    parser.add_argument("--out", type=Path, help="Output DOCX path")
    parser.add_argument("--fig-dir", type=Path, help="Directory for generated figures")
    parser.add_argument("--language", default="zh", choices=["en", "zh"], help="Report language")
    parser.add_argument("--page-limit", type=int, default=DEFAULT_PAGE_LIMIT, help="API page size used for full pagination")
    parser.add_argument("--max-pages", type=int, default=DEFAULT_MAX_PAGES, help="Safety cap for paginated API calls")
    parser.add_argument("--skip-trait-evidence", action="store_true", help="Skip per-trait search-by-trait-and-gene calls")
    parser.add_argument("--max-evidence-rows", type=int, default=0, help="Retained for backward compatibility; full evidence is written to CSV, not the DOCX body")
    parser.add_argument("--max-trait-rows", type=int, default=REPORT_TRAIT_ROWS, help=f"Max trait rows in DOCX summary table; default {REPORT_TRAIT_ROWS}")
    parser.add_argument("--mechanism-md", type=Path, help="Personalized PMID-backed Section 6 mechanism synthesis Markdown generated from this run's RiceMind evidence")
    parser.add_argument("--sidecars-only", action="store_true", help="Write payload/CSV/mechanism bundle/figures only; skip DOCX generation")
    parser.add_argument("--docx-only", action="store_true", help="Write only the DOCX and figures; do not rewrite payload/CSV/mechanism sidecars")
    parser.add_argument("--write-brief", action="store_true", help="Write the mechanism synthesis brief even when --mechanism-md is already supplied")
    args = parser.parse_args()

    if args.no_api and not args.input_json:
        parser.error("Provide --input-json when using --no-api")
    if args.sidecars_only and args.docx_only:
        parser.error("Use at most one of --sidecars-only and --docx-only")
    if args.mechanism_md and (
        not args.mechanism_md.is_file()
        or not args.mechanism_md.read_text(encoding="utf-8").strip()
    ):
        parser.error("--mechanism-md must point to a readable, non-empty Markdown file")

    bundle: Dict[str, Any] = {"gene": args.gene}
    if args.input_json:
        bundle.update(read_json(args.input_json))
    if not args.no_api:
        fetched = fetch_from_api(
            args.api_base,
            args.gene,
            args.endpoint_map,
            args.page_limit,
            args.max_pages,
            fetch_trait_evidence=not args.skip_trait_evidence,
        )
        bundle.update(fetched)

    out_path = args.out or Path(f"{safe_filename(args.gene)}_RiceMind_gene_mechanism_report.docx")
    data_dir = out_path.parent / f"{out_path.stem}_data"
    fig_dir = args.fig_dir or data_dir / "figures"

    traits = normalize_traits(bundle)
    evidence = normalize_evidence(bundle, traits)
    varieties = normalize_varieties(bundle)
    bundle["api_calls"] = compact_api_calls(bundle.get("api_calls") or [])

    traits_csv = data_dir / f"{out_path.stem}_normalized_traits.csv"
    evidence_csv = data_dir / f"{out_path.stem}_normalized_evidence.csv"
    varieties_csv = data_dir / f"{out_path.stem}_normalized_varieties.csv"
    payload_json = data_dir / f"{out_path.stem}_payload.json"
    mechanism_bundle_json = data_dir / f"{out_path.stem}_mechanism_evidence_bundle.json"
    mechanism_brief_md = data_dir / f"{out_path.stem}_mechanism_synthesis_brief.md"

    wrote_sidecars: List[Path] = []
    if not args.docx_only:
        if write_csv(
            traits_csv,
            traits,
            ["gene", "rap_id", "trait", "ontology_type", "ontology_id", "evidence_code", "source_db", "confidence", "support", "year"],
        ):
            wrote_sidecars.append(traits_csv)
        if write_csv(
            evidence_csv,
            evidence,
            ["trait", "ontology_type", "ontology_id", "evidence_code", "source_db", "confidence", "pmid", "sentence_id", "year", "title", "journal", "doi", "sentence"],
        ):
            wrote_sidecars.append(evidence_csv)
        if write_csv(varieties_csv, varieties, ["variety"]):
            wrote_sidecars.append(varieties_csv)
        if bundle_has_usable_data(bundle):
            write_json(payload_json, bundle)
            wrote_sidecars.append(payload_json)
        else:
            remove_output_file(payload_json)

        mechanism_bundle = build_mechanism_evidence_bundle(args.gene, traits, evidence, varieties)
        if evidence:
            write_json(mechanism_bundle_json, mechanism_bundle)
            wrote_sidecars.append(mechanism_bundle_json)
        else:
            remove_output_file(mechanism_bundle_json)
        if evidence and (args.write_brief or not args.mechanism_md):
            write_text(mechanism_brief_md, build_mechanism_prompt_markdown(args.gene, mechanism_bundle, evidence_csv.name, args.language))
            wrote_sidecars.append(mechanism_brief_md)
        else:
            remove_output_file(mechanism_brief_md)

    if evidence and not args.mechanism_md and not args.sidecars_only:
        if args.docx_only:
            next_step = (
                "Rerun without --docx-only (preferably with --sidecars-only) to create the normalized evidence "
                "and mechanism synthesis brief, then supply the completed Markdown with --mechanism-md."
            )
        else:
            next_step = (
                f"Use {mechanism_brief_md} and {evidence_csv} to write the mechanism review, "
                "then rerun with --mechanism-md."
            )
        print(
            "DOCX was not generated because sentence evidence requires a personalized --mechanism-md synthesis. "
            + next_step,
            file=sys.stderr,
        )
        for path in wrote_sidecars:
            print(f"Wrote {path}")
        remove_empty_directory(fig_dir)
        remove_empty_directory(data_dir)
        return 3

    figure_targets = {
        "confidence": fig_dir / "confidence_distribution.png",
        "ontology": fig_dir / "ontology_distribution.png",
        "top_traits": fig_dir / "top_traits_by_support.png",
        "evidence_code": fig_dir / "evidence_code_distribution.png",
        "source": fig_dir / "source_distribution.png",
        "years": fig_dir / "publication_year_trend.png",
        "journal": fig_dir / "journal_distribution.png",
    }
    for figure_path in figure_targets.values():
        remove_output_file(figure_path)
    fig_paths = {
        "confidence": plot_counter(Counter(row["confidence"] for row in traits if row["confidence"]), "Confidence-tier distribution", figure_targets["confidence"]),
        "ontology": plot_counter(Counter(row["ontology_type"] for row in traits if row["ontology_type"]), "Ontology distribution", figure_targets["ontology"]),
        "top_traits": plot_top_traits(traits, figure_targets["top_traits"]),
        "evidence_code": plot_counter(split_counter(traits + evidence, "evidence_code"), "Evidence-code distribution", figure_targets["evidence_code"]),
        "source": plot_counter(split_counter(traits + evidence, "source_db"), "Source database distribution", figure_targets["source"]),
        "years": plot_years((row["year"] for row in evidence), figure_targets["years"]),
        "journal": plot_counter(Counter(row["journal"] for row in evidence if row["journal"]), "Journal distribution", figure_targets["journal"]),
    }

    if not args.sidecars_only:
        try:
            out_path.parent.mkdir(parents=True, exist_ok=True)
            build_docx(args.gene, bundle, traits, evidence, varieties, out_path, fig_paths, args.language, args.max_evidence_rows, args.max_trait_rows, args.mechanism_md)
        except RuntimeError as exc:
            print(str(exc), file=sys.stderr)
            remove_empty_directory(fig_dir)
            remove_empty_directory(data_dir)
            return 2
        print(f"Wrote {out_path}")
    for path in wrote_sidecars:
        print(f"Wrote {path}")
    remove_empty_directory(fig_dir)
    remove_empty_directory(data_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
